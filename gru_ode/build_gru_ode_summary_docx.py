"""
Build a Word document summarizing the GRU-ODE: formulation and hyperparameters.

Two model variants are documented:
  1. Pooled 500 ms GRU-ODE  (gru_ode_pooled_by_region.py)
  2. 10 ms Poisson GRU-ODE  (gru_ode_10ms.py)

Output: data/GRU_ODE_summary.docx
"""
from pathlib import Path
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

REPO = Path(__file__).resolve().parent
OUT = REPO / 'data' / 'GRU_ODE_summary.docx'
OUT.parent.mkdir(parents=True, exist_ok=True)


def set_font(run, name='Calibri', size=11, bold=False, italic=False):
    run.font.name = name
    run.font.size = Pt(size)
    run.font.bold = bold
    run.font.italic = italic


def add_heading(doc, text, level=1):
    h = doc.add_heading('', level=level)
    r = h.add_run(text)
    r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, italic=False, size=11):
    p = doc.add_paragraph()
    r = p.add_run(text)
    set_font(r, size=size, bold=bold, italic=italic)
    return p


def add_bullets(doc, items):
    for it in items:
        p = doc.add_paragraph(style='List Bullet')
        r = p.add_run(it)
        set_font(r)


def add_code(doc, code):
    p = doc.add_paragraph()
    r = p.add_run(code)
    set_font(r, name='Consolas', size=10)


def add_table(doc, header, rows, widths_in=None):
    t = doc.add_table(rows=1 + len(rows), cols=len(header))
    t.style = 'Light Grid Accent 1'
    for j, h in enumerate(header):
        c = t.rows[0].cells[j]
        c.text = ''
        p = c.paragraphs[0]
        r = p.add_run(h)
        set_font(r, size=10, bold=True)
    for i, row in enumerate(rows):
        for j, val in enumerate(row):
            c = t.rows[i + 1].cells[j]
            c.text = ''
            p = c.paragraphs[0]
            r = p.add_run(str(val))
            set_font(r, size=10)
    if widths_in:
        for row in t.rows:
            for j, w in enumerate(widths_in):
                row.cells[j].width = Inches(w)
    return t


def main():
    doc = Document()

    # ---- Style defaults
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(11)

    # ---- Title
    title = doc.add_heading('', 0)
    r = title.add_run('GRU-ODE: Formulation and Hyperparameters')
    r.font.color.rgb = RGBColor(0x0B, 0x2A, 0x4D)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    rs = sub.add_run('Neuropixels Foraging Project — Aim 1 stepping stone for graph neural ODE')
    set_font(rs, italic=True, size=11)

    # ---- Section 1: Overview
    add_heading(doc, '1. Overview', 1)
    add_para(doc,
        'The Gated Neural ODE (GRU-ODE, also called "Gated Neural ODE") replaces the '
        'discrete-time GRU with a continuous-time ordinary differential equation that '
        'uses GRU-style gating to control the rate of change of the hidden state. '
        'Between observations the hidden state evolves smoothly via the ODE; when a new '
        'observation arrives, the hidden state is updated discretely via a standard GRU cell.')
    add_para(doc, 'Two components are explicitly separated:', bold=True)
    add_bullets(doc, [
        'Continuous dynamics (ODE): how the hidden state evolves between observations '
        '— the intrinsic dynamics of the neural population.',
        'Observation updates (discrete jumps): how each new neural observation corrects '
        'or redirects the trajectory.',
        'Prediction: the model forecasts the next time step using only the ODE — no '
        'observation — so prediction quality measures how autonomous the learned '
        'dynamics are.'
    ])
    add_para(doc, 'Reference: De Brouwer et al. 2019 (NeurIPS), "GRU-ODE-Bayes". ODE '
                  'integration uses torchdiffeq (Chen et al. 2018).')

    # ---- Section 2: Mathematical Formulation
    add_heading(doc, '2. Mathematical Formulation', 1)

    add_heading(doc, '2.1. Discrete GRU (reference)', 2)
    add_para(doc, 'The standard GRU updates the hidden state at each discrete time step k:')
    add_code(doc,
        'r_k = sigmoid(W_xr · x_k + W_hr · h_{k-1} + b_r)         (reset gate)\n'
        'z_k = sigmoid(W_xz · x_k + W_hz · h_{k-1} + b_z)         (update gate)\n'
        'n_k = tanh   (W_xn · x_k + r_k ⊙ (W_hn · h_{k-1} + b_hn)) (candidate)\n'
        'h_k = (1 − z_k) ⊙ n_k + z_k ⊙ h_{k-1}                    (hidden update)')
    add_para(doc, 'Here x_k is the input at step k, h_{k-1} is the previous hidden state, '
                  'and ⊙ denotes element-wise multiplication. Rewriting the hidden update '
                  'as a finite difference exposes the continuous-time analogue:')
    add_code(doc, 'h_k − h_{k-1} = (1 − z_k) ⊙ (n_k − h_{k-1})')

    add_heading(doc, '2.2. GRU-ODE (continuous-time dynamics)', 2)
    add_para(doc, 'Replacing the finite difference with a derivative gives the ODE that '
                  'governs the hidden state between observations:')
    add_code(doc,
        'dh/dt = (1 − z(h)) ⊙ (n(h) − h)\n\n'
        'with\n'
        '   z(h) = sigmoid(W_z2 · tanh(W_z1 · h + b_z1) + b_z2)         (update gate)\n'
        '   n(h) = tanh   (W_n2 · tanh(W_n1 · h + b_n1) + b_n2)         (candidate)')
    add_para(doc, 'Both gate networks are 2-layer MLPs with hidden width 64 (parameter '
                  'ODE_GATE_HIDDEN). Between observations there is no input x — the '
                  'gates depend only on the current hidden state h.')
    add_para(doc, 'Interpretation:', bold=True)
    add_bullets(doc, [
        'z(h) is the rate-of-change gate. Values near 1 → "stay where you are" (slow '
        'dynamics); values near 0 → "move toward the candidate" (fast dynamics).',
        'n(h) is the candidate state — the direction the dynamics point toward.',
        '(1 − z(h)) ⊙ (n(h) − h) is the velocity field of the hidden state.'
    ])

    add_heading(doc, '2.3. Observation update (discrete jump)', 2)
    add_para(doc, 'When a new observation x_k arrives, the hidden state is updated by a '
                  'standard nn.GRUCell operating in the shared latent space:')
    add_code(doc,
        'x_proj = W_input[session] · x_k + b_input[session]    (project to D_shared)\n\n'
        'r = sigmoid(W_xr · x_proj + W_hr · h + b_r)\n'
        'z = sigmoid(W_xz · x_proj + W_hz · h + b_z)\n'
        'n = tanh   (W_xn · x_proj + r ⊙ (W_hn · h + b_hn))\n'
        'h_new = (1 − z) ⊙ n + z ⊙ h')
    add_para(doc, 'Each session has its own input projection W_input[session] (which '
                  'maps the session-specific neuron count to the shared dimensionality '
                  'D_shared = 32) so that one shared ODE / GRUCell can serve all '
                  'sessions simultaneously.')

    add_heading(doc, '2.4. Full forward pass', 2)
    add_para(doc, 'For an input sequence of SEQ_LEN observations:')
    add_code(doc,
        'h ← 0   (zeros, shape (batch, hidden_size=32))\n\n'
        'For k = 1 … SEQ_LEN:\n'
        '    h_evolved ← ODESolve(dh/dt = (1−z(h))⊙(n(h)−h), h, t=[0, dt])\n'
        '    x_proj    ← Linear[session](x[:, k, :])\n'
        '    h         ← GRUCell(h_evolved, x_proj)\n\n'
        '# Prediction step: pure ODE, NO observation\n'
        'h_final ← ODESolve(dh/dt = (1−z(h))⊙(n(h)−h), h, t=[0, dt])\n\n'
        'shared_out ← Linear_shared(h_final)\n'
        'pred       ← Linear[session](shared_out)')
    add_para(doc, 'The prediction step uses ONLY the ODE — this is what makes prediction '
                  'quality a clean probe of how autonomous the learned dynamics are.')

    # ---- Section 3: Architecture
    add_heading(doc, '3. Architecture', 1)
    add_para(doc, 'The PooledGRUODE module groups all sessions of a region under a single '
                  'shared core. Session identity selects per-session input/output '
                  'projections:')
    add_code(doc,
        'PooledGRUODE\n'
        '  ├─ input_projections : ModuleDict\n'
        '  │     session_i: Linear(N_neurons_i → 32)\n'
        '  │\n'
        '  ├─ ode_func : GRUODEFunc                # continuous dynamics\n'
        '  │     update_gate: Linear(32→64) → Tanh → Linear(64→32) → Sigmoid\n'
        '  │     candidate  : Linear(32→64) → Tanh → Linear(64→32) → Tanh\n'
        '  │\n'
        '  ├─ obs_cell : nn.GRUCell(input=32, hidden=32)\n'
        '  │\n'
        '  ├─ fc_shared : Linear(32 → 32)           # shared decoder\n'
        '  │\n'
        '  └─ output_projections : ModuleDict\n'
        '        session_i: Linear(32 → N_neurons_i)')
    add_para(doc, 'For an LHA session with ~45 neurons:')
    add_table(doc, ['Component', 'Discrete GRU', 'GRU-ODE'],
        [
            ['Input projection', '1,472', '1,472'],
            ['Core dynamics', '6,336 (nn.GRU)', '4,192 (z) + 4,192 (n) = 8,384 (ODEFunc)'],
            ['Observation update', '— (in nn.GRU)', '6,336 (nn.GRUCell)'],
            ['Shared decoder', '1,056', '1,056'],
            ['Output projection', '1,485', '1,485'],
            ['Total per session', '~10,349', '~18,733'],
        ],
        widths_in=[1.8, 1.8, 2.4])
    add_para(doc, 'The GRU-ODE is roughly 2× the parameter count because it has both '
                  'continuous (ODE) and discrete (GRUCell) dynamics machinery.')

    # ---- Section 4: Hyperparameters — 500ms model
    add_heading(doc, '4. Hyperparameters', 1)
    add_heading(doc, '4.1. Pooled 500 ms GRU-ODE — gru_ode_pooled_by_region.py', 2)
    add_table(doc, ['Hyperparameter', 'Value', 'Notes'],
        [
            ['BIN_SIZE_MS', '500', 'Spike-count bin width.'],
            ['SEQ_LEN', '10', 'Input sequence length (10 × 500 ms = 5 s context).'],
            ['D_SHARED', '32', 'Shared latent space dimensionality.'],
            ['HIDDEN_SIZE', '32', 'Hidden state dimensionality.'],
            ['ODE_GATE_HIDDEN', '64', 'Internal width of the z(h) and n(h) MLPs.'],
            ['BATCH_SIZE', '64', '—'],
            ['NUM_EPOCHS', '150', 'Maximum training epochs.'],
            ['PATIENCE', '15', 'Early-stopping patience on validation loss.'],
            ['LEARNING_RATE', '1e-3', 'Adam optimizer.'],
            ['Optimizer', 'Adam', '—'],
            ['Loss', 'MSELoss', 'Targets are z-scored.'],
            ['TRAIN_FRAC', '0.8', 'Temporal split (first 80 % train, last 20 % test).'],
            ['GRAD_CLIP', '1.0', 'Max gradient norm — prevents ODE backprop blow-ups.'],
            ['ODE_SOLVER', "'rk4'", '4th-order Runge-Kutta (fixed step). ~6× faster than dopri5 in our tests.'],
            ['ODE_STEP_SIZE', '1.0', 'Single RK4 step per integration interval.'],
            ['ODE_DT', '1.0', 'Normalized time between observations (each 500 ms bin = 1 dt).'],
            ['Backprop', 'odeint (regular)', 'Adjoint not needed at hidden_size=32; faster forward.'],
            ['FS', '30000', 'Raw sampling rate (Hz); BIN_SAMPLES = BIN_SIZE_MS · FS / 1000.'],
        ],
        widths_in=[2.0, 1.4, 3.6])

    add_heading(doc, '4.2. 10 ms Poisson GRU-ODE — gru_ode_10ms.py', 2)
    add_para(doc, 'Fine-grained variant. Input is 10 ms z-scored bins; target is the raw '
                  'spike count summed over the next 100 ms (10 bins). Loss is Poisson '
                  'negative log-likelihood (count data appropriate). The ODE forecasts '
                  '10 steps ahead with no observations.')
    add_table(doc, ['Hyperparameter', 'Value', 'Notes'],
        [
            ['BIN_SIZE_MS', '10', '10 ms input bins.'],
            ['SEQ_LEN', '50', '50 × 10 ms = 500 ms context.'],
            ['PRED_WINDOW_MS', '100', 'Prediction horizon.'],
            ['PRED_BINS', '10', 'Sum next 10 raw-count bins → Poisson target.'],
            ['STRIDE', '50', 'Non-overlapping windows for test; can drop to 10 for full run.'],
            ['D_SHARED', '32', 'Shared latent space dim.'],
            ['HIDDEN_SIZE', '32', 'Hidden state dim.'],
            ['ODE_GATE_HIDDEN', '64', 'Same MLP width as 500 ms model.'],
            ['BATCH_SIZE', '64', '—'],
            ['NUM_EPOCHS', '150', '—'],
            ['PATIENCE', '20', 'Slightly longer patience than 500 ms.'],
            ['LEARNING_RATE', '1e-3', 'Adam.'],
            ['Optimizer', 'Adam', '—'],
            ['Loss', 'PoissonNLLLoss(log_input=True)', 'Model outputs log(λ) where λ = expected count.'],
            ['TRAIN_FRAC', '0.8', 'Temporal split.'],
            ['GRAD_CLIP', '1.0', '—'],
            ['ODE_SOLVER', "'rk4'", 'Same as 500 ms.'],
            ['ODE_STEP_SIZE', '1.0', '—'],
            ['ODE_DT', '1.0', 'Each 10 ms bin = 1 dt in ODE time.'],
            ['ODE calls / forward', '60', '50 input steps + 10 prediction steps.'],
        ],
        widths_in=[2.0, 1.6, 3.4])

    # ---- Section 5: ODE solver context
    add_heading(doc, '5. ODE solver — practical notes', 1)
    add_para(doc, 'The original plan called for the adaptive Dormand-Prince method '
                  '(dopri5) with rtol=1e-3, atol=1e-4 (standard defaults for neural '
                  'ODEs). In practice this was ~6× slower than fixed-step RK4 (~50 s/'
                  'epoch vs. ~8 s/epoch) at no measurable benefit in fit quality, so the '
                  'production scripts use rk4 with step_size=1.0. The dopri5 settings '
                  'are retained here because they may be useful if dynamics later '
                  'become stiff:')
    add_table(doc, ['Solver', 'When to use', 'rtol', 'atol', 'Comment'],
        [
            ['rk4 (default)', 'Smooth dynamics, fixed cost', '—', '—', 'Single 4th-order RK step per dt.'],
            ['dopri5', 'Stiff or fast-varying dynamics', '1e-3', '1e-4', 'Adaptive — slower, more accurate.'],
        ],
        widths_in=[1.4, 1.8, 0.8, 0.8, 2.2])
    add_para(doc, 'Adaptive solvers guarantee local error |err| < atol + rtol · |h|; '
                  'tighter tolerances increase the number of internal sub-steps the '
                  'solver takes (NFE — number of function evaluations). NFE per integration '
                  'step is itself a useful diagnostic: high NFE indicates stiff or '
                  'rapidly changing dynamics.')

    # ---- Section 6: Training setup
    add_heading(doc, '6. Training setup', 1)
    add_bullets(doc, [
        'Optimizer: Adam (default β1=0.9, β2=0.999, eps=1e-8 from PyTorch).',
        'Loss: MSELoss for the 500 ms model (z-scored targets); PoissonNLLLoss(log_input=True) for the 10 ms model.',
        'Early stopping: patience 15 (500 ms) / 20 (10 ms) on validation loss.',
        'Gradient clipping: max norm 1.0 (prevents ODE-backprop blow-ups).',
        'Train/test split: temporal — first 80 % train, last 20 % test (no shuffling across the boundary).',
        'Pooled architecture: one shared ode_func + obs_cell + fc_shared across all sessions of a region; each session gets its own input/output projection.',
        'Model variants per region: condition-specific Fed, condition-specific Fasted, and Combined (all sessions).',
        'Device: CUDA when available; falls back to CPU.',
    ])

    # ---- Section 7: Empirical results (concise reproduction)
    add_heading(doc, '7. Empirical highlights (reference)', 1)
    add_para(doc, '500 ms model — combined-pooled, fed vs fasted (per-session means; '
                  'Mann-Whitney U):')
    add_table(doc, ['Region', 'Metric', 'Fed', 'Fasted', 'p'],
        [
            ['LHA', 'R²', '0.033', '0.035', '0.686'],
            ['LHA', 'PR', '9.74', '6.48', '0.114'],
            ['LHA', 'Variance', '0.139', '0.113', '0.029'],
            ['LHA', 'Speed', '1.149', '1.158', '0.886'],
            ['RSP', 'R²', '0.145', '0.142', '0.686'],
            ['RSP', 'PR', '12.86', '9.50', '0.029'],
            ['RSP', 'Variance', '0.254', '0.231', '0.200'],
            ['RSP', 'Speed', '2.128', '1.828', '0.029'],
        ],
        widths_in=[0.9, 1.4, 1.0, 1.0, 0.8])
    add_para(doc, 'GRU-ODE matches discrete-GRU prediction accuracy (|ΔR²| < 0.003 in '
                  'all conditions) while reproducing all biological signals (RSP > LHA '
                  'predictability; fasting reduces RSP PR; LHA fasting effect on '
                  'variance/speed). Continuous-time formulation adds interpretability '
                  '(separable ODE vs. observation contributions) without sacrificing '
                  'fit quality.')

    # ---- Section 8: References
    add_heading(doc, '8. References', 1)
    add_bullets(doc, [
        'Chen, R.T.Q. et al. (2018). "Neural Ordinary Differential Equations." NeurIPS.',
        'De Brouwer, E. et al. (2019). "GRU-ODE-Bayes: Continuous modeling of sporadically-observed time series." NeurIPS.',
        'Rubanova, Y., Chen, R.T.Q., Duvenaud, D. (2019). "Latent ODEs for irregularly-sampled time series." NeurIPS.',
        'torchdiffeq library: https://github.com/rtqichen/torchdiffeq',
    ])

    # ---- Files
    add_heading(doc, '9. Code & outputs', 1)
    add_bullets(doc, [
        'Scripts: gru_ode_pooled_by_region.py (500 ms), gru_ode_10ms.py (10 ms Poisson).',
        'Results CSV (500 ms): data/gru_ode_pooled_by_region_results.csv.',
        'Results CSV (10 ms): figures/gru_ode_10ms_poisson_results.csv.',
        'Models: figures/gru_ode_pooled_{region}_{condition}_model.pt.',
        'Comparison figure: figures/gru_ode_vs_gru_comparison.png.',
        'Planning memory: memory/gru_ode_planning.md.',
    ])

    doc.save(OUT)
    print(f'Wrote {OUT}')


if __name__ == '__main__':
    main()
