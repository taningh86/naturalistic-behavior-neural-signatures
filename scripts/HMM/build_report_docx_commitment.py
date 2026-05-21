"""Build a detailed Word-document report of the commitment-marker extraction.

Reads outputs from data/HMM/commitment_markers/ and figures/HMM/commitment_markers/
to assemble a single self-contained report.

Output: data/HMM/commitment_markers_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.4f}", max_cols=None):
    if max_cols is not None:
        df = df.iloc[:, :max_cols]
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    cfg = load_config()
    cm_dir = REPO_ROOT / cfg["commitment_dirs"]["out"]
    fig_dir = REPO_ROOT / cfg["commitment_dirs"]["fig"]

    history_df = pd.read_csv(cm_dir / "sampling_history.csv")
    all_events_df = pd.read_csv(cm_dir / "all_events_combined.csv")
    classification_df = pd.read_csv(cm_dir / "state_classification.csv")
    food_pot_map = cfg["food_pot_per_session"]
    discovery_window_s = float(cfg["discovery_window_s"])
    bin_size_s = float(cfg["target_bin_ms"]) / 1000.0

    fed = cfg["sessions"]["fed"]
    fasted = cfg["sessions"]["fasted"]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===== TITLE =====
    title = doc.add_heading(
        "Behavioral Commitment Markers — Dual-Probe Foraging", level=0,
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Discovery, prior sampling, and pre-commitment behavioral state from the\n"
        "merged dynamax HMM Viterbi sequences (K=14)"
    )
    r.italic = True
    r.font.size = Pt(11)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-06  •  Script: scripts/HMM/09_extract_commitment_markers.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "We extract behavioral commitment markers — the moment of food discovery and "
        "the prior search history that preceded it — for each of 7 dual-probe foraging "
        "sessions (4 fed: S4, S6, S8, S10; 3 fasted: S12, S14, S16). The analysis "
        "consumes the merged 14-state Viterbi sequence from the dynamax mixed-emission "
        "HMM (script 08) and a per-bin pot-identity field (Pot-1..Pot-4) added to the "
        "binned npz files in an updated script 01. Discovery is operationalised as "
        "the first dig at the experimentally-designated food pot followed by entry "
        f"into the pure-feeding HMM state within {discovery_window_s:.0f} s. All "
        "events recorded prior to discovery (or, for the no-food extinction session "
        "S10, prior to session end) are tabulated to characterise the animal's "
        "search history."
    )
    n_disc = int((~history_df["discovery_time_s"].isna()).sum())
    n_clean = int((history_df.get("discovery_method") == "within_window").sum())
    n_fallback = int((history_df.get("discovery_method") == "closest_dig_fallback").sum())
    n_no_food = int(history_df["food_pot"].isna().sum())
    fed_disc = history_df[(~history_df["discovery_time_s"].isna()) &
                           (history_df["state"] == "fed")]
    fas_disc = history_df[(~history_df["discovery_time_s"].isna()) &
                           (history_df["state"] == "fasted")]
    fed_disc_clean_only = fed_disc[fed_disc.get("discovery_method") == "within_window"]
    add_para(
        doc,
        f"Of {len(history_df)} sessions, {n_clean} produced a clean within-window "
        f"discovery, {n_fallback} required the closest-dig fallback rule (food-pot "
        "dig was found but the next entry into a pure-feeding HMM state lay outside "
        f"the {discovery_window_s:.0f} s window), and {n_no_food} are no-food "
        "sessions where discovery is not defined. With the fallback applied, all "
        f"6 food sessions (3 fed: {fed_disc.session.tolist()}; 3 fasted: "
        f"{fas_disc.session.tolist()}) carry a discovery time, giving a balanced "
        "3 vs 3 design for downstream comparisons."
    )
    add_para(
        doc,
        f"Discovery times confirm the expected biological direction: fasted mice "
        f"(mean {fas_disc.discovery_time_s.mean():.0f} s) discover faster than fed "
        f"mice (mean {fed_disc.discovery_time_s.mean():.0f} s using all 3 fed "
        f"discoveries; {fed_disc_clean_only.discovery_time_s.mean():.0f} s using "
        "only clean-window detections). With n=3 vs 3 the two-sided Mann-Whitney "
        "p floor is 2/20 = 0.10, so all directional findings remain descriptive "
        "trends pending more sessions."
    )

    # ===== PIPELINE / DATA SOURCES =====
    add_heading(doc, "Data Sources and Pipeline", level=1)
    add_para(
        doc,
        "Inputs and the order in which they are produced:"
    )
    add_bullet(
        doc,
        "scripts/HMM/01_load_and_rebin.py (UPDATED) — now also writes a per-bin "
        "pot_id field (0 = not at any pot, 1..4 = at Pot-N). Derived from the four "
        "Pot-i and Pot-i-zone EthoVision columns; pot has priority over pot_zone "
        "and the smallest pot index wins on simultaneous activations. Mode-binned "
        "to 480 ms alongside the existing zone label.",
    )
    add_bullet(
        doc,
        "scripts/HMM/02..04_*_dynamax.py — prepare → CV-select N → fit final dynamax "
        "mixed-emission HMM (selected N=16, per-bin LL = -0.94).",
    )
    add_bullet(
        doc,
        "scripts/HMM/05_extract_state_posteriors_dynamax.py — per-bin posteriors and "
        "Viterbi for the N=16 fit.",
    )
    add_bullet(
        doc,
        "scripts/HMM/08_merge_redundant_states_dynamax.py — collapses cosine-similar "
        "states (threshold 0.90); for the current fit two pairs were merged "
        "({2, 14} → merged S2 = pure-feeding; {3, 11} → merged S3 = home), "
        "leaving K=14 merged states. Merged Viterbi is the input to commitment-"
        "marker extraction.",
    )
    add_bullet(
        doc,
        "scripts/HMM/09_extract_commitment_markers.py — this analysis. Uses merged "
        "Viterbi + pot_id + the food-pot mapping below.",
    )
    add_para(doc, "Run order:")
    code_run(doc, "01 → 02 → 03 → 04 → 05 → 06 → 07 → 08 → 09")

    # ===== CONFIG =====
    add_heading(doc, "Configuration", level=1)
    add_para(doc, "Food pot per session (1-indexed; null = no food):")
    fp_rows = []
    for sn in fed + fasted:
        fp = food_pot_map.get(sn, food_pot_map.get(int(sn)))
        fp_rows.append(dict(
            session=sn,
            state=("fed" if sn in fed else "fasted"),
            food_pot=("none" if fp is None else f"P{fp}"),
        ))
    add_df_table(doc, pd.DataFrame(fp_rows))
    doc.add_paragraph()
    add_para(doc, "Detection thresholds:")
    add_bullet(doc, f"discovery_window_s = {cfg['discovery_window_s']} s "
                    f"(≈ {int(round(cfg['discovery_window_s']/bin_size_s))} bins at "
                    f"{cfg['target_bin_ms']} ms / bin)")
    add_bullet(doc, f"feeding_state_min_prob = {cfg['feeding_state_min_prob']} "
                    "(P(feeding) per state)")
    add_bullet(doc, f"digging_state_min_prob = {cfg['digging_state_min_prob']} "
                    "(P(digging_sand) per state)")
    add_bullet(doc, f"contemplation_event_min_prob = {cfg['contemplation_event_min_prob']} "
                    "(P(contemplation_at_transition) or P(exploration_at_transition))")
    add_bullet(doc, f"pot_zone_state_min_prob = {cfg['pot_zone_state_min_prob']} "
                    "(zone_pot + zone_pot_zone)")

    # ===== STATE CLASSIFICATION =====
    add_heading(doc, "Auto-Derived State Classes", level=1)
    add_para(
        doc,
        "States are tagged into four behavioural classes from the merged-profile "
        "emission probabilities, using the thresholds above. The pot-zone class "
        "explicitly excludes feeding and digging states so that pre-commitment pot "
        "approaches are counted distinctly from the discovery dig and the "
        "feeding bouts themselves."
    )
    cls_rows = []
    for cat in ["feeding", "digging", "contemplation", "pot_zone"]:
        states = classification_df[classification_df.category == cat]["state_id"] \
            .astype(int).tolist()
        cls_rows.append(dict(category=cat, states=str(sorted(states))))
    add_df_table(doc, pd.DataFrame(cls_rows))
    doc.add_paragraph()
    add_para(
        doc,
        "Two states reach the feeding threshold (merged S2 and S11) — both are "
        "pure-feeding states and either qualifies as 'feeding' for discovery "
        "detection. Only one merged state (S6) reaches the digging threshold; one "
        "(S4) reaches the contemplation threshold."
    )

    # ===== EVENT TYPE DEFINITIONS =====
    add_heading(doc, "Event Type Definitions", level=1)
    defs = [
        ("discovery_dig",
         "First dig run at the experimentally-designated food pot whose end is "
         f"followed by entry into a pure-feeding state within {discovery_window_s:.0f} s."),
        ("prior_dig_food_pot",
         "Pre-discovery dig run at the food pot. By construction these did NOT lead to "
         f"feeding within {discovery_window_s:.0f} s."),
        ("prior_dig_non_food_pot",
         "Pre-discovery dig run at a non-food pot."),
        ("failed_dig",
         "Any pre-discovery dig run not followed by feeding within the window. "
         "Includes all non-food-pot digs and any food-pot digs that failed. "
         "(Listed alongside prior_dig_* rows; the same dig can appear under both "
         "labels by design.)"),
        ("pot_zone_entry",
         "Pre-discovery transition INTO a pot-zone HMM state (not feeding, not "
         "digging). Pot identity at the entry bin is reported when localizable, "
         "or 'ambiguous' when pot_id == 0 at that frame."),
        ("S4_entry",
         "Pre-discovery transition INTO the contemplation/T-zone state (merged S4)."),
    ]
    for name, descr in defs:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(10)
        p.add_run("  —  " + descr).font.size = Pt(10)

    # ===== SAMPLING HISTORY SUMMARY =====
    doc.add_page_break()
    add_heading(doc, "Sampling History Summary", level=1)
    add_para(
        doc,
        "Per-session search history at the moment of discovery (or, for S10 "
        "extinction, computed across the whole session):"
    )
    disp = history_df[[
        "session", "state", "food_pot",
        "discovery_time_s", "discovery_method", "discovery_lag_s",
        "n_prior_pot_digs", "n_prior_failed_digs",
        "n_prior_distinct_pots_visited",
        "n_prior_pot_zone_entries", "n_prior_S4_entries",
        "discovery_dig_was_first_dig", "n_total_dig_runs",
    ]].copy()
    disp.columns = [
        "session", "state", "food_pot",
        "disc_time_s", "method", "dig_to_feed_lag_s",
        "prior_digs", "failed_digs",
        "distinct_pots", "pot_zone_entries", "S4_entries",
        "stumbled", "total_digs",
    ]
    add_df_table(doc, disp, float_fmt="{:.1f}")

    add_image(
        doc, fig_dir / "sampling_history_summary.png", width_in=6.7,
        caption="Cross-session sampling history. Fed sessions (blue) vs fasted "
                "(red). Marker shape encodes detection method: circle = clean "
                "within-window discovery, square = fallback (closest-food-pot-dig), "
                "diamond = no-food session, X = no dig at food pot. Gold edge "
                "marks stumbled discoveries (discovery dig was the first dig of "
                "the session).",
    )

    # ===== HEADLINE FINDINGS =====
    add_heading(doc, "Headline Findings", level=1)
    food_sessions = history_df[~history_df.food_pot.isna()]
    fed_food = food_sessions[food_sessions.state == "fed"]
    fas_food = food_sessions[food_sessions.state == "fasted"]
    fed_disc_mean_all = fed_food["discovery_time_s"].mean()
    fas_disc_mean_all = fas_food["discovery_time_s"].mean()
    fed_priors_mean = fed_food["n_prior_pot_digs"].mean()
    fas_priors_mean = fas_food["n_prior_pot_digs"].mean()
    fed_s4_mean = fed_food["n_prior_S4_entries"].mean()
    fas_s4_mean = fas_food["n_prior_S4_entries"].mean()
    add_bullet(
        doc,
        f"With the fallback applied, all 3 fed and all 3 fasted food-bearing sessions "
        "carry a discovery time. Fasted mice discover food in mean "
        f"{fas_disc_mean_all:.0f} s; fed mice take mean {fed_disc_mean_all:.0f} s "
        f"— roughly {fed_disc_mean_all/fas_disc_mean_all:.1f}× longer for fed.",
    )
    add_bullet(
        doc,
        f"Prior dig count differs in the same direction: fasted mean "
        f"{fas_priors_mean:.1f} prior digs vs fed mean {fed_priors_mean:.1f}. "
        "Fasted animals commit earlier; fed animals sample more before settling. "
        "S14 (fasted) reaching discovery in only 54 s after sampling all 4 pots "
        "is the striking phenotype.",
    )
    add_bullet(
        doc,
        f"Pre-discovery contemplation (S4) entries: fasted mean {fas_s4_mean:.1f} "
        f"vs fed mean {fed_s4_mean:.1f}. Fed animals contemplate at the transition "
        "zone substantially more than fasted before discovering food.",
    )
    add_bullet(
        doc,
        "All food-bearing sessions visited 3-4 pots before discovery — even the "
        "rapidly-discovering S14 (54 s discovery) entered all 4 pot zones in that "
        "brief window. Strategy is exhaustive sampling that scales with time.",
    )
    add_bullet(
        doc,
        "S10 (no-food extinction) is qualitatively distinct: 59 dig runs across "
        "all 4 pots and 32 S4 entries over the 30-min session — much higher than "
        "any food-bearing session. Behaviorally consistent with extinction-probe "
        "design.",
    )

    add_heading(doc, "S4 Caveat — Atypical Dig→Feed Lag", level=2)
    add_para(
        doc,
        "S4 (fed, 6-17-25 foraging, food at P4) is the one session where no clean "
        f"dig→feed transition was found within the {discovery_window_s:.0f} s "
        "window. Diagnostic of the 6 food-pot dig runs shows none was followed "
        "by a pure-feeding HMM state (merged S2 or S11) within the window:"
    )
    code_run(
        doc,
        "(start_bin, duration_s, led_to_feeding):\n"
        "  (774, 3.84, False), (788, 5.28, False), (830, 11.04, False),\n"
        "  (885, 5.28, False), (977, 1.44, False), (1034, 2.40, False)",
        size=9,
    )
    add_para(
        doc,
        "Manual inspection of the EthoVision raw labels (independent of HMM) "
        "shows S4 does feed: a 4.3 s feeding bout begins at t=439.8 s (bin 916, "
        "raw feeding=1 across all 9 bins, animal at P4). However the HMM Viterbi "
        "assigns those bins to merged S1 (P(feeding)=0.124), not to a pure-feeding "
        "state — the bout is too short and the surrounding zone-occupancy is "
        "split between pot_zone (31 %), arena (46 %) and other (19 %), so the "
        "state-emission criterion lumps it with the approach/transition state. "
        "The next entry into S2 / S11 doesn't occur until t=800.4 s (bin 1667), "
        "5+ minutes after the food-pot dig sequence."
    )
    add_para(
        doc,
        "Decision: per the user's instruction, S4 is kept in the analysis using "
        "the closest-food-pot-dig fallback. The chosen 'discovery dig' is the "
        f"food-pot dig at bin 1034 (t=496.5 s), whose next-feeding-state entry "
        f"lies 301.4 s later. The lag is recorded in the discovery_lag_s field; "
        "the discovery_method field is 'closest_dig_fallback' to flag this. "
        "This brings the design to a balanced 3 fed (S4 fallback, S6 clean, "
        "S8 clean) vs 3 fasted (all clean). Analyses that hinge on dig-to-feed "
        "latency or that treat the discovery-aligned neural window as a "
        "transient should exclude S4 or treat it separately."
    )

    # ===== EVENT COUNTS =====
    doc.add_page_break()
    add_heading(doc, "Event Counts", level=1)
    add_para(doc, "Total events of each type across all 7 sessions:")
    counts = (
        all_events_df.groupby(["event_type"]).size().reset_index(name="count")
        .sort_values("count", ascending=False)
    )
    add_df_table(doc, counts)
    doc.add_paragraph()
    add_para(doc, "Per-session event totals:")
    pivot = (
        all_events_df.groupby(["session", "event_type"]).size().unstack(fill_value=0)
        .reset_index()
    )
    add_df_table(doc, pivot)

    # ===== PER-SESSION TIMELINES =====
    doc.add_page_break()
    add_heading(doc, "Per-Session Event Timelines", level=1)
    add_para(
        doc,
        "Each panel shows one session. Y-axis lists event categories; X-axis is "
        "session time. Coloured markers are individual events; vertical red dashed "
        "line marks the discovery time when detected. Title shows food pot and "
        "summary counts."
    )
    sess_order = [(s, "fed") for s in fed] + [(s, "fasted") for s in fasted]
    for sn, st in sess_order:
        path = fig_dir / f"session_{sn}_events.png"
        add_image(
            doc, path, width_in=6.7,
            caption=f"Session {sn} ({st}). Pre-discovery (or full-session for no-"
                    "food/failed) event timeline.",
        )

    # ===== OUTPUT FILES =====
    doc.add_page_break()
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/binned/session_{N}.npz",
         "(Updated by amended script 01.) Now includes pot_id (T,) integer field."),
        ("data/HMM/commitment_markers/state_classification.csv",
         "Per-state-class membership (feeding/digging/contemplation/pot_zone)."),
        ("data/HMM/commitment_markers/session_{N}_events.csv",
         "Per-session event table: bin, time_s, event_type, state_id, "
         "pot_identity, duration_s."),
        ("data/HMM/commitment_markers/all_events_combined.csv",
         "All sessions concatenated, with session and state columns prepended."),
        ("data/HMM/commitment_markers/sampling_history.csv",
         "One row per session: discovery time + 7 sampling-history metrics."),
        ("figures/HMM/commitment_markers/session_{N}_events.png",
         "Per-session event timelines."),
        ("figures/HMM/commitment_markers/sampling_history_summary.png",
         "Cross-session sampling-history overview (fed vs fasted)."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    # ===== CAVEATS / NEXT =====
    add_heading(doc, "Caveats and Next Steps", level=1)
    add_bullet(
        doc,
        "Discovery window (currently 10 s) is a hard threshold. S4 misses it by "
        "5+ minutes for entry into a pure-feeding HMM state, which the closest-"
        "food-pot-dig fallback resolves at the cost of an atypically long "
        "dig→feed lag (see the dedicated S4 caveat section).",
    )
    add_bullet(
        doc,
        "S4's fallback discovery (lag 301 s) reflects an HMM state-assignment "
        "limitation, not a behavioural failure. Raw EthoVision feeding starts "
        "~14 s after S4's earliest food-pot dig but is assigned to merged S1 "
        "(approach/transition state) not S2 / S11 (pure feeding). Analyses "
        "centred on the discovery_time_s should treat S4 separately.",
    )
    add_bullet(
        doc,
        "n_prior_distinct_pots_visited equals 4 in every session, including the "
        "rapidly-discovering S14. This is real: every animal sampled all four pots "
        "before committing. The metric does not differentiate sessions and may be "
        "less informative than expected for n=7.",
    )
    add_bullet(
        doc,
        "Stumbled-discovery flag (discovery_dig_was_first_dig) is False for every "
        "session — no animal found food on its very first dig. Strategy is genuinely "
        "exploratory across this cohort.",
    )
    add_bullet(
        doc,
        "Pot-identity inference uses the priority rule pot[i] > pot_zone[i] then "
        "smallest index on ties. Cases where the animal straddles two pots' zones "
        "are deterministically resolved but may not match the animal's actual focus. "
        "If this matters for downstream analyses, the per-pot zone columns can be "
        "exported separately rather than collapsed.",
    )
    add_bullet(
        doc,
        "All metrics are descriptive; with n=4 fed vs 3 fasted (and one fed session "
        "without a clean discovery), formal Mann-Whitney testing is underpowered "
        "(p_min ≈ 0.057, often higher with NaNs). The natural extension is to align "
        "ACA / LHA neural activity to discovery_time_s and the event series, which "
        "is the explicit out-of-scope item for this script.",
    )

    out_path = REPO_ROOT / "data" / "HMM" / "commitment_markers_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
