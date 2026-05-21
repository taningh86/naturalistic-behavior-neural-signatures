"""Build a detailed Word-document report of the dynamax mixed-emission HMM refit.

Reads outputs from data/HMM/ and figures/HMM/ to assemble a single
self-contained report.

Output: data/HMM/HMM_pipeline_report_dynamax.docx
"""
import json
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


# ---------- doc helpers ----------
def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.4f}", max_cols=None):
    if max_cols is not None:
        df = df.iloc[:, :max_cols]
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


# ---------- auto state gloss ----------
def auto_gloss_state(row, zone_labels, event_names):
    """Generate a one-sentence behavioral gloss from a state profile row."""
    # Dominant zone
    zone_probs = {z: row[f"zone_{z}_prob"] for z in zone_labels}
    top_zone = max(zone_probs, key=zone_probs.get)
    top_zone_p = zone_probs[top_zone]
    # All zones with > 0.2
    notable_zones = sorted(
        [(z, p) for z, p in zone_probs.items() if p > 0.2 and z != top_zone],
        key=lambda x: -x[1],
    )

    # Dominant events
    event_probs = {e: row[f"event_{e}_prob"] for e in event_names}
    notable_events = sorted(
        [(e, p) for e, p in event_probs.items() if p > 0.15],
        key=lambda x: -x[1],
    )

    # Speed
    sp = row["speed_z_mean"]
    sp_label = "fast" if sp > 0.5 else "slow" if sp < -0.5 else "moderate-speed"
    # Distance to pot
    dz = row["dist_z_mean"]
    dist_label = "far from pot" if dz > 0.5 else "near pot" if dz < -0.5 else "mid-distance"

    # Dwell
    dwell = row["mean_dwell_s"]
    dwell_label = (
        "long-dwell" if dwell > 8 else
        "short-dwell" if dwell < 1.5 else
        "medium-dwell"
    )

    occ = row["soft_occupancy"]

    # Build sentence
    zone_str = top_zone
    if top_zone_p < 0.95 and notable_zones:
        zone_str = f"{top_zone}/{notable_zones[0][0]}"

    event_str = ""
    if notable_events:
        ev_descs = []
        for e, p in notable_events[:3]:
            ev_descs.append(f"{e} (P={p:.2f})")
        event_str = "; events: " + ", ".join(ev_descs)
    else:
        event_str = "; no dominant event"

    return (f"{occ * 100:.1f}% occupancy; {dwell_label} ({dwell:.1f} s); "
            f"{sp_label}, {dist_label}; zone: {zone_str} (P={top_zone_p:.2f}){event_str}.")


# ---------- main ----------
def main():
    cfg = load_config()
    data_dir = REPO_ROOT / "data" / "HMM"
    fig_dir = REPO_ROOT / "figures" / "HMM"
    params_dir = data_dir / "final_model_params_dynamax"

    with open(params_dir / "meta.json") as f:
        meta = json.load(f)

    cv_df = pd.read_csv(data_dir / "cv_results_dynamax.csv")
    profiles_df = pd.read_csv(data_dir / "state_profiles_dynamax.csv")
    occ_df = pd.read_csv(data_dir / "fed_vs_fasted_dynamax.csv")
    stats_df = pd.read_csv(data_dir / "fed_vs_fasted_stats_dynamax.csv")
    sess_meta = pd.read_csv(data_dir / "prepared_dynamax" / "session_metadata.csv")

    zone_labels = meta["zone_labels"]
    event_names = meta["event_names"]
    N = meta["N"]

    doc = Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===== TITLE =====
    title = doc.add_heading(
        "Behavioral HMM Pipeline (dynamax refit) — Dual-Probe Foraging", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Custom mixed-emission HMM (Gaussian + Categorical + Bernoulli) on EthoVision behavior\n"
        "Mouse01 Coordinates-1 — fed vs fasted foraging sessions"
    )
    r.italic = True
    r.font.size = Pt(11)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-05  •  Refit follows the original ssm pipeline (2026-04-29)")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "We refit the behavioral HMM with proper factorized emissions, replacing the "
        "ssm jittered-Gaussian workaround in the original pipeline. Each per-bin "
        "observation is decomposed into three conditionally-independent groups given "
        "the hidden state: a 2-D diagonal Gaussian over speed and distance to nearest "
        "pot, a single Categorical over the 6-zone label (home / transition / pot / "
        "pot-zone / arena / other), and a 7-vector of Bernoulli events (incomplete "
        "home returns, quick-loop-at-home, digging-sand, feeding, rearing, transition-"
        "wall exploration, contemplation at T-zone). The pipeline runs end-to-end on "
        f"the same 7 dual-probe foraging sessions ({len(meta['fed_sessions'])} fed: "
        f"{meta['fed_sessions']}; {len(meta['fasted_sessions'])} fasted: "
        f"{meta['fasted_sessions']}) and was extended to N ∈ "
        f"{cfg['dynamax_N_range']} for state selection."
    )
    add_para(
        doc,
        f"Cross-validation (3 stratified folds, 5 random initializations per fold) "
        f"selected N = {N} states by the 1-SE rule. The final model converged in "
        f"{meta['n_iters_to_convergence']} EM iterations on the best of "
        f"{meta['n_inits_tried']} initializations, with a per-bin log-likelihood of "
        f"{meta['per_bin_log_likelihood']:.4f}. Because emissions are now genuine "
        "Bernoulli/Categorical instead of Gaussian-on-binary, log-likelihoods are "
        "not directly comparable to the prior ssm fit (the ssm per-bin LL of +13 "
        "was an artifact of variance collapse on near-deterministic columns)."
    )
    add_para(
        doc,
        "The headline biological finding from the ssm fit replicates and is sharpened: "
        "the original ssm 'long-dwell pot/feed/dig' state cleanly splits here into "
        "separate pure-feeding states (S2, S12, S14, P(feeding) ≈ 1.0) and a pure-"
        "digging state (S6, P(digging_sand) = 1.0), made possible by the factorized "
        "emissions. The strongest fed-vs-fasted contrast is now S14 (a pure-feeding "
        "state at the pot), at 17.5% in fasted vs 2.5% in fed sessions — perfect rank "
        "separation across the n=4 vs n=3 group sizes (Mann-Whitney U = 0). With "
        "those sample sizes the smallest two-sided p-value possible is 2/35 ≈ 0.057, "
        "so all directional findings remain descriptive trends rather than confirmed "
        "effects."
    )

    # ===== KEY DESIGN DECISIONS =====
    add_heading(doc, "Pipeline Design", level=1)

    add_heading(doc, "Why a refit was needed", level=2)
    add_para(
        doc,
        "The original pipeline concatenated all behavioral features into one Gaussian "
        "observation vector — the standard MoSeq / Wiltschko 2015 convention used because "
        "ssm does not natively support factorized emissions. Two undesirable consequences "
        "followed: (i) Gaussian variance collapsed on binary / one-hot columns when a "
        "state specialized on a single category, requiring a hand-tuned Gaussian-jitter "
        "fix (σ = 0.05) before EM was numerically stable; and (ii) the per-state "
        "Bernoulli-equivalent emission is encoded only in the Gaussian mean, so two "
        "behaviors that co-occur in time (e.g., 'feeding while at pot') cannot be "
        "decoupled into independent emission probabilities. The refit removes both "
        "issues by modeling each group with its native distribution."
    )

    add_heading(doc, "Mixed emission model", level=2)
    add_para(
        doc,
        "For each state k, the emission likelihood factorizes as:"
    )
    code_run(
        doc,
        "log P(x_t | z_t = k) = sum_d log N(x_cont_td | mu_kd, sigma_kd)\n"
        "                     + log Cat(x_zone_t | p_k_zone)\n"
        "                     + sum_e [x_e * log q_ke + (1 - x_e) * log(1 - q_ke)]",
        size=9,
    )
    add_para(
        doc,
        "Conditional independence between the three groups is assumed given the hidden "
        "state, so emission log-likelihoods sum across groups. M-step updates are "
        "closed-form weighted MLEs with smoothing priors:"
    )
    add_bullet(doc, "Gaussian: variance floor 1e-3 (prevents singular states).")
    add_bullet(doc, "Categorical zone: Dirichlet pseudocount α = 1 (prevents log(0)).")
    add_bullet(doc, "Bernoulli events: Beta(0.5, 0.5) pseudocount per event "
                    "(critical because some sessions have zero events of a given "
                    "type — e.g., S4 has no rearing).")

    add_heading(doc, "Implementation notes", level=2)
    add_para(
        doc,
        "EM is implemented in scripts/HMM/mixed_hmm.py. Forward filtering uses "
        "dynamax.hidden_markov_model.inference.hmm_filter; the backward pass and "
        "two-slice marginals are implemented in log-space directly in JAX. We bypass "
        "dynamax's hmm_smoother because dynamax 0.1.5 has a bug — `compute_trans_probs` "
        "is missing from the function's `static_argnames`, so any value of that kwarg "
        "is traced by JIT and triggers TracerBoolConversionError. Viterbi decoding "
        "uses dynamax's hmm_posterior_mode."
    )
    add_para(
        doc,
        "JAX runs on CPU only (Windows has no JAX GPU build); each EM iteration on "
        "all 7 sessions (~26k bins) takes ~100 ms at N=8 and ~600 ms at N=20, so "
        "the full CV sweep (135 fits) completes in roughly 25 minutes."
    )

    add_heading(doc, "Sessions and bin granularity", level=2)
    add_para(
        doc,
        f"Same 7 foraging-labeled dual-probe sessions as the ssm fit: fed = "
        f"{meta['fed_sessions']}, fasted = {meta['fasted_sessions']}. Session 18 "
        "remains excluded (no sorted neural data on 7/30 recording date). Native "
        f"EthoVision bin is 40 ms; rebinned by factor 12 → {cfg['target_bin_ms']} ms "
        f"bins. Median session length is {sess_meta['n_bins'].median():.0f} HMM "
        f"bins ({sess_meta['n_bins'].sum()} total)."
    )

    add_heading(doc, "Cross-validation strategy", level=2)
    add_para(
        doc,
        "Three stratified folds, each holding out one fed and one fasted session: "
        "(S4, S12), (S6, S14), (S8, S16). For each N in {4, 6, 8, 10, 12, 14, 16, 18, "
        "20} and each fold, 5 random initializations are run; the init with the "
        "highest training LL is kept and scored on the held-out test sessions. "
        "Recommended N is the smallest within 1 SE of the maximum mean held-out "
        "LL/bin (1-SE rule)."
    )

    # ===== PIPELINE SCRIPTS =====
    add_heading(doc, "Pipeline Scripts (dynamax variant)", level=1)
    pipeline_steps = [
        ("01_load_and_rebin.py",
         "Unchanged from the original pipeline. Reads each session xlsx, computes "
         "per-bin nearest-pot distance, priority-ordered zone label, Bernoulli events, "
         "and rebins to 480 ms. Output: data/HMM/binned/session_{N}.npz."),
        ("02_prepare_for_hmm_dynamax.py",
         "Loads pooled continuous values, computes pooled mean/SD, z-scores. Keeps "
         "zone as integer index (no one-hot) and events as 0/1 ints. No jitter. "
         "Saves three separate arrays per session — X_continuous (T, 2), X_zone (T,), "
         "X_events (T, 7) — to data/HMM/prepared_dynamax/session_{N}.npz, plus a "
         "zone-label-mapping CSV for downstream interpretation."),
        ("mixed_hmm.py",
         "Library module (not run directly). Defines MixedHMMParams dataclass, "
         "init_params(), e_step(), m_step(), fit(), held_out_loglik(), "
         "smoothed_posteriors(), viterbi_states(). All inference uses JAX; "
         "forward via dynamax.hmm_filter, backward + two-slice marginals "
         "hand-rolled in log-space."),
        ("03_fit_and_select_states_dynamax.py",
         "Builds 3 stratified CV folds. For each (N, fold) runs 5 random inits to "
         "convergence (max 500 iters, tol 1e-4), keeps the init with highest train "
         "LL, scores held-out per-bin LL. Aggregates mean ± SE across folds, applies "
         "the 1-SE rule. Output: data/HMM/cv_results_dynamax.csv, "
         "figures/HMM/state_selection_dynamax.png."),
        ("04_fit_final_model_dynamax.py",
         "CLI: --N N. Pooled fit on all 7 sessions, 5 random inits, picks best by "
         "training LL. Saves numpy parameter arrays (pi, A, mu, sigma, p_zone, q_events) "
         "to data/HMM/final_model_dynamax.npz, plus human-readable CSVs "
         "(initial_distribution, transition_matrix, emissions_continuous/zone/events) "
         "and meta.json in data/HMM/final_model_params_dynamax/."),
        ("05_extract_state_posteriors_dynamax.py",
         "Per session: forward-backward smoothed posteriors and Viterbi-decoded "
         "state. CSV per session in data/HMM/posteriors_dynamax/."),
        ("06_validate_states_dynamax.py",
         "Posterior-weighted per-state behavioral profile across all sessions; "
         "consolidated state×features heatmap (continuous z, zone P, event P). "
         "Per-session timelines (posterior heatmap + Viterbi ribbon + dig/feed event "
         "ticks). Auto-warnings for low-occupancy (<2%), redundant pairs (cosine "
         "similarity > 0.95), flickering states (mean dwell < 2 bins)."),
        ("07_compare_metabolic_states_dynamax.py",
         "Per-session soft/hard occupancy, Viterbi mean dwell, empirical transition "
         "matrix; per-state Mann-Whitney U on session-level fed-vs-fasted occupancy "
         "and dwell; bar plots with session scatter; transition-matrix triptych "
         "(fed mean, fasted mean, difference)."),
    ]
    for name, descr in pipeline_steps:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run("  —  " + descr)
        r2.font.size = Pt(10)

    add_para(doc, "Run order: 01 → 02 → 03 → 04 → 05 → 06 → 07. Use:")
    code_run(
        doc,
        'PYTHONIOENCODING=utf-8 "C:\\Users\\Gregg\\anaconda3\\envs\\si_env\\python.exe" '
        '-u scripts/HMM/0X_*_dynamax.py',
    )
    add_para(
        doc,
        "PYTHONIOENCODING=utf-8 is required (scripts use the Δ and → glyphs); -u "
        "disables Python stdout buffering so progress prints appear in real time when "
        "piped to tee/log.",
        size=10,
    )

    # ===== STATE SELECTION =====
    doc.add_page_break()
    add_heading(doc, "State Selection (Cross-Validation)", level=1)
    add_para(
        doc,
        "For each (N, fold), the best of 5 random inits is selected by training LL "
        "and used to score held-out per-bin LL. The aggregate is the mean and SE "
        "across folds of those best-init scores."
    )
    best = (cv_df.sort_values("train_ll", ascending=False)
                  .groupby(["N", "fold"], as_index=False).first())
    cv_agg = best.groupby("N", as_index=False).agg(
        mean_ll_per_bin=("heldout_ll_per_bin", "mean"),
        se=("heldout_ll_per_bin",
            lambda x: x.std(ddof=1) / np.sqrt(len(x))),
        n_folds=("fold", "count"),
        mean_iters=("n_iter", "mean"),
    )
    cv_agg.columns = ["N", "mean ll/bin", "SE", "folds", "mean EM iters"]
    add_df_table(doc, cv_agg, float_fmt="{:.3f}")
    doc.add_paragraph()

    max_mean = cv_agg["mean ll/bin"].max()
    se_at_max = float(cv_agg.loc[cv_agg["mean ll/bin"].idxmax(), "SE"])
    add_para(
        doc,
        f"The mean held-out LL/bin rises monotonically from "
        f"{cv_agg['mean ll/bin'].min():.2f} at N={int(cv_agg.iloc[0]['N'])} to "
        f"{max_mean:.2f} at N={int(cv_agg.loc[cv_agg['mean ll/bin'].idxmax(),'N'])}, "
        f"and is still rising at the top of the swept range. The 1-SE rule "
        f"(threshold = max − SE_at_max = {max_mean - se_at_max:.3f}) "
        f"selects the smallest N whose mean LL/bin exceeds the threshold: "
        f"N = {N}. State counts above 20 may yield further gains; this can be "
        "tested by extending dynamax_N_range in config.yaml."
    )
    add_image(
        doc, fig_dir / "state_selection_dynamax.png", width_in=5.5,
        caption="Held-out log-likelihood per bin vs N (mean ± SE across 3 folds). "
                "Dashed line: 1-SE threshold below max-mean. Dotted vertical: "
                "recommended N.",
    )

    # ===== FINAL MODEL =====
    doc.add_page_break()
    add_heading(doc, "Final Model", level=1)
    add_para(
        doc,
        f"Fit on all 7 sessions pooled at N = {N}; "
        f"emission groups = {{Gaussian (D=2), Categorical (K_zone={meta['K_zone']}), "
        f"Bernoulli (n_events={meta['n_events']})}}; method = EM; max_iters = "
        f"{meta['em_max_iters']}; tolerance = {meta['em_tol']}; "
        f"variance floor = {meta['var_floor']}; zone Dirichlet α = "
        f"{meta['zone_dirichlet']}; event Beta β = {meta['event_beta']}. "
        f"{meta['n_inits_tried']} random initializations were run; the model "
        f"converged in {meta['n_iters_to_convergence']} EM iterations on the "
        f"chosen init (idx {meta['chosen_init_idx']}, seed "
        f"{meta['chosen_seed']}). Total log-likelihood = "
        f"{meta['final_log_likelihood']:.0f} over {meta['total_bins']} bins, "
        f"i.e. {meta['per_bin_log_likelihood']:.4f} per bin."
    )

    init_recs = pd.DataFrame(meta["all_init_records"])
    if len(init_recs):
        add_para(doc, "All initializations:", size=10)
        init_recs = init_recs[["init_idx", "seed", "ll", "n_iter", "time_s"]]
        init_recs.columns = ["init", "seed", "final ll", "EM iters", "time (s)"]
        add_df_table(doc, init_recs, float_fmt="{:.2f}")

    add_heading(doc, "Per-state behavioral profile heatmap", level=2)
    add_para(
        doc,
        "Three panels: continuous emissions (state means in z-units, RdBu_r), "
        "zone categorical probabilities (Blues), and Bernoulli event probabilities "
        "(Reds). Each row is one state; cells are posterior-weighted summaries "
        "computed across all 7 sessions."
    )
    add_image(
        doc, fig_dir / "state_profiles_dynamax.png", width_in=6.7,
        caption="State × feature heatmap. Continuous (z), zone P, event P.",
    )

    # State summary table — concise key columns
    add_heading(doc, "State summary table", level=2)
    state_summary = profiles_df[[
        "state", "soft_occupancy", "mean_dwell_s",
        "speed_z_mean", "dist_z_mean",
        "zone_home_prob", "zone_transition_prob",
        "zone_pot_prob", "zone_pot_zone_prob", "zone_arena_prob",
        "event_digging_sand_prob", "event_feeding_prob",
        "event_rearing_prob",
        "event_exploration_at_transition_prob",
        "event_contemplation_at_transition_prob",
    ]].copy()
    state_summary.columns = [
        "state", "occ", "dwell_s",
        "speed_z", "dist_z",
        "P(home)", "P(trans)", "P(pot)", "P(potZ)", "P(arena)",
        "P(dig)", "P(feed)", "P(rear)", "P(explT)", "P(contT)",
    ]
    add_df_table(doc, state_summary, float_fmt="{:.2f}")

    # Auto-generated per-state gloss
    add_heading(doc, "Auto-generated per-state interpretations", level=2)
    add_para(
        doc,
        "These descriptions are generated mechanically from the profile (occupancy, "
        "dominant zone, dominant events with P > 0.15, speed/distance signs, and "
        "Viterbi mean dwell). They are starting points for interpretation, not "
        "manual labels."
    )
    for _, row in profiles_df.iterrows():
        k = int(row["state"])
        gloss = auto_gloss_state(row, zone_labels, event_names)
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"State {k}: ")
        r.bold = True
        p.add_run(gloss)

    # Validation warnings
    add_heading(doc, "Validation warnings (auto-detected)", level=2)
    flagged_low = profiles_df[profiles_df["soft_occupancy"] < 0.02]
    flagged_flicker = profiles_df[profiles_df["mean_dwell_bins"] < 2]
    if len(flagged_low):
        add_bullet(
            doc,
            f"Low-occupancy (< 2%): " +
            ", ".join(f"S{int(s)} ({o*100:.2f}%)"
                       for s, o in zip(flagged_low["state"], flagged_low["soft_occupancy"]))
            + "."
        )
    else:
        add_bullet(doc, "No low-occupancy states (< 2%).")
    if len(flagged_flicker):
        add_bullet(
            doc,
            f"Flickering (mean dwell < 2 bins): " +
            ", ".join(f"S{int(s)} ({d:.2f} bins)"
                       for s, d in zip(flagged_flicker["state"], flagged_flicker["mean_dwell_bins"]))
            + "."
        )
    else:
        add_bullet(doc, "No flickering states (mean dwell < 2 bins).")
    add_bullet(
        doc,
        "Redundant pair: S2 and S14 (cosine similarity 0.951 on full profile vector). "
        "Both are pure-feeding-at-pot states with very similar zone and event "
        "probabilities; the difference is in the continuous Gaussian (slightly "
        "different speed/distance distributions). A re-fit at N=15 might absorb "
        "them into a single feeding state without losing predictive performance."
    )

    # ===== FED VS FASTED =====
    doc.add_page_break()
    add_heading(doc, "Fed vs Fasted Comparison", level=1)
    add_para(
        doc,
        "Comparisons are session-level: per state, one occupancy value and one mean "
        "dwell value per session, then a two-sided Mann-Whitney U on n=4 fed vs n=3 "
        "fasted session means. The smallest two-sided p-value possible at these "
        "sample sizes is 2/35 ≈ 0.0571 (perfect rank separation), so individual "
        "p-values should be read as ranks rather than as confirmatory significance."
    )

    add_heading(doc, "Per-state statistics", level=2)
    occ_stats = stats_df[stats_df["metric"] == "soft_occupancy"][[
        "state", "fed_mean", "fasted_mean", "U", "p"
    ]].copy().sort_values("p")
    occ_stats.columns = ["state", "occ_fed", "occ_fasted", "U", "p"]
    dwell_stats = stats_df[stats_df["metric"] == "mean_dwell_s"][[
        "state", "fed_mean", "fasted_mean", "U", "p"
    ]].copy().sort_values("p")
    dwell_stats.columns = ["state", "dwell_fed (s)", "dwell_fasted (s)", "U", "p"]

    add_para(doc, "Soft occupancy:", size=10, bold=True)
    add_df_table(doc, occ_stats.reset_index(drop=True), float_fmt="{:.3f}")
    doc.add_paragraph()
    add_para(doc, "Mean Viterbi dwell (s):", size=10, bold=True)
    add_df_table(doc, dwell_stats.reset_index(drop=True), float_fmt="{:.3f}")

    add_heading(doc, "Headline findings", level=2)
    add_bullet(
        doc,
        "S14 (pure pot+feeding state, P(feeding)=1.0, low speed): occupancy 17.5% "
        "fasted vs 2.5% fed (U=0, p=0.057, perfect rank separation); mean dwell "
        "5.0 s fed vs 18.7 s fasted (U=1, p=0.108). The factorized refit pulls this "
        "state out as cleanly food-locked, where the ssm fit conflated it with "
        "the digging state."
    )
    add_bullet(
        doc,
        "S6 (pure pot+digging state, P(digging_sand)=1.0): mean dwell 9.81 s fasted "
        "vs 3.90 s fed (U=0, p=0.057). Digging is now its own state instead of being "
        "co-modeled with feeding as in the ssm fit."
    )
    add_bullet(
        doc,
        "S4 (transition zone with contemplation/exploration): mean dwell 4.95 s "
        "fasted vs 3.53 s fed (U=0, p=0.057). Modest occupancy difference; this "
        "is a candidate 'deliberation' state."
    )
    add_bullet(
        doc,
        "S7 (small low-event state, ~2% occupancy, flicker-prone): occupancy 4.96% "
        "fed vs 2.82% fasted (U=12, p=0.057, perfect separation in the fed direction)."
    )
    add_bullet(
        doc,
        "Multiple feeding-rich states (S2, S5, S12, S14) with slightly different "
        "speed/distance profiles: these split the original ssm 'feeding' signal "
        "across several behavioral subtypes."
    )
    add_bullet(
        doc,
        "No state reaches p < 0.05 (the 0.057 floor remains). All findings remain "
        "descriptive trends; statistical confirmation requires more sessions."
    )

    add_image(
        doc, fig_dir / "fed_vs_fasted_dynamax.png", width_in=6.7,
        caption="Fed vs fasted: per-state soft occupancy (top) and Viterbi mean "
                "dwell (bottom). Bars are group mean ± SEM; dots are session-level "
                "values; '.' marks p<0.10, '*' marks p<0.05 (none reach with n=4 vs 3).",
    )
    add_image(
        doc, fig_dir / "fed_vs_fasted_transitions_dynamax.png", width_in=6.7,
        caption="Empirical Viterbi transition matrices. Left: fed mean P. "
                "Centre: fasted mean P. Right: fasted − fed (red = more in fasted, "
                "blue = more in fed). The diagonal dominates as expected; off-"
                "diagonal asymmetries indicate which inter-state transitions are "
                "shifted by metabolic state.",
    )

    # ===== Comparison to ssm fit =====
    doc.add_page_break()
    add_heading(doc, "Comparison to ssm Pipeline (2026-04-29)", level=1)
    add_para(
        doc,
        "The earlier ssm pipeline (12 states, jittered-Gaussian emissions) and the "
        "present dynamax refit (16 states, factorized emissions) agree on the "
        "qualitative biological story but differ in granularity and in the "
        "interpretability of individual states. Direct LL comparison is not "
        "meaningful (jittered-Gaussian on binary columns inflates positive LL "
        "via near-zero variance); the comparison is on state structure and "
        "fed-vs-fasted contrast."
    )
    cmp = pd.DataFrame([
        ["Number of states (CV-selected)", "12", "16"],
        ["Held-out LL/bin (CV best mean)", "+13.42 (artifact)", "−1.00 to −1.16"],
        ["Strongest fasted state", "S6: pot/feed/dig combined", "S14 (pure feed) + S6 (pure dig) split"],
        ["Fasted occupancy of strongest state", "51.6% vs 19.2% fed", "S14: 17.5% vs 2.5% fed"],
        ["Mean dwell of strongest state", "30.9 s fasted vs 12.3 s fed", "S6: 9.8 s vs 3.9 s; S14: 18.7 s vs 5.0 s"],
        ["Smallest p-value attainable (n=4 vs 3)", "0.057", "0.057"],
        ["Number of p < 0.06 hits", "4 states", "≥ 4 states (S6 dwell, S14 occ, S4 dwell, S7 occ)"],
        ["Variance-collapse fix needed", "Yes (jitter σ=0.05)", "No (true Bernoulli/Categorical)"],
        ["Behaviors decoupled per state", "Tied via Gaussian mean", "Independent Bernoulli probs per event"],
    ], columns=["Aspect", "ssm fit (N=12)", "dynamax refit (N=16)"])
    add_df_table(doc, cmp, float_fmt="{:.4f}")
    doc.add_paragraph()
    add_para(
        doc,
        "Practical upshot: the dynamax fit is preferable for any downstream analysis "
        "that needs interpretable per-behavior emission probabilities — for example, "
        "tagging neural-population traces with whether a bin is 'feeding-rich' vs "
        "'digging-rich' independently. The ssm fit remains valid as a coarser "
        "behavior segmentation but should not be used to claim that a single state "
        "encodes a multi-behavior composite (which the factorized fit reveals to "
        "be an artifact of tied Gaussian means)."
    )

    # ===== TIMELINES =====
    doc.add_page_break()
    add_heading(doc, "Per-Session Timelines", level=1)
    add_para(
        doc,
        "Each panel shows a single session. Top: per-bin posterior heatmap "
        "(states on y, time on x, posterior probability viridis). Bottom: "
        "Viterbi-decoded state as a colored ribbon. Vertical red ticks: "
        "digging-sand events; vertical orange ticks: feeding events."
    )
    sess_order = [(s, "fed") for s in meta["fed_sessions"]] + \
                 [(s, "fasted") for s in meta["fasted_sessions"]]
    for sn, st in sess_order:
        path = fig_dir / "timelines_dynamax" / f"session_{sn}.png"
        add_image(
            doc, path, width_in=6.7,
            caption=f"Session {sn} ({st}). 30-min recording; HMM posteriors over time.",
        )

    # ===== OUTPUTS =====
    doc.add_page_break()
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/binned/session_{N}.npz",
         "(Unchanged — produced by 01.) Per-session rebinned arrays."),
        ("data/HMM/prepared_dynamax/session_{N}.npz",
         "X_continuous (T,2), X_zone (T,), X_events (T,7) — factorized inputs."),
        ("data/HMM/prepared_dynamax/zone_label_mapping.csv",
         "Maps zone integer index → human-readable label."),
        ("data/HMM/cv_results_dynamax.csv",
         "Per (N, fold, init) train LL, held-out LL, fit time."),
        ("data/HMM/final_model_dynamax.npz",
         "Final fitted HMM parameters (numpy arrays: pi, A, mu, sigma, p_zone, q_events)."),
        ("data/HMM/final_model_params_dynamax/",
         "initial_distribution.csv, transition_matrix.csv, "
         "emissions_continuous.csv, emissions_zone.csv, "
         "emissions_events.csv, meta.json."),
        ("data/HMM/posteriors_dynamax/session_{N}.csv",
         "Per-bin smoothed posteriors p_state_0..N-1 + Viterbi assignment + time."),
        ("data/HMM/state_profiles_dynamax.csv",
         "Per-state behavioral profile (means, std, all zone/event probs, occupancy, dwell)."),
        ("data/HMM/state_dwell_occupancy_dynamax.csv",
         "Compact occupancy + mean dwell per state."),
        ("data/HMM/fed_vs_fasted_dynamax.csv",
         "Per-session × state occupancy + mean dwell."),
        ("data/HMM/fed_vs_fasted_transitions_dynamax.csv",
         "Per-session × (i, j) empirical transition probabilities."),
        ("data/HMM/fed_vs_fasted_stats_dynamax.csv",
         "Per-state Mann-Whitney U on occupancy and dwell."),
        ("figures/HMM/state_selection_dynamax.png",
         "Cross-validation curve (mean ± SE)."),
        ("figures/HMM/state_profiles_dynamax.png",
         "States × features heatmap (3-panel: continuous, zone, events)."),
        ("figures/HMM/timelines_dynamax/session_{N}.png",
         "Per-session posterior + Viterbi timeline with event overlays."),
        ("figures/HMM/fed_vs_fasted_dynamax.png",
         "Bar plots of fed-vs-fasted occupancy and dwell with session scatter."),
        ("figures/HMM/fed_vs_fasted_transitions_dynamax.png",
         "Mean transition matrices for fed, fasted, and fasted − fed difference."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    # ===== CAVEATS =====
    add_heading(doc, "Caveats and Next Steps", level=1)
    add_bullet(
        doc,
        "Underpowered: the n=4 vs 3 Mann-Whitney floor (p_min ≈ 0.057) is unchanged. "
        "All claims remain descriptive trends. Adding even one more session per "
        "group (5 vs 4 → p_min ≈ 0.016) would let the strongest hits clear 0.05."
    )
    add_bullet(
        doc,
        "LL still rising at N=20: the CV curve has not yet plateaued. Extending "
        "dynamax_N_range to 24 or 28 may favor a larger model. The 1-SE rule "
        "tends to be conservative when the curve is monotone."
    )
    add_bullet(
        doc,
        "Redundant feeding states (S2 and S14, cosine 0.951): a re-fit at N=15 "
        "might collapse them into one feeding state without losing predictive "
        "performance, and may make the metabolic-state contrast even cleaner."
    )
    add_bullet(
        doc,
        "No null control: a shuffled-state-sequence baseline would establish that "
        "the fed-vs-fasted occupancy differences exceed what one would see from "
        "trivially relabelled states. This is the obvious next step before claiming "
        "biology even at p < 0.06."
    )
    add_bullet(
        doc,
        "Neural alignment: with the per-bin Viterbi or posterior streams in "
        "posteriors_dynamax/session_{N}.csv, behavior-state-conditioned ACA / LHA "
        "population activity can now be averaged and contrasted across states. The "
        "factorized fit is well suited to this — a 'state 14 = pot+feeding' label "
        "is a much more specific neural-alignment regressor than the ssm fit's "
        "composite long-dwell pot/feed/dig state."
    )
    add_bullet(
        doc,
        "Dependency note: si_env now pins jax==0.4.30 + jaxlib==0.4.30 + "
        "dynamax==0.1.5 (installed via --no-deps to avoid pulling numpy 2.x, "
        "which would break cebra, kilosort, numba, hdbscan, pynwb, tensorflow, "
        "elephant, and giotto-tda). JAX runs on CPU only on Windows."
    )

    # Save
    out_path = REPO_ROOT / "data" / "HMM" / "HMM_pipeline_report_dynamax.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
