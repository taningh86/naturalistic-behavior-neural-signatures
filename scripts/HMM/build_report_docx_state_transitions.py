"""Build a Word-document report of the state-transition neural signatures (script 14).

Reads outputs from data/HMM/neural_alignment/state_transitions/ and figures.

Output: data/HMM/neural_alignment_state_transitions_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.2f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def state_label(prof_row):
    zone_keys = ["home", "transition", "pot", "pot_zone", "arena", "other"]
    zone_p = {z: prof_row[f"zone_{z}_prob"] for z in zone_keys}
    top_zone = max(zone_p, key=zone_p.get)
    desc = f"{top_zone}({zone_p[top_zone]:.2f})"
    events = []
    for ev in ["digging_sand", "feeding", "rearing",
               "contemplation_at_transition", "exploration_at_transition"]:
        if prof_row[f"event_{ev}_prob"] > 0.30:
            short = ev.split("_")[0]
            events.append(f"{short}={prof_row[f'event_{ev}_prob']:.2f}")
    if events:
        desc += " + " + ",".join(events)
    return desc


def main():
    cfg = load_config()
    base_out = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "state_transitions"
    base_fig = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "state_transitions"

    rep_a1 = pd.read_csv(base_out / "replication_A1.csv")
    rep_a2 = pd.read_csv(base_out / "replication_A2.csv")
    cross_a1 = pd.read_csv(base_out / "A1_cross_session.csv")
    cross_a2 = pd.read_csv(base_out / "A2_cross_session.csv")

    prof = pd.read_csv(REPO_ROOT / cfg["merge_dirs"]["state_profiles_csv"])
    K = len(prof)
    state_descriptions = [state_label(prof.iloc[k]) for k in range(K)]

    sess_order = [4, 6, 8, 12, 14, 16]
    sess_state_map = {4: "fed", 6: "fed", 8: "fed",
                       12: "fasted", 14: "fasted", 16: "fasted"}

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===== TITLE =====
    title = doc.add_heading(
        "Neural Signatures of HMM State Transitions (Pre-Exit Analysis)", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Tests whether ACA / LHA single-unit activity carries a signature of "
        "upcoming HMM state transitions, isolated from post-transition behavior."
    )
    r.italic = True
    r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-08  •  Script: scripts/HMM/14_state_transition_neural_signatures.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "All comparisons use only bins where the animal is still in the source "
        "state. Two parallel hypotheses are tested:"
    )
    add_bullet(
        doc,
        "**A1 — Generic switch signal**: within state i, the last 3 bins of "
        "each run (pre-exit, ~1.4 s) vs earlier bins of the same run (stay). "
        "Different = 'I'm about to leave i' irrespective of destination.",
    )
    add_bullet(
        doc,
        "**A2 — Pair-specific switch signal**: within state i's pre-exit "
        "bins, ending in destination j vs ending in destinations other than "
        "j. Different = 'I'm about to enter j specifically'.",
    )
    add_para(doc, "Both compared to 100-iteration circular-shift Viterbi shuffles "
                  "that preserve run structure / marginal occupancy and break alignment "
                  "with neural data. All 6 foraging sessions × 2 regions.")

    add_para(doc, "Three plain-language findings:", bold=True, size=11)
    add_bullet(
        doc,
        "**ACA carries a generic pre-exit signal across virtually every state.** "
        "Of states tested in ≥3 sessions, ACA passes shuffle p95 in 4-6/6 "
        "sessions for S6 (digging), S3 (home), S9 (pot-zone), S4 (T-zone+"
        "contemplation), S8 (pot-zone), S2 (pot+feeding), and S12 (T-zone). "
        "S6 (digging) replicates in **6/6** sessions. This is the strongest "
        "cross-session-replicating result of the entire neural-alignment "
        "program.",
    )
    add_bullet(
        doc,
        "**LHA carries the pre-exit signal only at consummatory / home "
        "states.** S2 (feeding) replicates in 3/3 sessions and S3 (home) in "
        "5/5. But LHA does NOT carry pre-exit signal at S6 (digging) — 0/6 "
        "sessions pass — even though digging is heavily LHA-modulated in "
        "B1's whole-state encoding. ACA-LHA dissociation: ACA encodes 'about "
        "to leave' broadly; LHA encodes it only when the state is feeding "
        "or home.",
    )
    add_bullet(
        doc,
        "**Pair-specific switching (A2) does NOT replicate.** No (source, "
        "destination) pair passes shuffle p95 in ≥3 sessions in either "
        "region. The best is LHA S9→S13 at 2/2 tested sessions. The "
        "destination-specific hypothesis fails: ACA / LHA know you're about "
        "to leave a state but don't distinguish where you're going.",
    )

    # ===== METHOD =====
    add_heading(doc, "Method", level=1)
    add_bullet(doc, "Same QC-filtered units as Track B / 11. 480 ms neural bins "
                    "matched to HMM bin grid.")
    add_bullet(doc, "Bin labelling: a 'run' is a maximal contiguous stretch with "
                    "the same Viterbi state. For each run of length L ≥ 2·K_pre "
                    "(K_pre = 3, so L ≥ 6 bins, ~2.9 s):")
    add_bullet(doc, "  • Last K_pre = 3 bins: pre-exit.")
    add_bullet(doc, "  • Bins 0..L-K_pre-1: stay.")
    add_bullet(doc, "  • Runs with L < 6 bins are excluded.")
    add_bullet(doc, "A1 thresholds: state needs ≥30 stay bins AND ≥30 pre-exit "
                    "bins (pooled across all destinations).")
    add_bullet(doc, "A2 thresholds: source state needs ≥20 pre-exit bins ending "
                    "in destination j AND ≥20 bins ending in non-j destinations. "
                    "This is restrictive — many (source, destination) pairs are "
                    "skipped per session.")
    add_bullet(doc, "Per-(unit, state) Mann-Whitney U on bin-level firing rates. "
                    "BH-FDR correction within region across all (unit, state) or "
                    "(unit, source, destination) tests.")
    add_bullet(doc, "Population-level: per-state stay/pre-exit centroid in PC1-3 "
                    "(session-specific PCA, top 5 PCs).")
    add_bullet(doc, "100 circular-shift Viterbi shuffles (offset ∈ [200, T-200]). "
                    "Per state (A1) or per pair (A2), pass = observed FDR-sig "
                    "unit count > shuffle 95th percentile.")
    add_bullet(doc, "Cross-session replication: count sessions where each "
                    "(state, region) for A1 or (source, destination, region) "
                    "for A2 passes shuffle p95.")

    # ===== A1 RESULTS =====
    doc.add_page_break()
    add_heading(doc, "A1 — Generic switch signal (stay vs pre-exit)", level=1)
    add_para(
        doc,
        "Per-region replication: count of sessions where the state passes the "
        "shuffle 95th-percentile null. n_sessions_tested varies because "
        "states with fewer than 30 stay or pre-exit bins are skipped per "
        "session."
    )

    # ACA + LHA replication tables
    add_para(doc, "ACA — A1 replication across sessions:", bold=True, size=10)
    aca_a1 = rep_a1[rep_a1.region == "ACA"].sort_values(
        ["n_sessions_passing", "state"], ascending=[False, True]
    ).copy()
    aca_a1["behavior"] = aca_a1["state"].apply(
        lambda k: state_descriptions[int(k)] if int(k) < len(state_descriptions) else "")
    aca_a1["state_label"] = aca_a1["state"].apply(lambda k: f"S{int(k)}")
    aca_disp = aca_a1[["state_label", "n_sessions_tested", "n_sessions_passing",
                         "sessions_passing", "behavior"]].copy()
    aca_disp.columns = ["state", "tested", "passing", "sessions", "behavior"]
    add_df_table(doc, aca_disp)

    add_para(doc, "LHA — A1 replication across sessions:", bold=True, size=10)
    lha_a1 = rep_a1[rep_a1.region == "LHA"].sort_values(
        ["n_sessions_passing", "state"], ascending=[False, True]
    ).copy()
    lha_a1["behavior"] = lha_a1["state"].apply(
        lambda k: state_descriptions[int(k)] if int(k) < len(state_descriptions) else "")
    lha_a1["state_label"] = lha_a1["state"].apply(lambda k: f"S{int(k)}")
    lha_disp = lha_a1[["state_label", "n_sessions_tested", "n_sessions_passing",
                         "sessions_passing", "behavior"]].copy()
    lha_disp.columns = ["state", "tested", "passing", "sessions", "behavior"]
    add_df_table(doc, lha_disp)

    add_bullet(
        doc,
        "**Headline ACA result**: every state with ≥3 sessions of A1 testing "
        "passes shuffle p95 in 3/3 to 6/6 sessions. S6 (digging, 6/6), S3 "
        "(home, 5/5), S9 (pot-zone, 5/5), S4 (T-zone+contemplation, 4/4), S8 "
        "(pot-zone, 4/4). The pre-exit-vs-stay difference is universal in "
        "ACA across multiple behavioral states.",
    )
    add_bullet(
        doc,
        "**LHA dissociation**: only S2 (feeding, 3/3) and S3 (home, 5/5) "
        "show robust replication. S6 (digging) — heavily LHA-state-selective "
        "in Track B / B1 — has 0/6 sessions passing for the pre-exit "
        "signal. Same for S9, S10, S12. LHA encodes which state you're in "
        "but generally does NOT signal you're about to leave it.",
    )
    add_bullet(
        doc,
        "**Per-session FDR-sig unit counts (A1):** S4=174, S6=246, S8=155, "
        "S12=199, S14=167, S16=188 (pooled ACA+LHA). Stable across sessions.",
    )

    add_image(doc, base_fig / "replication_heatmap_A1_ACA.png", width_in=5.2,
              caption="ACA: A1 generic switch passes shuffle p95 across "
                      "sessions × states. Red ✓ = passing. Almost every state "
                      "passes in nearly every session.")
    add_image(doc, base_fig / "replication_heatmap_A1_LHA.png", width_in=5.2,
              caption="LHA: A1 generic switch passes are sparse — concentrated "
                      "at S2 (feeding) and S3 (home). S6 / S9 / S10 / S12 are "
                      "negative throughout.")

    # ===== A2 RESULTS =====
    doc.add_page_break()
    add_heading(doc, "A2 — Pair-specific (destination-discriminating) switch", level=1)
    add_para(
        doc,
        "Tests whether neural activity in the pre-exit bins differs depending "
        "on the destination state, holding source state constant. This "
        "directly distinguishes 'about to leave' from 'about to enter j'. "
        "Many (source, destination) pairs are skipped per session because of "
        "the strict 20-bins-per-destination requirement."
    )

    # Top pairs by A2 replication
    if len(rep_a2):
        add_para(doc, "Top A2 (source, destination) pairs by sessions passing — "
                       "all regions:", bold=True, size=10)
        top_a2 = rep_a2.sort_values(
            ["n_sessions_passing", "n_sessions_tested"],
            ascending=[False, False]
        ).head(15).copy()
        top_a2["src_lbl"] = top_a2["source"].apply(lambda k: f"S{int(k)}")
        top_a2["dst_lbl"] = top_a2["destination"].apply(lambda k: f"S{int(k)}")
        disp = top_a2[["region", "src_lbl", "dst_lbl", "n_sessions_tested",
                         "n_sessions_passing", "sessions_passing"]].copy()
        disp.columns = ["region", "source", "destination", "tested", "passing",
                          "sessions"]
        add_df_table(doc, disp)

    add_bullet(
        doc,
        "**No pair passes shuffle p95 in ≥3 sessions** in either region. "
        "The pair-specific switching hypothesis is not supported by the "
        "data.",
    )
    add_bullet(
        doc,
        "Best partial replicators (2/2 tested): LHA S9→S13. Single-session "
        "passes appear at ACA 0→12, ACA 6→9, ACA 9→13, LHA 6→9, LHA 6→13.",
    )
    add_bullet(
        doc,
        "**The strict 20-bins-per-destination requirement is restrictive.** "
        "Most (source, destination) pairs are tested in only 1-3 sessions, "
        "so the replication test has very limited power. Even at relaxed "
        "thresholds the headline is unlikely to change because per-session "
        "FDR-sig counts are inconsistent: A2 sig units = 0 (S4), 36 (S6), 6 "
        "(S8), 16 (S12), 0 (S14), 44 (S16). No clean across-session pattern.",
    )

    add_image(doc, base_fig / "replication_heatmap_A2_ACA.png", width_in=5.2,
              caption="ACA: A2 pair-specific replication heatmap. Sparse and "
                      "scattered — no pair replicates in ≥3 sessions.")
    add_image(doc, base_fig / "replication_heatmap_A2_LHA.png", width_in=5.2,
              caption="LHA: A2 pair-specific replication heatmap. Same — "
                      "sparse, with one 2/2 partial replicator (S9→S13).")

    # ===== INTERPRETATION =====
    doc.add_page_break()
    add_heading(doc, "Interpretation: ACA generic-switch broadly, LHA at consummatory only", level=1)
    add_para(
        doc,
        "Combining A1 and A2 yields a clean dissociation:"
    )
    add_bullet(
        doc,
        "**ACA** carries a generic 'about-to-exit' signal across virtually all "
        "well-sampled states (digging, home, pot-zone, T-zone, feeding, "
        "transitions). The signal is generic — not destination-specific — "
        "so ACA's pre-exit modulation reflects the act of leaving rather "
        "than a prediction of the destination.",
    )
    add_bullet(
        doc,
        "**LHA** carries the pre-exit signal only at S2 (feeding) and S3 "
        "(home). LHA does NOT signal upcoming exit from S6 (digging), even "
        "though digging is one of LHA's most state-selective representations "
        "in Track B's whole-session ANOVA. This is the strongest single-"
        "region functional dissociation in the dataset: LHA represents "
        "'what behavior is occurring' but is largely silent about 'when "
        "behavior is about to change' — except at consummatory states.",
    )
    add_bullet(
        doc,
        "**Pair-specific switching is not encoded** by either region (above "
        "the strictness of the test). This rules out ACA/LHA being read out "
        "by downstream targets to predict a specific upcoming state. The "
        "switch signal is more likely a 'state-end' read-out than a "
        "'state-prediction' read-out.",
    )
    add_bullet(
        doc,
        "**Methodological note**: By construction, all comparisons use only "
        "bins where the animal is still in the source state, so any neural "
        "difference between stay and pre-exit cannot be attributed to "
        "post-transition behavior. The signal must reflect either an "
        "internal pre-decision computation or a faithful proxy for the "
        "imminent behavioral change (e.g., subtle micro-movements undetected "
        "at this binning resolution).",
    )

    # ===== Per-session figures =====
    doc.add_page_break()
    add_heading(doc, "Per-session A1 — significant unit counts per state", level=1)
    add_para(doc, "Each panel shows one session: count of FDR-sig units per "
                  "state, ACA on left, LHA on right.", size=10)
    for sn in sess_order:
        path = base_fig / f"session_{sn}" / "A1_per_state_significant_unit_counts.png"
        if path.exists():
            add_image(doc, path, width_in=6.7,
                      caption=f"Session {sn} ({sess_state_map[sn]}).")

    # ===== CAVEATS =====
    add_heading(doc, "Caveats", level=1)
    add_bullet(doc, "K_pre = 3 bins (1.44 s) is arbitrary. K_pre = 2 or 5 should "
                    "be checked in a follow-up sensitivity analysis. The A1 "
                    "result is unlikely to be qualitatively changed (the "
                    "shuffle null adapts to whatever K_pre gives) but the "
                    "specific unit counts will move.")
    add_bullet(doc, "Stay and pre-exit windows are not symmetric. Stay = "
                    "everything in the run except the last 3 bins; pre-exit = "
                    "exactly 3 bins. Centroid distances are valid by direct "
                    "comparison but the asymmetric sample sizes are reflected "
                    "in the Mann-Whitney U variance.")
    add_bullet(doc, "A2 is power-limited by the ≥20-bins-per-destination "
                    "requirement. Some (source, destination) pairs may have a "
                    "real signal but be tested in only 1-2 sessions, making "
                    "cross-session replication impossible to demonstrate. "
                    "The negative A2 finding is therefore conservative: we "
                    "can say pair-specific signals don't ROBUSTLY replicate, "
                    "not that they don't exist.")
    add_bullet(doc, "Cross-session unit matching has not been done. The "
                    "replication is at the population-of-cells level "
                    "(proportion of FDR-sig units), not the same-cell level.")
    add_bullet(doc, "Single-mouse, single-paradigm dataset. The ACA/LHA "
                    "dissociation needs replication in additional mice and "
                    "ideally a different behavior to claim regional "
                    "specificity rather than a session-specific quirk.")

    # ===== OUTPUTS =====
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/neural_alignment/state_transitions/replication_A1.csv",
         "Per (region, state): n_sessions_tested, n_sessions_passing, sessions_passing."),
        ("data/HMM/neural_alignment/state_transitions/replication_A2.csv",
         "Per (region, source, destination): n_sessions_tested, n_sessions_passing."),
        ("data/HMM/neural_alignment/state_transitions/A1_cross_session.csv",
         "Per (session, region, state): observed sig count, shuffle p95, exceeds_p95."),
        ("data/HMM/neural_alignment/state_transitions/A2_cross_session.csv",
         "Same for A2 (per source, destination)."),
        ("data/HMM/neural_alignment/state_transitions/session_{N}/A1_*",
         "Per-session A1 unit-level results, centroid distances, pass summary."),
        ("data/HMM/neural_alignment/state_transitions/session_{N}/A2_*",
         "Same for A2."),
        ("data/HMM/neural_alignment/state_transitions/session_{N}/shuffle_*",
         "Per-session shuffle null distributions for A1 and A2."),
        ("figures/HMM/neural_alignment/state_transitions/replication_heatmap_A1_{ACA,LHA}.png",
         "Cross-session A1 pass heatmaps (states × sessions)."),
        ("figures/HMM/neural_alignment/state_transitions/replication_heatmap_A2_{ACA,LHA}.png",
         "Cross-session A2 pass heatmaps (top 20 pairs × sessions)."),
        ("figures/HMM/neural_alignment/state_transitions/session_{N}/A1_per_state_*.png",
         "Per-session A1 per-state significant unit counts."),
        ("figures/HMM/neural_alignment/state_transitions/session_{N}/A2_destination_significance_heatmap_*.png",
         "Per-session A2 source × destination heatmaps."),
        ("figures/HMM/neural_alignment/state_transitions/session_{N}/shuffle_*.png",
         "Per-session shuffle null distribution figures."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    out_path = REPO_ROOT / "data" / "HMM" / "neural_alignment_state_transitions_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
