"""Build a Word-document report of the CORRECTED LFP spectral pipeline (script 17).

Reads outputs from data/HMM/neural_alignment/lfp_spectral/ and figures.

Output: data/HMM/lfp_spectral_v2_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO_ROOT = Path(__file__).resolve().parents[2]

BASE_OUT = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "lfp_spectral"
BASE_FIG = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "lfp_spectral"

SESSION_STATE = {4: "fed", 6: "fed", 8: "fed",
                 12: "fasted", 14: "fasted", 16: "fasted"}

OLD_SE_R = 0.948    # single-ended baseline from S12 LFP diagnostics
DIAG_BIPOLAR_R = 0.009  # S12 diagnostic bipolar mean |r|

OUT_DOCX = REPO_ROOT / "data" / "HMM" / "lfp_spectral_v2_report.docx"


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def df_table(doc, df, float_fmt="{:.3f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    sanity = pd.read_csv(BASE_OUT / "sanity_bipolar_r_cross_session.csv")
    m1_cross = pd.read_csv(BASE_OUT / "M1_cross_session.csv")
    m1_rep = pd.read_csv(BASE_OUT / "M1_replication.csv")
    m3_cross = pd.read_csv(BASE_OUT / "M3_cross_session.csv")
    m3_rep = pd.read_csv(BASE_OUT / "M3_replication.csv")
    m4_cross = pd.read_csv(BASE_OUT / "M4_cross_session.csv")
    m4_rep = pd.read_csv(BASE_OUT / "M4_replication.csv")
    m4_sign = pd.read_csv(BASE_OUT / "M4_sign_test.csv")
    m2_paired = pd.read_csv(BASE_OUT / "M2_cross_session_paired_t.csv")

    # Per-session event counts from M2 npz files
    spec_dir = BASE_OUT / "spectrograms"
    m2_events = []
    for sn in [4, 6, 8, 12, 14, 16]:
        for reg in ("ACA", "LHA"):
            for st in [2, 3, 4, 6, 8, 9, 12]:
                fp = spec_dir / f"session_{sn}_state_{st}_{reg}.npz"
                if fp.exists():
                    d = np.load(fp)
                    m2_events.append(dict(session=sn, region=reg, state=st,
                                          n_events=int(d["n_events"])))
    m2_events = pd.DataFrame(m2_events)

    doc = Document()

    # Title
    t = doc.add_heading("LFP Spectral Analysis at HMM State Transitions", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sr = sub.add_run("Corrected pipeline (v2): static-geometry bipolar referencing — "
                       "all 6 foraging sessions (S4/6/8 fed, S12/14/16 fasted)")
    sr.italic = True
    sr.font.size = Pt(11)

    para(doc, "Script: scripts/HMM/17_lfp_spectral_analysis.py", size=9)
    para(doc, f"Outputs: {BASE_OUT.as_posix()}", size=9)
    para(doc, "Generated 2026-05-11", size=9)
    doc.add_paragraph()

    # Executive Summary
    heading(doc, "Executive Summary", level=1)
    para(doc, "Three plain-language findings.", bold=True)

    bullet(
        doc,
        f"Bipolar correction worked. Pre-correction (single-ended) ACA-LHA delta "
        f"r ≈ {OLD_SE_R:.3f} on S12 (LFP diagnostics report). With "
        f"geometry-correct bipolar referencing the regional ACA-LHA Pearson r "
        f"drops to a mean of {sanity['regional_ACA_vs_LHA_pearson_r'].mean():.3f} "
        f"across the 6 foraging sessions "
        f"(range {sanity['regional_ACA_vs_LHA_pearson_r'].min():.3f} - "
        f"{sanity['regional_ACA_vs_LHA_pearson_r'].max():.3f}). "
        f"Common-mode contamination removed.",
    )
    bullet(
        doc,
        "Band-power and coherence state-transition effects are negative once "
        "common-mode is removed. M1 (per-region per-state per-band power, stay vs "
        "pre-exit Mann-Whitney + FDR + 100 circular-shift shuffles) gives nearly "
        "all 0/6 sessions passing in every ACA cell and 0/n_tested in every LHA "
        "cell, with only scattered single-session hits at noise level. M3 "
        "(ACA-LHA coherence per state x band) shows the same pattern. The "
        "uncorrected pipeline's apparent state-locked band power and coherence "
        "effects were dominated by common-mode contamination.",
    )
    bullet(
        doc,
        "LFP envelope Granger is bidirectional with NO consistent ACA-leading "
        "direction at the LFP level. Per-session F values are large in BOTH "
        "directions (e.g. S16 delta: ACA→LHA F=2932, LHA→ACA F=20; S8 delta: "
        "343 vs 52). But the sign test across 6 sessions (count of sessions "
        "where ACA→LHA F > LHA→ACA F) is null in every band: best is delta "
        "4/6 ACA-leads, binomial p=0.69. This contrasts with the spike-level "
        "Granger result (script 16, PC1 6/6 ACA-leads, binom p=0.031): the "
        "ACA→LHA directional signal exists at the spike-population level but "
        "NOT at the LFP envelope level. ",
    )

    # Method
    heading(doc, "Method", level=1)
    bullet(doc, "Sessions: 6 dual-probe foraging — S4, S6, S8 (fed), S12, S14, S16 (fasted).")
    bullet(doc, "Source: CatGT-extracted .lf.bin at 2500 Hz (Cat_GT_Out/catgt_DOUBLE_PROBE_*_FOR_g0/*imec{0,1}/*.lf.bin).")
    bullet(
        doc,
        "Geometry: static IMRO geometry parsed once via "
        "scripts/HMM/lfp_parse_geometry.py from the 6_17_25 EXP reference .ap.meta "
        "(NP2013, 4-shank, 250 µm shank pitch). 370 imec0 ACA pairs, 184 imec1 "
        "LHA pairs (y<2500 µm), 184 imec1 RSP pairs (y≥2500 µm); all pairs are "
        "within-shank ≤30 µm nearest neighbors.",
    )
    bullet(
        doc,
        "Preprocessing per probe (chunked 30 s blocks, memmap int16): gain to µV → "
        "iirnotch at 60 and 120 Hz → bipolar via static pair indices → "
        "mean within region → 8th-order Butterworth LP at 200 Hz (anti-alias) → "
        "decimate ×5 to 500 Hz → artifact mask at |x| > 5 × MAD proxy SD. "
        "Saves ACA, LHA, RSP regional traces.",
    )
    bullet(
        doc,
        "M1 band power: Welch PSD per HMM bin (nperseg=240 samples), band-mean "
        "power per (delta 1-4, theta 4-12, beta 15-30, low_gamma 30-60, "
        "high_gamma 60-100 Hz), stay-vs-pre_exit Mann-Whitney + FDR within "
        "region + 100 circular-shift shuffles for replication.",
    )
    bullet(
        doc,
        "M2 spectrograms: multitaper (DPSS NW=3, K=5), 0.5 s window, 0.1 s step, "
        "log2 fold-change vs [-3,-1] s baseline, event-aligned to state exits, "
        "aggregated across sessions.",
    )
    bullet(
        doc,
        "M3 ACA-LHA coherence: Welch CSD per HMM bin, band-mean coherence, "
        "stay-vs-pre_exit Mann-Whitney + 100 shuffles per (state, band).",
    )
    bullet(
        doc,
        "M4 LFP Granger: bivariate VAR on Hilbert envelope per band over S3 "
        "stay + 5 s post-exit segments, BIC lag selection (1-20 samples at 500 Hz "
        "= 2-40 ms), analytical F and 100 circular-shift shuffles. Sign test "
        "across sessions = binomial(n_sessions ACA→LHA F > LHA→ACA F).",
    )

    # Sanity table
    heading(doc, "Sanity check: bipolar correction", level=1)
    para(
        doc,
        "Pearson r between regional bipolar ACA and LHA over the full foraging "
        "session, computed on the 500 Hz decimated regional means. The S12 "
        f"LFP diagnostics report showed single-ended ACA-LHA delta r ≈ {OLD_SE_R:.3f} "
        f"and S12 bipolar mean |r| ≈ {DIAG_BIPOLAR_R:.3f}. The corrected pipeline "
        "should show |r| well below 0.5 if bipolar referencing is doing its job.",
    )

    sanity_view = sanity.copy()
    sanity_view["session"] = sanity_view["session"].astype(int)
    sanity_view["state"] = sanity_view["session"].map(SESSION_STATE)
    sanity_view = sanity_view[["session", "state", "regional_ACA_vs_LHA_pearson_r"]]
    sanity_view = sanity_view.rename(columns={"regional_ACA_vs_LHA_pearson_r": "r"})
    df_table(doc, sanity_view, float_fmt="{:.3f}")

    para(
        doc,
        f"All 6 sessions ≪ 0.5. Mean r = {sanity['regional_ACA_vs_LHA_pearson_r'].mean():.3f}. "
        f"S16 highest at {sanity['regional_ACA_vs_LHA_pearson_r'].max():.3f} but still "
        "well within acceptable range. The geometry-correct bipolar referencing "
        "successfully removed the common-mode contamination.",
    )

    # M1
    heading(doc, "M1. Band Power per State × Band (stay vs pre-exit)", level=1)
    para(
        doc,
        "Per-region (ACA, LHA), per-state (ACA: 2,3,4,6,8,9,12; LHA: 2,3), "
        "per-band (5 bands) Mann-Whitney + FDR across the band × state cells. "
        "Replication = number of sessions where the cell observed FDR-sig in "
        "the real data AND the per-shuffle FDR-sig rate < 5%. Below: replication "
        "tables per region.",
    )

    heading(doc, "M1 replication: ACA", level=2)
    aca_m1 = m1_rep[m1_rep.region == "ACA"][
        ["state", "band", "n_tested", "n_passing"]
    ].sort_values(["state", "band"]).reset_index(drop=True)
    df_table(doc, aca_m1, float_fmt="{:.0f}")

    heading(doc, "M1 replication: LHA", level=2)
    lha_m1 = m1_rep[m1_rep.region == "LHA"][
        ["state", "band", "n_tested", "n_passing"]
    ].sort_values(["state", "band"]).reset_index(drop=True)
    df_table(doc, lha_m1, float_fmt="{:.0f}")

    n_aca_pass = int(aca_m1["n_passing"].sum())
    n_lha_pass = int(lha_m1["n_passing"].sum())
    n_aca_cells = int(len(aca_m1))
    n_lha_cells = int(len(lha_m1))
    para(
        doc,
        f"Summary: {n_aca_pass}/{n_aca_cells} ACA cells have any session passing; "
        f"{n_lha_pass}/{n_lha_cells} LHA cells. Almost all 0/n_tested. The single "
        "ACA hit (state 9 theta) and the scattered LHA single-session hits are "
        "at noise level. Negative result.",
    )

    # M2 spectrograms
    heading(doc, "M2. Event-Aligned Spectrograms (state exits)", level=1)
    para(
        doc,
        "Per-state, per-region event-aligned spectrograms around HMM state exits. "
        "Window: [-3, +3] s relative to exit. Multitaper PSD (DPSS NW=3, K=5 "
        "tapers), 0.5 s window, 0.1 s step. Each spectrogram is normalized as "
        "log2(power / baseline_mean) where baseline = [-3, -1] s. Aggregated "
        "across the n sessions where the state was observed at least 3 times.",
    )

    # M2.a Per-session event counts
    heading(doc, "M2.a Per-session event counts (state exits)", level=2)
    para(
        doc,
        "Number of exit events per (state, session). Each exit is a transition "
        "where the Viterbi changes from `state` to anything else, with run length "
        "≥ 2 × K_PRE (6 bins, ≈2.9 s). Counts are identical between ACA and LHA "
        "because the exit definition is region-agnostic.",
    )
    ev_pivot = m2_events[m2_events.region == "ACA"].pivot_table(
        index="state", columns="session", values="n_events", aggfunc="first",
    ).fillna(0).astype(int).reset_index()
    df_table(doc, ev_pivot, float_fmt="{:.0f}")
    total_events = int(m2_events[m2_events.region == "ACA"]["n_events"].sum())
    para(
        doc,
        f"Total: {total_events} state-exit events across the 6 sessions × 7 states "
        f"(state-region grid is duplicated: same {total_events} exits used for "
        "both regions' spectrograms).",
    )

    # M2.b Cross-session paired-t per band
    heading(doc, "M2.b Cross-session paired t-test, pre-exit vs baseline band-mean",
              level=2)
    para(
        doc,
        "For each (region, state, band) cell, the per-session band-mean log2 "
        "fold-change in the pre-exit window [-1.4, 0] s is paired-tested against "
        "the per-session band-mean log2 fold-change in the baseline window "
        "[-3, -1] s (paired t across 6 sessions). Reports mean log2-FC values "
        "for each window, their delta, t statistic, and uncorrected p value. "
        "Note the baseline mean is near zero by construction (the spec was "
        "normalized to the [-3, -1] s baseline) — the test detects whether the "
        "pre-exit window deviates from the baseline.",
    )

    heading(doc, "M2.b ACA", level=3)
    aca_m2 = m2_paired[m2_paired.region == "ACA"].copy()
    aca_m2 = aca_m2[["state", "band", "mean_log2fc_pre_exit", "delta", "t", "p"]]
    aca_m2 = aca_m2.sort_values(["state", "band"]).reset_index(drop=True)
    df_table(doc, aca_m2, float_fmt="{:.3f}")

    heading(doc, "M2.b LHA", level=3)
    lha_m2 = m2_paired[m2_paired.region == "LHA"].copy()
    lha_m2 = lha_m2[["state", "band", "mean_log2fc_pre_exit", "delta", "t", "p"]]
    lha_m2 = lha_m2.sort_values(["state", "band"]).reset_index(drop=True)
    df_table(doc, lha_m2, float_fmt="{:.3f}")

    sig = m2_paired[m2_paired.p < 0.05]
    para(
        doc,
        f"Uncorrected significance (p<0.05): {len(sig)} cells out of {len(m2_paired)} "
        f"tested ({100*len(sig)/len(m2_paired):.1f}%). At α=0.05 the false "
        f"positive rate expectation is exactly 0.05 × {len(m2_paired)} = "
        f"{0.05*len(m2_paired):.1f} cells. The observed count is at chance.",
    )

    if len(sig):
        heading(doc, "M2.b Cells reaching uncorrected p<0.05", level=3)
        sig_view = sig[["region", "state", "band", "mean_log2fc_baseline",
                          "mean_log2fc_pre_exit", "delta", "t", "p"]]\
                     .sort_values("p").reset_index(drop=True)
        df_table(doc, sig_view, float_fmt="{:.3f}")
        para(
            doc,
            "Note: ACA state-8 delta and LHA state-8 delta have opposite signs "
            "(ACA pre-exit Δ=+0.105 log2-FC, LHA pre-exit Δ=-0.016 log2-FC). "
            "Neither survives Bonferroni or FDR across the 70 tests "
            "(Bonferroni-corrected α = 0.05/70 ≈ 0.0007; the p values are "
            "0.036 and 0.038). No defensible state-locked pre-exit band-power "
            "change at the M2 timescale.",
        )

    para(
        doc,
        "Verdict: M2 corroborates M1 and M3 — no replicable state-locked LFP "
        "power deviation at exit transitions once geometry-correct bipolar "
        "removes common-mode. The largest cell-level effects (ACA delta state 8 "
        "and 12 around 0.10 log2-FC) are within the spread of the null "
        "distribution at this sample size.",
    )

    heading(doc, "M2.c Aggregate spectrograms (figures)", level=2)
    para(
        doc,
        "Visual inspection of the cross-session aggregates per (region, state). "
        "Each plot is the mean log2 fold-change across sessions, plotted from "
        "1 to 100 Hz, with t=0 at state exit. Color scale is symmetric (red = "
        "pre-exit power increase, blue = decrease).",
    )
    states_to_show = [3, 6, 8, 9, 12]
    for st in states_to_show:
        for reg in ("ACA", "LHA"):
            fp = BASE_FIG / f"M2_aggregate_state_{st}_{reg}.png"
            if fp.exists():
                add_image(doc, fp, width_in=5.5,
                          caption=f"M2 aggregate: state {st}, {reg}")

    # M3 coherence
    heading(doc, "M3. ACA-LHA Coherence per State × Band", level=1)
    para(
        doc,
        "Replication table — number of sessions where the (state, band) cell "
        "exceeds shuffle p95 in the observed delta (pre_exit mean coherence "
        "minus stay mean coherence).",
    )
    m3_rep_view = m3_rep[["state", "band", "n_tested", "n_passing"]].sort_values(
        ["state", "band"]
    ).reset_index(drop=True)
    df_table(doc, m3_rep_view, float_fmt="{:.0f}")
    n_m3_pass = int(m3_rep["n_passing"].sum())
    n_m3_cells = int(len(m3_rep))
    para(
        doc,
        f"Summary: {n_m3_pass}/{n_m3_cells} cells have any session passing — "
        "isolated single-session hits at noise level (S6 beta and theta 1/5; "
        "S8 low_gamma 1/3; S12 theta 1/3). No state × band combination "
        "replicates across sessions. Negative.",
    )

    # M4 Granger
    heading(doc, "M4. LFP Granger Envelope (S3 stay + 5 s post-exit)", level=1)
    para(
        doc,
        "Bivariate VAR per band on Hilbert envelope, lag selected by BIC. Both "
        "directions (ACA→LHA, LHA→ACA) tested per session. Sign test counts "
        "across sessions which direction has the larger F.",
    )

    heading(doc, "M4 per-session F values (observed F per band × direction)", level=2)
    pivot = m4_cross.pivot_table(
        index=["band", "direction"], columns="session",
        values="observed_F", aggfunc="first",
    ).reset_index()
    df_table(doc, pivot, float_fmt="{:.2f}")

    heading(doc, "M4 sign test across sessions", level=2)
    sign_view = m4_sign.copy()
    df_table(doc, sign_view, float_fmt="{:.3f}")
    para(
        doc,
        "Reading: in every band the sign test is null. The largest asymmetry "
        "is delta with 4/6 sessions ACA-leading (binom p=0.69, two-sided). "
        "high_gamma even reverses (1/6 ACA-leads, 5/6 LHA-leads, binom p=0.22). "
        "There is no consistent ACA→LHA directionality at the LFP envelope "
        "level. F values are large in both directions in most sessions; "
        "coupling is bidirectional and symmetric.",
    )

    heading(doc, "M4 replication: shuffle-pass and F-pass counts per band × direction",
              level=2)
    m4_rep_view = m4_rep[["band", "direction", "n_tested",
                            "n_pass_shuf", "n_pass_F"]].reset_index(drop=True)
    df_table(doc, m4_rep_view, float_fmt="{:.0f}")

    # Comparison
    heading(doc, "Comparison vs script 16 (spike-population Granger)", level=1)
    para(
        doc,
        "Script 16 ran the same Granger analysis on the population PC1 "
        "constructed from spike rates (per-unit z-scored, PCA across units). "
        "Result: ACA→LHA leads in 6/6 sessions, binomial p=0.031 (script 16 "
        "memory). The LFP envelope Granger run here does NOT replicate that "
        "asymmetry in any band.",
    )
    para(
        doc,
        "Interpretation: the directional ACA→LHA signal at home state exits "
        "is captured by the population spike PCA but is washed out at the LFP "
        "envelope level. Two non-exclusive explanations:",
    )
    bullet(
        doc,
        "Selectivity: spike-PC1 is a learned projection that emphasizes the "
        "subset of units carrying the pre-exit signal. The LFP envelope is a "
        "spatial average over the whole bipolar pair set in each region and "
        "mixes that signal with all the rest of the local population activity.",
    )
    bullet(
        doc,
        "Timescale: spike PC1 may carry sub-band-resolution information (e.g. "
        "phase-specific or transient features) that survives a 480 ms HMM bin "
        "but does not appear as a sustained envelope effect in one of the 5 "
        "broad bands used here.",
    )

    # Interpretation
    heading(doc, "Interpretation / Recommendation", level=1)
    para(
        doc,
        "The corrected pipeline (script 17 v2) takes the surviving claims from "
        "this whole HMM/LFP track and reduces them substantially:",
    )
    bullet(doc, "M1 band power → negative (was: extensive state-locked band-power changes)")
    bullet(doc, "M3 coherence → negative (was: state-locked coherence changes)")
    bullet(
        doc,
        "M4 Granger → bidirectional bulk coupling, no asymmetry (was: ACA→LHA "
        "leading in fed sessions like S4)",
    )
    bullet(
        doc,
        "The script 16 spike-level ACA→LHA Granger result is preserved as the "
        "single robust directional finding at home state exits.",
    )
    para(
        doc,
        "Recommendation: any claim that requires LFP power or coherence to "
        "change at HMM state transitions must be dropped or reformulated. The "
        "directionality story should be framed as a spike-level result only.",
    )

    # Caveats
    heading(doc, "Caveats", level=1)
    bullet(
        doc,
        "Static geometry: parsed from 6_17_25 EXP .ap.meta. We trust that all "
        "foraging sessions share the same IMRO map (per project notes). "
        "A per-session geometry parse would catch any drift but was out of scope.",
    )
    bullet(
        doc,
        "y=2500 µm LHA/RSP split is along the IMRO physical y coordinate, not "
        "kilosort spike depth. Real LHA tissue spans y < 345 µm and RSP spans "
        "y > 4680 µm in spike depths; the y=2500 boundary chosen for bipolar "
        "averaging splits the probe down the middle, so the 'LHA' regional "
        "mean includes channels in white matter / non-LHA tissue between the "
        "two regions. RSP regional mean is saved but not analyzed.",
    )
    bullet(
        doc,
        "Welch nperseg = 240 samples at 500 Hz gives ~2 Hz frequency resolution "
        "and ~0.5 s temporal resolution per HMM bin. Fine for band-mean "
        "comparisons but does not resolve narrow rhythms (e.g. theta phase).",
    )
    bullet(
        doc,
        "n_sessions = 6 is the same statistical floor as the rest of this "
        "pipeline. Binomial test minimum two-sided p with n=6 is 2/64 = 0.031 "
        "(6-of-6 same direction).",
    )
    bullet(
        doc,
        "All p values reported are unadjusted across analyses M1-M4; within "
        "each analysis FDR is applied as documented. There is no cross-analysis "
        "multiple-comparison correction.",
    )
    bullet(
        doc,
        "Common-input confound: bidirectional Granger may reflect a shared "
        "third driver (e.g. global state) rather than direct ACA-LHA coupling. "
        "This is unavoidable in a two-region recording.",
    )

    # Output Files
    heading(doc, "Output Files", level=1)
    code(doc, "data/HMM/neural_alignment/lfp_spectral/")
    bullet(doc, "sanity_bipolar_r_cross_session.csv — per-session ACA-LHA bipolar r")
    bullet(doc, "M1_cross_session.csv — per-session-per-cell pass flags (M1)")
    bullet(doc, "M1_replication.csv — replication counts per (region, state, band)")
    bullet(doc, "M3_cross_session.csv, M3_replication.csv — coherence")
    bullet(doc, "M4_cross_session.csv — per-session F values per (band, direction)")
    bullet(doc, "M4_replication.csv — shuffle/F pass counts per (band, direction)")
    bullet(doc, "M4_sign_test.csv — binomial test ACA→LHA vs LHA→ACA per band")
    bullet(doc, "discovered_lfp_files.csv — input file inventory")
    bullet(doc, "preprocessed_v2/session_{N}_{ACA,LHA,RSP}_regional.npy + artifact masks")
    bullet(doc, "session_{N}/ — per-session detail tables, psd_cache.npz, csd_cache.npz")
    bullet(doc, "spectrograms/ — per-session per-state event-aligned spectrograms + aggregates")
    code(doc, "figures/HMM/neural_alignment/lfp_spectral/")
    bullet(doc, "M2_aggregate_state_{N}_{ACA,LHA}.png — cross-session event-aligned spectrograms")
    bullet(doc, "session_{N}/session_{N}_M2_state_{S}_{R}.png — per-session spectrograms")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Saved {OUT_DOCX}")


if __name__ == "__main__":
    main()
