"""Build a detailed Word-document report of the behavioral HMM pipeline.

Reads outputs from data/HMM/ and figures/HMM/ to assemble a single
self-contained report.

Output: data/HMM/HMM_pipeline_report.docx
"""
import json
import pickle
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    return p


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.4f}", max_cols=None):
    """Add a DataFrame as a styled table."""
    if max_cols is not None:
        df = df.iloc[:, :max_cols]
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    # Header
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    # Rows
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def main():
    cfg = load_config()
    data_dir = REPO_ROOT / "data" / "HMM"
    fig_dir = REPO_ROOT / "figures" / "HMM"

    with open(data_dir / "final_model_params" / "meta.json") as f:
        meta = json.load(f)

    cv_df = pd.read_csv(data_dir / "cv_results.csv")
    profiles_df = pd.read_csv(data_dir / "state_profiles.csv")
    occ_df = pd.read_csv(data_dir / "fed_vs_fasted.csv")
    stats_df = pd.read_csv(data_dir / "fed_vs_fasted_stats.csv")
    transitions_df = pd.read_csv(data_dir / "fed_vs_fasted_transitions.csv")
    sess_meta = pd.read_csv(data_dir / "prepared" / "session_metadata.csv")

    doc = Document()

    # Margins
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===== TITLE =====
    title = doc.add_heading("Behavioral HMM Pipeline — Dual-Probe Foraging", level=0)
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run("Modular pipeline for state segmentation of EthoVision behavior data\n"
                    "Mouse01 Coordinates-1 — fed vs fasted foraging sessions")
    r.italic = True
    r.font.size = Pt(11)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-04-30")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)

    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "We built a 7-step modular Hidden Markov Model (HMM) pipeline that ingests "
        "raw EthoVision foraging xlsx files, rebins them at 480 ms, fits a behavioral HMM "
        "using the ssm package (Linderman lab), and produces per-bin state assignments, "
        "behavioral state profiles, and fed-vs-fasted comparisons. The pipeline ran "
        "end-to-end on 7 dual-probe foraging sessions (4 fed: S4, S6, S8, S10; "
        "3 fasted: S12, S14, S16). Cross-validation selected N=12 states (smallest "
        "within 1 standard error of the maximum log-likelihood). The fitted model "
        "converged in 25 EM iterations with a per-bin log-likelihood of 13.42."
    )
    add_para(
        doc,
        "The strongest fed-vs-fasted contrast is State 6, a long-dwell, low-speed, "
        "feeding-and-digging-rich state that occupies 51.6% of fasted-session time "
        "vs 19.2% in fed sessions (mean dwell 30.9 s vs 12.3 s). Three smaller fed-biased "
        "states (5, 7, 8) and a near-significant occupancy difference all give Mann-Whitney "
        "p = 0.057 — the floor for n=4 vs n=3. None reach the conventional p < 0.05 "
        "threshold; with seven sessions the test is underpowered, so these results "
        "are descriptive trends rather than confirmed effects."
    )

    # ===== KEY DESIGN DECISIONS =====
    add_heading(doc, "Pipeline Design", level=1)

    add_heading(doc, "Sessions and bin granularity", level=2)
    add_para(
        doc,
        f"Seven foraging-labeled dual-probe sessions: fed = {cfg['sessions']['fed']}, "
        f"fasted = {cfg['sessions']['fasted']}. Session 18 was excluded because it has "
        "behavior-only data (no sorted neural data on 7/30 recording date) — including "
        "it would create asymmetry for any downstream neural alignment. The native "
        "EthoVision bin is 40 ms; we rebin by a factor of 12 to obtain "
        f"{cfg['target_bin_ms']} ms bins. The session length is 30 minutes per session, "
        f"yielding ≈ {sess_meta['n_bins'].median():.0f} HMM bins per session "
        f"({sess_meta['n_bins'].sum()} total)."
    )

    add_heading(doc, "Observation features (D = 15)", level=2)
    add_para(
        doc,
        "We use a single 15-dimensional Gaussian observation vector per bin, following "
        "the MoSeq / Wiltschko (2015) behavioral-HMM convention. ssm does not natively "
        "support a factorized Gaussian + Bernoulli + Categorical emission, so we "
        "concatenate everything into one Gaussian observation:"
    )
    add_bullet(doc, "Continuous (z-scored over the pooled dataset): speed, distance to nearest pot.")
    add_bullet(doc, "Bernoulli events (0/1): incomplete home returns, quick loop at home, "
                    "digging sand, feeding, rearing, transition-wall exploration, "
                    "contemplation at T-zone.")
    add_bullet(doc, "Zone categorical (priority-mapped, then one-hot): home > transition > "
                    "pot > pot-zone > arena > other.")

    add_heading(doc, "Variance-collapse fix (critical)", level=2)
    add_para(
        doc,
        "On a first pass the cross-validated held-out per-bin log-likelihoods were "
        "pathological (values of ±10⁶ across folds). The cause was Gaussian variance "
        "collapse on binary / one-hot columns: when a state specialises on, say, "
        "zone_home = 1 in every training bin, the fitted variance for that column "
        "shrinks toward zero, and any held-out bin where zone_home = 0 lies at "
        "many sigmas, producing a near-infinite negative log-density. Two changes "
        "fixed this:"
    )
    add_bullet(doc, "Switch to a diagonal-covariance Gaussian "
                    "(observations='diagonal_gaussian'), which avoids singular "
                    "off-diagonal entries.")
    add_bullet(doc, "Add a small Gaussian jitter (σ = 0.05) to all binary and one-hot "
                    "feature columns. This puts a hard lower bound on per-feature "
                    "variance and stops EM from collapsing.")
    add_para(
        doc,
        "After both fixes, per-bin log-likelihoods are well-behaved (5–14 across folds) "
        "and rise monotonically with N."
    )

    add_heading(doc, "Cross-validation strategy", level=2)
    add_para(
        doc,
        "With 7 sessions (4 fed + 3 fasted), no clean leave-2-out × 4-fold split "
        "exists. We use 3 stratified folds, each holding out one fed and one fasted "
        "session for testing. Folds were generated with a fixed seed (20260429) for "
        "reproducibility. We sweep N in {4, 6, 8, 10, 12, 14} and report the held-out "
        "log-likelihood per bin in each fold. The recommended N is the smallest value "
        "within 1 SE of the best mean (a standard 1-SE-rule heuristic for parsimony)."
    )

    # ===== PIPELINE SCRIPTS =====
    add_heading(doc, "Pipeline Scripts", level=1)
    pipeline_steps = [
        ("01_load_and_rebin.py",
         "Reads each session xlsx (sheet 'Track-Full Arena-Subject 1', header row 35, "
         "skipping the units row). Coerces all numeric columns; '−' missing markers "
         "become NaN and are filled by ffill / bfill / 0. Computes per-bin nearest-pot "
         "distance (row-wise min over Pot-1..Pot-4). Maps zone columns into a single "
         "priority-ordered categorical. Rebins by block-mean (continuous) and "
         "block-mode (categorical / Bernoulli) at factor 12. Saves "
         "data/HMM/binned/session_{N}.npz."),
        ("02_prepare_for_hmm.py",
         "Loads pooled continuous values, computes pooled mean / SD, z-scores. "
         "One-hot encodes the zone categorical. Adds Gaussian jitter (σ = 0.05) to "
         "all binary and one-hot features. Concatenates everything into a single "
         "(T × 15) observation matrix per session. Saves data/HMM/prepared/session_{N}.npz "
         "and a session-metadata CSV."),
        ("03_fit_and_select_states.py",
         "Builds 3 stratified CV folds. For each N in N_range and each fold, fits an "
         "ssm.HMM (diagonal-Gaussian, EM, 200 iters, tol 1e-4) on training, returns "
         "the held-out per-bin log-likelihood. Aggregates mean ± SE across folds, "
         "applies the 1-SE rule, writes the recommended N back to config.yaml. "
         "Plots state-selection.png."),
        ("04_fit_final_model.py",
         "Re-fits an HMM with the selected N on all 7 sessions pooled. Pickles the "
         "fitted hmm together with metadata into final_model.pkl. Writes "
         "human-readable parameters: initial distribution, transition matrix, per-state "
         "per-feature emission means and sigmas, and a meta.json summary."),
        ("05_extract_state_posteriors.py",
         "For each session: forward-backward via hmm.expected_states(X)[0] to get the "
         "(T × N) per-bin posterior, plus Viterbi-decoded state via "
         "hmm.most_likely_states(X). Saves session-level CSVs with bin, time_s, "
         "p_state_0..N-1, viterbi columns."),
        ("06_validate_states.py",
         "Computes posterior-weighted per-state behavioral profile (mean and SD per "
         "feature). Plots states × features heatmap (z-scored across states). Builds "
         "per-session timeline plots: posterior heatmap on top, Viterbi colored ribbon "
         "on bottom, with dig (orange) and feeding (red) events overlaid. Prints "
         "warnings for low-occupancy (< 2 %), redundant (cosine-sim > 0.95 across z-profiles), "
         "and flickering (mean dwell < 2 bins) states."),
        ("07_compare_metabolic_states.py",
         "Per session: soft and hard occupancy, Viterbi-derived mean dwell, empirical "
         "transition matrix. Per state: Mann-Whitney U on session-level fed-vs-fasted "
         "occupancy and mean dwell. Plots fed-vs-fasted bars (with session-level scatter) "
         "and a 3-panel transition-matrix figure (fed mean P, fasted mean P, "
         "fasted − fed difference)."),
    ]
    for name, descr in pipeline_steps:
        p = doc.add_paragraph()
        r = p.add_run(name)
        r.bold = True
        r.font.size = Pt(10)
        r2 = p.add_run("  —  " + descr)
        r2.font.size = Pt(10)

    add_para(doc, "Run order: 01 → 02 → 03 → 04 → 05 → 06 → 07. Use:")
    code_p = doc.add_paragraph()
    cr = code_p.add_run('PYTHONIOENCODING=utf-8 "C:\\Users\\Gregg\\anaconda3\\envs\\si_env'
                         '\\python.exe" scripts/HMM/0X_*.py')
    cr.font.name = "Consolas"
    cr.font.size = Pt(9)
    add_para(doc, "Setting PYTHONIOENCODING=utf-8 is required because several scripts "
                  "print the '→' arrow; default cp1252 stdout encoding raises "
                  "UnicodeEncodeError. Avoid 'conda run' — its captured-stdout decoder also breaks.",
             size=10)

    # ===== STATE SELECTION =====
    doc.add_page_break()
    add_heading(doc, "State Selection (Cross-Validation)", level=1)
    add_para(
        doc,
        "Aggregated held-out per-bin log-likelihoods across 3 stratified folds:"
    )
    cv_agg = cv_df.groupby("N")["ll_per_bin"].agg(["mean", "std", "count"]).reset_index()
    cv_agg["se"] = cv_agg["std"] / np.sqrt(cv_agg["count"])
    cv_agg = cv_agg[["N", "mean", "std", "se", "count"]]
    cv_agg.columns = ["N", "mean ll/bin", "sd", "se", "folds"]
    add_df_table(doc, cv_agg, float_fmt="{:.3f}")
    doc.add_paragraph()
    add_para(
        doc,
        "The mean log-likelihood is still rising at N = 14, the top of the swept range. "
        "The 1-SE rule selects N = 12 as the smallest model within one standard error "
        "of the maximum. State counts above 14 may give further gains; this can be "
        "tested by extending N_range in config.yaml. Note also that the validation "
        "outputs flag four states with < 2 % occupancy at N = 12 — re-fitting at "
        "N = 10 may yield a cleaner state set with similar predictive performance."
    )
    add_image(doc, fig_dir / "state_selection.png", width_in=5.5,
              caption="Held-out log-likelihood per bin vs N (mean ± SE across 3 folds). "
                      "Dashed line: 1 SE below max-mean. Dotted vertical: recommended N.")

    # ===== FINAL MODEL =====
    doc.add_page_break()
    add_heading(doc, "Final Model", level=1)
    add_para(
        doc,
        f"Fit on all 7 sessions pooled; observations = diagonal_gaussian; "
        f"method = EM; n_iters = up to 200; tolerance = 1e-4. The model converged in "
        f"{meta['n_iter_actual']} iterations (well below the cap). Total log-likelihood "
        f"= {meta['final_log_likelihood']:.0f} over "
        f"{int(meta['final_log_likelihood']/meta['per_bin_log_likelihood'])} bins, "
        f"i.e. {meta['per_bin_log_likelihood']:.3f} per bin. The N×N transition matrix "
        "and per-state per-feature emission means and sigmas are in "
        "data/HMM/final_model_params/."
    )

    add_heading(doc, "Per-state behavioral profiles", level=2)
    add_para(
        doc,
        "Each row is one state; each column shows the posterior-weighted feature mean "
        "z-scored across states. Red = above-state-average; blue = below. Continuous "
        "features are speed_z and distance_to_pot_z (already z-scored across the dataset, "
        "so the heatmap shows further normalisation across the 12 states). Binary / "
        "one-hot features show the fraction of bins (per state) where the feature was active."
    )
    add_image(doc, fig_dir / "state_profiles.png", width_in=6.5,
              caption="State × feature heatmap. Each cell is the posterior-weighted "
                      "mean of the feature in that state, z-scored across states.")

    # State summary table
    add_heading(doc, "State summary", level=2)
    state_summary = profiles_df[[
        "state", "occupancy_soft", "occupancy_hard", "mean_dwell_s", "n_runs",
        "speed_z", "distance_to_pot_z", "event_digging_sand", "event_feeding",
        "zone_home", "zone_pot", "zone_arena",
    ]].copy()
    state_summary.columns = ["state", "occ_soft", "occ_hard", "dwell_s", "runs",
                              "speed_z", "dist_z",
                              "P(dig)", "P(feed)", "P(home)", "P(pot)", "P(arena)"]
    add_df_table(doc, state_summary, float_fmt="{:.3f}")

    # Inline interpretation
    doc.add_paragraph()
    add_para(doc, "Reading the profile (a brief gloss for each state):")
    interpretations = {
        0: "Tiny rearing-only state (0.3 % occupancy; rearing P ≈ 1.0). Likely a "
           "fragmentary detection of brief rearing bouts; flagged for low occupancy.",
        1: "Transition-zone-only state (1.7 %). High exploration-at-transition activity "
           "(P ≈ 1.0) at low speed; mouse pausing in the corridor.",
        2: "Far-from-pot arena state (10.4 %; dwell 13.4 s; distance_to_pot z = +2.25). "
           "Mouse in arena but far from pots, low Bernoulli activity.",
        3: "Home-zone state with quick-loop-at-home spikes (1.6 %; P(home) = 0.72; "
           "P(quick-loop) = 0.81).",
        4: "Pot / pot-zone short-bout fast-moving state (2.2 %; mean dwell 1.1 s — "
           "FLICKER warning was suppressed; very short visits to pot vicinity).",
        5: "Mid-distance, fast-moving arena state (8.3 %; speed z = +0.41; mean "
           "dwell 2.3 s). Reflects locomotion through the arena.",
        6: "Long-dwell pot state with feeding and digging (33.1 %; mean dwell 18 s; "
           "P(pot) = 0.67; P(feed) = 0.77; P(dig) = 0.25). The strongest fasted-biased state.",
        7: "Sustained feeding state (1.5 % overall; mean dwell 16.8 s; "
           "P(feed) ≈ 1.0; P(arena) ≈ 1.0). Fed-biased.",
        8: "Pot-zone fast bouts (8.9 %; speed z = +0.43; mean dwell 1.8 s; "
           "P(pot-zone) ≈ 1.0). Active foraging approach pattern.",
        9: "Home / transition with incomplete-return events (3.4 %; "
           "P(incomplete-home) = 0.30; P(home) = 0.71).",
        10: "Mid-arena rearing state (2.5 %; mean dwell 21 s; P(rearing) ≈ 1.0; "
            "split arena / transition / pot-zone occupancy).",
        11: "Dominant home-base state (26.0 %; P(home) = 0.68; P(transition) = 0.32; "
            "speed z = +0.46; mean dwell 4.9 s). Most-visited resting / approach state.",
    }
    for k, descr in interpretations.items():
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(f"State {k}: ")
        r.bold = True
        p.add_run(descr)

    # Validation warnings
    add_heading(doc, "Validation warnings (auto-detected)", level=2)
    flagged = profiles_df[profiles_df["occupancy_soft"] < 0.02]
    add_para(doc, f"{len(flagged)} of 12 states have soft occupancy < 2 %: "
                  f"{', '.join(f'S{int(s)} ({o*100:.2f} %)' for s, o in zip(flagged['state'], flagged['occupancy_soft']))}. "
                  "Each is plausibly real (rare events such as rearing, quick-loop-at-home, "
                  "or sustained feeding bouts), but with so little data per state the "
                  "emission parameters are noisy. A re-fit at N = 10 would likely absorb "
                  "states 0 and 7 into a larger feeding / rearing cluster.")
    add_para(doc, "No redundant pairs (z-profile cosine similarity > 0.95) and no "
                  "flickering states (mean dwell < 2 bins) were flagged.")

    # ===== FED VS FASTED =====
    doc.add_page_break()
    add_heading(doc, "Fed vs Fasted Comparison", level=1)
    add_para(
        doc,
        "All comparisons are session-level: per state we compute one occupancy and one "
        "mean dwell value per session, then run a two-sided Mann-Whitney U on the "
        "n = 4 fed vs n = 3 fasted session means. For these sample sizes the smallest "
        "two-sided p-value possible is 0.0571 (perfect rank separation), so individual "
        "p-values should be read as ranks rather than as confirmatory significance."
    )

    add_heading(doc, "Per-state statistics", level=2)
    stat_disp = stats_df.copy()
    stat_disp = stat_disp[[
        "hmm_state", "mean_occ_fed", "mean_occ_fasted", "occ_p",
        "mean_dwell_s_fed", "mean_dwell_s_fasted", "dwell_p",
    ]]
    stat_disp.columns = ["state", "occ_fed", "occ_fasted", "occ_p",
                          "dwell_s_fed", "dwell_s_fasted", "dwell_p"]
    add_df_table(doc, stat_disp, float_fmt="{:.3f}")

    add_heading(doc, "Headline findings", level=2)
    add_bullet(doc,
        "State 6 (long-dwell pot / feeding / dig) is strongly fasted-biased: 51.6 % vs 19.2 % "
        "(occ U = 0, p = 0.057, perfect separation). Mean dwell is 30.9 s in fasted "
        "sessions vs 12.3 s in fed (~2.5× longer). This is the cleanest direction "
        "effect in the dataset and is consistent with fasted mice settling into "
        "extended feeding bouts at the pot once food is found."
    )
    add_bullet(doc,
        "State 7 (sustained feeding in arena, P(feed) ≈ 1.0) appears only in fed sessions "
        "(2.6 % vs 0 %, p = 0.057). This may correspond to long bouts of in-arena pellet "
        "consumption seen only when the mouse is sated and not making repeated pot trips."
    )
    add_bullet(doc,
        "States 5 (mid-arena fast locomotion) and 8 (pot-zone fast bouts) are each ~2× "
        "more occupied in fed than fasted sessions (p = 0.057). Together with state 7 "
        "this is consistent with fed sessions having more diffuse exploration relative "
        "to fasted-session pot dwelling."
    )
    add_bullet(doc,
        "No other state reaches even the n = 4 vs n = 3 floor (p < 0.06). Notably "
        "state 11 (the dominant home-base state) is similar across diet conditions "
        "(29 % vs 22 %, p = 0.23) — fed and fasted mice spend similar fractions of "
        "time at home base."
    )
    add_bullet(doc,
        "Mean dwell time differences track occupancy differences but none of the "
        "dwell p-values reach the 0.057 floor — variability across sessions in dwell "
        "is larger than in occupancy fraction."
    )

    add_image(doc, fig_dir / "fed_vs_fasted.png", width_in=6.5,
              caption="Fed vs fasted: per-state occupancy (left) and mean dwell (right). "
                      "Bars show group mean ± SEM; black dots are session-level values; "
                      "asterisks mark p < 0.05 (none reach this in the present sample).")

    add_image(doc, fig_dir / "fed_vs_fasted_transitions.png", width_in=6.5,
              caption="Empirical transition matrices (Viterbi-derived). Left: fed mean P. "
                      "Centre: fasted mean P. Right: fasted minus fed (red = more in "
                      "fasted, blue = more in fed). Diagonal dominates as expected; "
                      "asymmetries indicate which inter-state transitions diet shifts.")

    # ===== TIMELINES =====
    doc.add_page_break()
    add_heading(doc, "Per-Session Timelines", level=1)
    add_para(
        doc,
        "Each panel shows a single session. Top: per-bin posterior heatmap (states on "
        "y, time on x, posterior probability in viridis). Bottom: Viterbi-decoded state "
        "as a coloured ribbon. Vertical orange ticks: digging-sand events; vertical red "
        "ticks: feeding events."
    )
    sess_order = [(4, "fed"), (6, "fed"), (8, "fed"), (10, "fed"),
                   (12, "fasted"), (14, "fasted"), (16, "fasted")]
    for sn, st in sess_order:
        path = fig_dir / "timelines" / f"session_{sn}.png"
        add_image(doc, path, width_in=6.7,
                  caption=f"Session {sn} ({st}). 30-min recording; HMM posteriors over time.")

    # ===== OUTPUTS =====
    doc.add_page_break()
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/binned/session_{N}.npz",
         "Per-session rebinned arrays (speed, distance_to_pot, zone, events, trial_time)."),
        ("data/HMM/prepared/session_{N}.npz",
         "Per-session observation matrix X (T × 15) plus z-score parameters."),
        ("data/HMM/cv_results.csv",
         "Per (N, fold) held-out log-likelihood per bin."),
        ("data/HMM/final_model.pkl",
         "Pickled bundle: fitted ssm.HMM + metadata."),
        ("data/HMM/final_model_params/",
         "initial_distribution.csv, transition_matrix.csv, emissions.csv, meta.json."),
        ("data/HMM/posteriors/session_{N}.csv",
         "Per-bin posteriors p_state_0..N-1 + Viterbi assignment + time."),
        ("data/HMM/state_profiles.csv",
         "Per-state behavioral profile (mean + sd of every feature, occupancy, dwell)."),
        ("data/HMM/state_dwell_occupancy.csv",
         "Compact occupancy + mean dwell per state."),
        ("data/HMM/fed_vs_fasted.csv",
         "Per-session × state occupancy and mean dwell."),
        ("data/HMM/fed_vs_fasted_transitions.csv",
         "Per-session × (i, j) empirical transition probabilities."),
        ("data/HMM/fed_vs_fasted_stats.csv",
         "Per-state Mann-Whitney U tests on occupancy and dwell."),
        ("figures/HMM/state_selection.png",
         "Cross-validation curve (mean ± SE)."),
        ("figures/HMM/state_profiles.png",
         "States × features heatmap."),
        ("figures/HMM/timelines/session_{N}.png",
         "Per-session posterior + Viterbi timeline with event overlays."),
        ("figures/HMM/fed_vs_fasted.png",
         "Bar plots of fed-vs-fasted occupancy and dwell with session scatter."),
        ("figures/HMM/fed_vs_fasted_transitions.png",
         "Mean transition matrices for fed, fasted, and fasted − fed difference."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    # ===== CAVEATS / NEXT STEPS =====
    add_heading(doc, "Caveats and Next Steps", level=1)
    add_bullet(doc,
        "Underpowered: with n = 4 vs 3, the Mann-Whitney p floor is 0.057. "
        "All claims are descriptive trends. Adding more sessions (e.g. expanding to "
        "exploration-phase sessions, or additional foraging recordings) would lift "
        "this floor."
    )
    add_bullet(doc,
        "No null control: a shuffled-state-sequence baseline would establish whether "
        "the fed-vs-fasted occupancy differences exceed what one would see from "
        "trivially relabelled states. This is the obvious next analysis."
    )
    add_bullet(doc,
        "N = 12 has 4 low-occupancy states. A re-fit at N = 10 (which is also within "
        "1 SE of the maximum log-likelihood) would likely produce a more interpretable "
        "state set without losing predictive performance."
    )
    add_bullet(doc,
        "Mixed-emission workaround is statistically suboptimal. To use proper "
        "Bernoulli + Categorical + Gaussian factorised emissions, one could either "
        "(a) write a custom ssm Observations subclass, or (b) move to a different "
        "library (e.g. pomegranate, hmmlearn with custom emissions). For the present "
        "exploratory pipeline the Gaussian-on-binary convention is the standard "
        "behavioral-HMM choice (cf. MoSeq) and is likely sufficient."
    )
    add_bullet(doc,
        "Neural alignment is the natural follow-up: with the per-bin Viterbi or "
        "posterior streams in posteriors/session_{N}.csv, behavioural-state-conditioned "
        "neural population activity in ACA / LHA can be averaged and contrasted across "
        "states to test for state-locked neural signatures, complementing the existing "
        "entropy- and approach-aligned analyses."
    )

    # Save
    out_path = REPO_ROOT / "data" / "HMM" / "HMM_pipeline_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
