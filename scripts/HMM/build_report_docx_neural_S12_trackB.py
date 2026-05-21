"""Build a Word-document report of Track B (state-conditioned) neural results for S12.

Reads outputs from data/HMM/neural_alignment/state_conditioned_S12/ and figures
from figures/HMM/neural_alignment/state_conditioned_S12/.

Output: data/HMM/neural_alignment_S12_trackB_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.3f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    cfg = load_config()
    cm_dir = REPO_ROOT / cfg["commitment_dirs"]["out"]
    na_dir = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "state_conditioned_S12"
    fig_dir = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "state_conditioned_S12"

    history = pd.read_csv(cm_dir / "sampling_history.csv")
    s12 = history[history.session == 12].iloc[0]

    # ---- Load Track B CSVs ----
    sel = pd.read_csv(na_dir / "B1_selectivity_summary.csv")
    b2_aca = pd.read_csv(na_dir / "B2_pre_vs_post_per_state_ACA.csv")
    b2_lha = pd.read_csv(na_dir / "B2_pre_vs_post_per_state_LHA.csv")
    b3_aca = pd.read_csv(na_dir / "B3_glm_coefficients_ACA.csv")
    b3_lha = pd.read_csv(na_dir / "B3_glm_coefficients_LHA.csv")
    b3_summary = pd.read_csv(na_dir / "B3_significant_coefficients_summary.csv")
    shifts_aca = pd.read_csv(na_dir / "B4_pre_post_centroid_shift_ACA.csv")
    shifts_lha = pd.read_csv(na_dir / "B4_pre_post_centroid_shift_LHA.csv")

    K = int(sel["preferred_state"].max()) + 1

    # ---- Aggregate stats per region ----
    def b1_stats(sub):
        return dict(
            n_units=len(sub),
            n_uncorr=int(sub["sig_uncorr"].sum()),
            n_fdr=int(sub["sig_fdr"].sum()),
            pref_counts=sub["preferred_state"].value_counts().sort_index().to_dict(),
        )

    sel_aca = sel[sel.region == "ACA"]
    sel_lha = sel[sel.region == "LHA"]
    s_b1_aca = b1_stats(sel_aca)
    s_b1_lha = b1_stats(sel_lha)

    def b2_stats(df):
        n_units_total = df["unit_id"].nunique()
        sig_units = df.loc[df["sig_fdr"], "unit_id"].nunique()
        per_state = df[df["sig_fdr"]].groupby("state").size()
        return dict(n_units=n_units_total, sig_units=int(sig_units),
                    per_state=per_state.reindex(np.arange(K), fill_value=0).to_dict())

    s_b2_aca = b2_stats(b2_aca)
    s_b2_lha = b2_stats(b2_lha)

    def b3_stats(df):
        n_units_total = df["unit_id"].nunique()
        sig_units = df.loc[df["sig"], "unit_id"].nunique()
        per_state = df[df["sig"]].groupby("state").size()
        return dict(n_units=n_units_total, sig_units=int(sig_units),
                    per_state=per_state.reindex(np.arange(K), fill_value=0).to_dict())

    s_b3_aca = b3_stats(b3_aca)
    s_b3_lha = b3_stats(b3_lha)

    # ---- Build doc ----
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_heading(
        "Neural Alignment Track B — State-Conditioned (Session 12, fasted)", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "State-selectivity (B1), pre/post within-state shift (B2), Poisson GLM "
        "on state posteriors (B3), and PCA trajectories colored by state (B4) "
        "for ACA and LHA in S12."
    )
    r.italic = True
    r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-06  •  Script: scripts/HMM/10b_neural_alignment_state_conditioned_S12.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        f"Track B asks whether the K={K} merged HMM states are encoded in ACA "
        "and LHA single-unit firing rates and population structure across the "
        f"whole 30-min session 12 (fasted, food P{int(s12['food_pot'])}, "
        f"discovery t={s12['discovery_time_s']:.1f} s). Neural rates are re-binned "
        "from Track A's 100 ms grid to the 480 ms HMM bin grid by averaging "
        "100 ms counts within each HMM bin and converting to Hz."
    )
    add_bullet(
        doc,
        f"State encoding is pervasive: B1 ANOVA flags "
        f"{s_b1_aca['n_fdr']}/{s_b1_aca['n_units']} ACA units and "
        f"{s_b1_lha['n_fdr']}/{s_b1_lha['n_units']} LHA units as state-selective "
        "after FDR. The Poisson GLM (B3) detects ≥1 significant state "
        f"coefficient in {s_b3_aca['sig_units']}/{s_b3_aca['n_units']} ACA and "
        f"{s_b3_lha['sig_units']}/{s_b3_lha['n_units']} LHA units. Essentially "
        "every recorded unit cares about state."
    )
    add_bullet(
        doc,
        f"Single-state preference: 33/{s_b1_lha['n_units']} (37%) LHA units "
        "prefer merged S2 (pure-feeding, P(feeding)=1.0) — the largest "
        "concentration of preference on any single state in either region. "
        "ACA preferred states are more spread; the largest single bin is S6 "
        f"(28/{s_b1_aca['n_units']} units, the pure-digging state)."
    )
    add_bullet(
        doc,
        "B2 within-state pre/post shifts: "
        f"{s_b2_aca['sig_units']}/{s_b2_aca['n_units']} ACA units "
        f"({s_b2_aca['sig_units']/s_b2_aca['n_units']*100:.0f}%) and "
        f"{s_b2_lha['sig_units']}/{s_b2_lha['n_units']} LHA units "
        f"({s_b2_lha['sig_units']/s_b2_lha['n_units']*100:.0f}%) show FDR-sig "
        "firing-rate change in ≥1 state when comparing pre- vs post-discovery "
        "bins of the same state. Even when the mouse is in 'the same' "
        "behavioral state, neural activity differs once the food location is "
        "known."
    )
    add_bullet(
        doc,
        "B4 PCA: top-5 components capture 22.6% (ACA) / 19.4% (LHA) of "
        "variance. Pre/post discovery centroid shifts in PC1-PC2-PC3 are "
        "largest in pot-zone states **S8 and S9** (1.5-1.6 units in both "
        "regions). The same pot-zone behavior occupies different neural "
        "subspaces before vs after the mouse knows where food is — the "
        "cleanest candidate signature of cognitive context in this session."
    )
    add_bullet(
        doc,
        "All four analyses are descriptive on a single session. The pervasive "
        "state-selectivity result needs a shuffle control before interpretation; "
        "the pot-zone pre/post centroid shifts are the most defensible "
        "candidate finding pending replication across sessions."
    )

    # ===== SETUP =====
    add_heading(doc, "Setup", level=1)
    add_bullet(doc, "Same QC-filtered units as Track A: 165 ACA (probe 0, KSLabel='good', "
                    "FR>0.2 Hz) and 89 LHA (probe 1, KSLabel='good', FR>0.2 Hz, "
                    "AMP>43 µV, depth 0-345 µm).")
    add_bullet(doc, "Neural binning: 480 ms (matched to HMM behavior bin), aggregated by "
                    "averaging the 100 ms count series from Track A. Spike-count series "
                    "for the GLM are re-binned directly at 480 ms from spike times.")
    add_bullet(doc, f"HMM source: K={K} merged Viterbi/posteriors at "
                    f"data/HMM/merged_posteriors_dynamax/session_12.csv "
                    f"(3750 bins × 1800 s).")
    add_bullet(doc, f"Discovery cutoff: bin {int(s12['discovery_bin'])}, "
                    f"t={s12['discovery_time_s']:.1f} s. "
                    "Pre-discovery: 1244 bins; post-discovery: 2506 bins.")
    add_bullet(doc, "Multiple-comparison correction: BH-FDR within region for B1, B2, B3.")

    # ===== B1 =====
    doc.add_page_break()
    add_heading(doc, "B1 — Per-unit state selectivity (one-way ANOVA)", level=1)
    add_para(
        doc,
        "For each unit, the firing rates across the bins assigned to each of "
        f"the K={K} merged HMM states form K groups; we run a one-way ANOVA "
        "across those groups. A unit is 'state-selective' if it varies "
        "significantly with state. Heatmap rows are z-scored per unit so "
        "high-firing units don't dominate the visualization."
    )
    b1_tbl = pd.DataFrame([
        dict(region="ACA", n_units=s_b1_aca["n_units"],
             sig_uncorr=f"{s_b1_aca['n_uncorr']}/{s_b1_aca['n_units']}",
             sig_FDR=f"{s_b1_aca['n_fdr']}/{s_b1_aca['n_units']}",
             pct_FDR=f"{s_b1_aca['n_fdr']/s_b1_aca['n_units']*100:.1f}%"),
        dict(region="LHA", n_units=s_b1_lha["n_units"],
             sig_uncorr=f"{s_b1_lha['n_uncorr']}/{s_b1_lha['n_units']}",
             sig_FDR=f"{s_b1_lha['n_fdr']}/{s_b1_lha['n_units']}",
             pct_FDR=f"{s_b1_lha['n_fdr']/s_b1_lha['n_units']*100:.1f}%"),
    ])
    add_df_table(doc, b1_tbl)

    add_para(doc, "Preferred-state distribution (count of units assigning each state as "
                  "their preferred — highest mean firing rate — state):", size=10)
    pref_rows = []
    for k in range(K):
        pref_rows.append(dict(
            state=f"S{k}",
            ACA=s_b1_aca["pref_counts"].get(k, 0),
            LHA=s_b1_lha["pref_counts"].get(k, 0),
        ))
    add_df_table(doc, pd.DataFrame(pref_rows))
    add_bullet(doc, "ACA preferences are spread across all 13 visited states. The biggest "
                    f"single concentrations are S6 ({s_b1_aca['pref_counts'].get(6,0)}, "
                    "pure-digging) and S5/S2/S4 (~20 each).")
    add_bullet(doc, f"LHA preferences pile heavily on S2 (pure-feeding): "
                    f"{s_b1_lha['pref_counts'].get(2,0)}/{s_b1_lha['n_units']} units "
                    f"({s_b1_lha['pref_counts'].get(2,0)/s_b1_lha['n_units']*100:.0f}%) "
                    "have S2 as their preferred state. Consistent with LHA's known role "
                    "in feeding/consummatory behavior.")
    add_bullet(doc, "S11 was not preferred by any unit in either region — that small "
                    "feeding state is sub-dominant to S2 in driving unit activity.")

    add_image(doc, fig_dir / "B1_heatmap_ACA.png", width_in=6.7,
              caption="ACA: rows = units sorted by preferred state; columns = "
                      "K=14 merged HMM states; values = z-scored mean FR per state "
                      "(red = above unit's average, blue = below).")
    add_image(doc, fig_dir / "B1_heatmap_LHA.png", width_in=6.7,
              caption="LHA: same layout. Note the visible block of units with strong "
                      "S2 preference (large red column).")
    add_image(doc, fig_dir / "B1_state_preference_counts.png", width_in=6.7,
              caption="Per-region histogram of preferred states. Red overlay marks "
                      "FDR-significant units (essentially every unit). Grey numbers "
                      "in parentheses = total bins assigned to that state in S12.")

    # ===== B2 =====
    doc.add_page_break()
    add_heading(doc, "B2 — Pre vs post-discovery within-state firing rate", level=1)
    add_para(
        doc,
        "For each unit and each state with ≥30 bins on both sides of discovery, "
        "Mann-Whitney U on bin-level firing rates pre vs post. Tests whether "
        "the SAME behavioral state evokes different neural activity once the "
        "mouse knows where food is, holding behavior approximately constant."
    )
    b2_tbl = pd.DataFrame([
        dict(region="ACA",
             n_units=s_b2_aca["n_units"],
             units_with_any_FDR_sig=f"{s_b2_aca['sig_units']}/{s_b2_aca['n_units']} "
                                     f"({s_b2_aca['sig_units']/s_b2_aca['n_units']*100:.0f}%)"),
        dict(region="LHA",
             n_units=s_b2_lha["n_units"],
             units_with_any_FDR_sig=f"{s_b2_lha['sig_units']}/{s_b2_lha['n_units']} "
                                     f"({s_b2_lha['sig_units']/s_b2_lha['n_units']*100:.0f}%)"),
    ])
    add_df_table(doc, b2_tbl)

    add_para(doc, "Per-state count of FDR-significant (unit, state) pairs:", size=10)
    per_state_rows = []
    for k in range(K):
        per_state_rows.append(dict(
            state=f"S{k}",
            ACA=s_b2_aca["per_state"].get(k, 0),
            LHA=s_b2_lha["per_state"].get(k, 0),
        ))
    add_df_table(doc, pd.DataFrame(per_state_rows))
    add_bullet(doc, "States with the most pre/post-modulated units in BOTH regions: "
                    f"S2 ({s_b2_aca['per_state'].get(2,0)} ACA, "
                    f"{s_b2_lha['per_state'].get(2,0)} LHA — feeding) and "
                    f"S9 ({s_b2_aca['per_state'].get(9,0)} ACA, "
                    f"{s_b2_lha['per_state'].get(9,0)} LHA — pot-zone).")
    add_bullet(doc, "S0, S1, S5 don't have ≥30 bins in either pre or post window, so no "
                    "test is run for those states — they're absent from this table.")
    add_bullet(doc, "The 86%/83% positive rate is high but expected: with 165/89 units "
                    "× many states, FDR within region still lets through any genuine "
                    "shift. The interpretive value is the per-state distribution, not "
                    "the unit count.")

    add_image(doc, fig_dir / "B2_delta_heatmap_ACA.png", width_in=6.7,
              caption="ACA: ΔFR (post − pre) per unit per state. Rows sorted by mean "
                      "Δ across states; red = unit fires more post-discovery in that "
                      "state, blue = less.")
    add_image(doc, fig_dir / "B2_delta_heatmap_LHA.png", width_in=6.7,
              caption="LHA: same layout.")
    add_image(doc, fig_dir / "B2_significant_unit_counts.png", width_in=6.7,
              caption="Per-region count of units with FDR-significant pre/post change "
                      "in each state.")

    # ===== B3 =====
    doc.add_page_break()
    add_heading(doc, "B3 — Poisson GLM on state posteriors", level=1)
    add_para(
        doc,
        "For each unit, fit a Poisson GLM with spike count per HMM bin as the "
        "response, soft state posteriors as predictors, and log(bin width) as "
        "an offset. The most-occupied state is dropped as the reference (S2 "
        "for both ACA and LHA in S12). The β coefficients describe how much "
        "the unit's log firing rate increases when posterior mass moves from "
        "the reference state to state k."
    )
    b3_tbl = pd.DataFrame([
        dict(region="ACA", reference_state="S2",
             n_units=s_b3_aca["n_units"],
             units_with_any_sig=f"{s_b3_aca['sig_units']}/{s_b3_aca['n_units']} "
                                 f"({s_b3_aca['sig_units']/s_b3_aca['n_units']*100:.0f}%)"),
        dict(region="LHA", reference_state="S2",
             n_units=s_b3_lha["n_units"],
             units_with_any_sig=f"{s_b3_lha['sig_units']}/{s_b3_lha['n_units']} "
                                 f"({s_b3_lha['sig_units']/s_b3_lha['n_units']*100:.0f}%)"),
    ])
    add_df_table(doc, b3_tbl)
    add_para(doc, "Per-state count of significant (unit, state) coefficient pairs "
                  "(FDR q<0.05 AND |z|>2.5):", size=10)
    glm_rows = []
    for k in range(K):
        glm_rows.append(dict(
            state=f"S{k}",
            ACA=s_b3_aca["per_state"].get(k, 0),
            LHA=s_b3_lha["per_state"].get(k, 0),
        ))
    add_df_table(doc, pd.DataFrame(glm_rows))
    add_bullet(doc, "Every unit (165/165 ACA, 89/89 LHA) has at least one significant "
                    "state coefficient. Confirms B1: the GLM, with finer soft-posterior "
                    "predictors, picks up state encoding in every recorded cell.")
    add_bullet(doc, "S2 (the reference, dropped) is implicit in the intercept — "
                    "coefficients describe deviation from the feeding state. States with "
                    "the most universal modulation are S3, S4, S6, S7, S8, S9, S12 "
                    "(>120 ACA units each). S5 (74% feeding), S11 (other feeding) and "
                    "S0 give weaker coefficient counts — less differentiated from S2 by "
                    "the linear model.")

    add_image(doc, fig_dir / "B3_coefficient_heatmap_ACA.png", width_in=6.7,
              caption="ACA: Poisson GLM β coefficients per unit per state. "
                      "Reference state (S2) is dropped. Rows sorted by max |β|.")
    add_image(doc, fig_dir / "B3_coefficient_heatmap_LHA.png", width_in=6.7,
              caption="LHA: same layout.")

    # ===== B4 =====
    doc.add_page_break()
    add_heading(doc, "B4 — PCA with state coloring and pre/post overlay", level=1)
    add_para(
        doc,
        "PCA on z-scored unit firing rates (each row in the bins × units matrix "
        "is one HMM bin). Top-5 components capture 22.6% (ACA) and 19.4% (LHA) "
        "of total variance. Each bin is plotted as a point in PC space, "
        "colored by its Viterbi-assigned state. The pre/post overlay re-plots "
        "the same axes with pre-discovery bins faint and post-discovery bins "
        "solid, both colored by state — visualizes whether the same state "
        "occupies the same neural sub-space before vs after discovery."
    )
    var_tbl = pd.DataFrame([
        dict(region="ACA", PC1="8.6%", PC2="5.2%", PC3="3.5%", PC4="2.8%", PC5="2.4%",
             cumulative="22.6%"),
        dict(region="LHA", PC1="7.9%", PC2="3.4%", PC3="2.9%", PC4="2.8%", PC5="2.5%",
             cumulative="19.4%"),
    ])
    add_df_table(doc, var_tbl)
    add_para(doc, "Pre vs post-discovery centroid shift per state (Euclidean distance in "
                  "PC1-PC2-PC3 space; states with <5 bins on either side are NaN):", size=10)
    shift_tbl = pd.DataFrame(dict(
        state=[f"S{k}" for k in range(K)],
        ACA_shift=[
            (shifts_aca.loc[shifts_aca.state == k, "centroid_shift_PC123"].iloc[0]
             if (shifts_aca.state == k).any() else np.nan)
            for k in range(K)],
        ACA_n_pre=[
            (int(shifts_aca.loc[shifts_aca.state == k, "n_pre"].iloc[0])
             if (shifts_aca.state == k).any() else 0)
            for k in range(K)],
        ACA_n_post=[
            (int(shifts_aca.loc[shifts_aca.state == k, "n_post"].iloc[0])
             if (shifts_aca.state == k).any() else 0)
            for k in range(K)],
        LHA_shift=[
            (shifts_lha.loc[shifts_lha.state == k, "centroid_shift_PC123"].iloc[0]
             if (shifts_lha.state == k).any() else np.nan)
            for k in range(K)],
    ))
    add_df_table(doc, shift_tbl, float_fmt="{:.2f}")
    add_bullet(doc, "Largest pre/post centroid shifts: ACA — S0 (1.87), S8 (1.60), "
                    "S9 (1.58); LHA — S8 (1.61), S1 (1.48), S9 (1.48). **S8 and S9 "
                    "(both pot-zone states) appear in the top 3 in both regions.**")
    add_bullet(doc, "Interpretation: the same behavioral state — being in a pot-zone — "
                    "carries different neural content before vs after the mouse knows "
                    "the food location. This is the cleanest cognitive-state signature "
                    "in the session.")
    add_bullet(doc, "S0 has the largest ACA shift (1.87) but only 25 post-discovery "
                    "bins, so the centroid estimate is noisy. S8 and S9 have larger "
                    "samples (60-200 bins/side) and are more robust.")

    add_image(doc, fig_dir / "B4_pca_state_colored_ACA.png", width_in=6.7,
              caption="ACA: PC1 vs PC2 (left) and PC2 vs PC3 (right), colored by "
                      "Viterbi-assigned state. Each dot = one 480 ms bin.")
    add_image(doc, fig_dir / "B4_pca_state_colored_LHA.png", width_in=6.7,
              caption="LHA: same axes, colored by state.")
    add_image(doc, fig_dir / "B4_pca_pre_post_ACA.png", width_in=6.7,
              caption="ACA pre vs post-discovery overlay. Pre-discovery bins are "
                      "faint, post-discovery bins are solid; both colored by state. "
                      "Drift in same-color clouds across the pre/post axis is the "
                      "centroid shift.")
    add_image(doc, fig_dir / "B4_pca_pre_post_LHA.png", width_in=6.7,
              caption="LHA pre/post overlay.")

    # ===== CROSS-ANALYSIS =====
    doc.add_page_break()
    add_heading(doc, "Cross-Analysis Patterns", level=1)
    add_bullet(
        doc,
        "State-encoding is essentially universal in this session (B1: 99% FDR-"
        "significant; B3: 100%). The 'who codes state' question has a trivial "
        "answer: every unit. The interpretive weight shifts to (a) WHICH state "
        "each unit prefers, (b) HOW pre/post-discovery context modifies "
        "within-state firing.",
    )
    add_bullet(
        doc,
        "LHA-vs-ACA preference asymmetry is the most striking single-region "
        "feature: 37% of LHA units prefer the pure-feeding state S2 versus "
        "only 12% of ACA units. ACA preferences are more diffuse — multiple "
        "state types are equally attractive to ACA cells.",
    )
    add_bullet(
        doc,
        "Pot-zone states (S8, S9) carry the largest pre/post neural shift in "
        "both B2 (sig-unit counts: ~95-100 ACA units, ~30-40 LHA units across "
        "S2+S9+S8) and B4 (PC centroid shifts ~1.5-1.6 in both regions). The "
        "mouse's neural state during 'arriving at a pot' depends on whether "
        "it knows the food location — same behavior, different neural content.",
    )
    add_bullet(
        doc,
        "Track A (event-locked transients) showed at most ~10-25 Hz peak diffs "
        "on individual bins. Track B's whole-session encoding (B1, B3) is much "
        "stronger because it pools across all bins of each state, not just a "
        "±1 s window. Single-trial transient changes are small; tonic state-"
        "conditioned modulation is large. This suggests downstream analyses "
        "should weight tonic context heavily.",
    )

    # ===== CAVEATS =====
    add_heading(doc, "Caveats", level=1)
    add_bullet(doc, "Single session — every result is descriptive on n=1 mouse-session.")
    add_bullet(doc, "Pervasive significance (B1: 99%, B3: 100%) is partly a power "
                    "artifact: with 3750 bins and K=14 groups, even small per-state "
                    "differences are detectable. A SHUFFLE CONTROL (permute Viterbi "
                    "labels and re-run B1/B3) is the obvious follow-up before claiming "
                    "any unit is genuinely state-selective.")
    add_bullet(doc, "B2 and B3 share substrate (both test pre/post and state effects); "
                    "their counts are not independent.")
    add_bullet(doc, "B4 PCA top-5 captures ~20% variance — the rest of the variance "
                    "lives in higher dimensions. Pre/post centroid shifts in PC1-3 may "
                    "miss higher-dim shifts; a Mahalanobis or full-rank distance is "
                    "the natural extension.")
    add_bullet(doc, "Re-binning 100 ms → 480 ms via bin-center averaging introduces "
                    "minor jitter on the 0.1 s scale; for B-track analyses this is "
                    "negligible.")
    add_bullet(doc, "Multi-session replication (especially the LHA-S2 preference and "
                    "the S8/S9 pre/post shift) is the obvious next step.")

    # ===== OUTPUTS =====
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/neural_alignment/state_conditioned_S12/B1_state_selectivity_matrix_{ACA,LHA}.csv",
         "Per-unit firing rate per state (raw, units × states)."),
        ("data/HMM/neural_alignment/state_conditioned_S12/B1_state_selectivity_zscored_{ACA,LHA}.csv",
         "Same matrix, z-scored row-wise for visualization."),
        ("data/HMM/neural_alignment/state_conditioned_S12/B1_selectivity_summary.csv",
         "ANOVA results, preferred state, FDR-sig flags per unit."),
        ("data/HMM/neural_alignment/state_conditioned_S12/B2_pre_vs_post_per_state_{ACA,LHA}.csv",
         "Per-(unit, state) pre/post FR + Mann-Whitney p, FDR p, sig flag."),
        ("data/HMM/neural_alignment/state_conditioned_S12/B3_glm_coefficients_{ACA,LHA}.csv",
         "Per-(unit, state) Poisson GLM β, SE, z, p, p_FDR, sig flag."),
        ("data/HMM/neural_alignment/state_conditioned_S12/B3_significant_coefficients_summary.csv",
         "Per-unit list of significant states."),
        ("data/HMM/neural_alignment/state_conditioned_S12/B4_pca_loadings_{ACA,LHA}.csv",
         "Top-5 PC unit loadings."),
        ("data/HMM/neural_alignment/state_conditioned_S12/B4_centroid_distances_{ACA,LHA}.csv",
         "Pairwise PC1-3 centroid distances between states."),
        ("data/HMM/neural_alignment/state_conditioned_S12/B4_pre_post_centroid_shift_{ACA,LHA}.csv",
         "Per-state pre/post centroid shift in 3D PC space."),
        ("figures/HMM/neural_alignment/state_conditioned_S12/B1_heatmap_{ACA,LHA}.png",
         "Per-region z-scored state-selectivity heatmaps."),
        ("figures/HMM/neural_alignment/state_conditioned_S12/B1_state_preference_counts.png",
         "Per-region preferred-state count bar chart."),
        ("figures/HMM/neural_alignment/state_conditioned_S12/B2_delta_heatmap_{ACA,LHA}.png",
         "Per-region (post − pre) ΔFR heatmaps."),
        ("figures/HMM/neural_alignment/state_conditioned_S12/B2_significant_unit_counts.png",
         "Per-region per-state count of FDR-sig pre/post units."),
        ("figures/HMM/neural_alignment/state_conditioned_S12/B3_coefficient_heatmap_{ACA,LHA}.png",
         "Per-region GLM β heatmaps."),
        ("figures/HMM/neural_alignment/state_conditioned_S12/B4_pca_state_colored_{ACA,LHA}.png",
         "PC1-PC2 and PC2-PC3 plots colored by state."),
        ("figures/HMM/neural_alignment/state_conditioned_S12/B4_pca_pre_post_{ACA,LHA}.png",
         "Same axes with pre-discovery faint, post-discovery solid."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    out_path = REPO_ROOT / "data" / "HMM" / "neural_alignment_S12_trackB_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
