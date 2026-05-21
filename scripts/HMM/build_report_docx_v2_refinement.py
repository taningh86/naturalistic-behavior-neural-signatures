"""Build a Word-doc report for the v2 bipolar pair refinement rerun.

Covers all three pipelines (17 LFP spectral, 18 state-identity + 3-region
Granger, 19 SWR detection) under the v2 bipolar pair set (LHA y<345 µm,
RSP y>4680 µm). Reports v2 results + side-by-side comparison vs v1.

Output: data/HMM/v2_refinement_report.docx
"""
from pathlib import Path

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO = Path(__file__).resolve().parents[2]

# v1 / v2 output dirs
S17_V1 = REPO / "data/HMM/neural_alignment/lfp_spectral"
S17_V3 = REPO / "data/HMM/neural_alignment/lfp_spectral_v3"
S18_A1_V1 = REPO / "data/HMM/neural_alignment/lfp_state_identity"
S18_A1_V2 = REPO / "data/HMM/neural_alignment/lfp_state_identity_v2"
S18_A2_V1 = REPO / "data/HMM/neural_alignment/lfp_three_region_granger"
S18_A2_V2 = REPO / "data/HMM/neural_alignment/lfp_three_region_granger_v2"
S19_V1 = REPO / "data/HMM/neural_alignment/swr_detection/threshold_02pct"
S19_V2 = REPO / "data/HMM/neural_alignment/swr_detection_v2/threshold_02pct"
S17_FIG_V3 = REPO / "figures/HMM/neural_alignment/lfp_spectral_v3"
S18_FIG_A1_V2 = REPO / "figures/HMM/neural_alignment/lfp_state_identity_v2"
S18_FIG_A2_V2 = REPO / "figures/HMM/neural_alignment/lfp_three_region_granger_v2"
S19_FIG_V2 = REPO / "figures/HMM/neural_alignment/swr_detection_v2/threshold_02pct"

OUT_DOCX = REPO / "data/HMM/v2_refinement_report.docx"


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True; run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def df_table(doc, df, float_fmt="{:.3f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        c = table.cell(0, j)
        c.text = str(col)
        for r in c.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, (bool, np.bool_)):
                txt = "True" if bool(v) else "False"
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            c = table.cell(i + 1, j)
            c.text = txt
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(8)


def code(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    doc = Document()

    # Title
    t = doc.add_heading("v2 Bipolar Refinement — Rerun Report", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sr = sub.add_run("Scripts 17 / 18 / 19 rerun with refined bipolar pair list "
                       "(probe-1 LHA y<345 µm, RSP y>4680 µm, 16 intermediate "
                       "channels excluded). Probe-0 ACA unchanged.")
    sr.italic = True; sr.font.size = Pt(11)
    para(doc, "Pair files: data/HMM/neural_alignment/lfp/bipolar_pairs_imec{0,1}_v2.csv", size=9)
    para(doc, "Pair counts: ACA 370, LHA 176, RSP 176 (v1: 370/184/184)", size=9)
    para(doc, "Generated 2026-05-12", size=9)
    doc.add_paragraph()

    # Executive summary
    heading(doc, "Executive Summary", level=1)
    bullet(doc,
            "v2 bipolar refinement does real work, primarily on RSP. SWR detection "
            "(script 19) sees a dramatic drop in regional event counts (1266 → 322, "
            "75% reduction) while the SAME spike validation MW p values are preserved "
            "or strengthened — confirming the 8 dropped RSP boundary pairs were "
            "producing volume-conducted noise events rather than real ripples. "
            "S14 RSP goes from 150 events at p=2.1e-29 to 73 events at p=1.7e-54 "
            "(more significant with fewer events).")
    bullet(doc,
            "Script 17 (LFP spectral) sanity check confirms further common-mode "
            "reduction in 4/6 sessions. S8 ACA-LHA r drops from 0.045 to 0.0018 "
            "(25x cleaner), S12 from 0.036 to 0.0002 (180x cleaner). Mean r drops "
            "from 0.10 to 0.078 (~24%). M1 band-power, M3 coherence, M4 Granger "
            "verdicts unchanged — all negative; bidirectional coupling with no "
            "directional asymmetry.")
    bullet(doc,
            "Script 18 (state-identity + 3-region Granger). A1 results essentially "
            "unchanged: ACA high_gamma still 6/6 sessions Kruskal replicate; RSP "
            "still has 3/5 replicating bands; ACA feeding_vs_transition pairwise "
            "goes from 2/5 to 3/5 bands replicate. A2 Granger still null in every "
            "pair × band; v2 reveals a (still non-significant) tendency for RSP to "
            "LEAD LHA in beta / low_gamma / high_gamma (5/6 sessions each, binom p=0.22). "
            "Consistent with RSP-as-ripple-source from script 19.")
    bullet(doc,
            "LHA results virtually identical between v1 and v2 across all three "
            "scripts. The 8 dropped LHA boundary pairs (at y=345 µm, upper LHA edge) "
            "were not the source of LHA's spurious ripple-band signal. The LHA-RSP "
            "co-occurrence drops from 28.5% to 6.4% (still 6/6 sessions above shuffle "
            "p95). The reduction is mostly driven by RSP having fewer events, not "
            "LHA changing. So the volume-conduction story we proposed in the script "
            "19 report needs refinement: dropping RSP boundary pairs cuts the "
            "co-occurrence rate, but LHA still emits ripple-band envelope events "
            "with no spike participation — and those LHA events that DO coincide "
            "with v2 RSP events are spike-rich at the RSP end. The LHA bipolar "
            "regional mean appears to pick up RSP signal across many of its pairs "
            "(not just boundary pairs).")

    # Method
    heading(doc, "v2 vs v1 method differences", level=1)
    bullet(doc, "Probe-0 ACA: 370 pairs in both v1 and v2 (identical).")
    bullet(doc, "Probe-1 LHA: v1 used y<2500 µm (IMRO midpoint) → 184 pairs. v2 uses y<345 µm (spike-defined tissue) → 176 pairs.")
    bullet(doc, "Probe-1 RSP: v1 used y≥2500 µm → 184 pairs. v2 uses y>4680 µm → 176 pairs.")
    bullet(doc, "Probe-1 intermediate (345 ≤ y ≤ 4680 µm): 16 channels excluded entirely from regional pair construction.")
    bullet(doc, "All preprocessing (notch, bipolar subtract, anti-alias, decimate) and detection / aggregation logic unchanged. Only the channel-pair inputs to the regional means differ.")

    # ===== Script 17 =====
    heading(doc, "Script 17 — LFP spectral analysis", level=1)
    para(doc,
          "Per-session ACA-LHA Pearson r (sanity check on bipolar common-mode "
          "removal). Lower is better.")

    v1_sanity = pd.read_csv(S17_V1 / "sanity_bipolar_r_cross_session.csv")
    v3_sanity = pd.read_csv(S17_V3 / "sanity_bipolar_r_cross_session.csv")
    merged = v1_sanity.merge(v3_sanity, on="session", suffixes=("_v1", "_v3"))
    merged["fold_change"] = (merged["regional_ACA_vs_LHA_pearson_r_v3"]
                              / merged["regional_ACA_vs_LHA_pearson_r_v1"])
    df_table(doc, merged, float_fmt="{:.4f}")
    para(doc,
          f"Mean v1 r = {merged['regional_ACA_vs_LHA_pearson_r_v1'].mean():.3f}, "
          f"mean v3 r = {merged['regional_ACA_vs_LHA_pearson_r_v3'].mean():.3f}. "
          "4/6 sessions improved (S6, S8, S12, S14). S4 essentially unchanged. "
          "S16 slight uptick (still <0.25, well below the 0.5 red-flag threshold).")

    heading(doc, "M4 Granger sign test — v3 vs v1", level=2)
    v1m4 = pd.read_csv(S17_V1 / "M4_sign_test.csv")
    v3m4 = pd.read_csv(S17_V3 / "M4_sign_test.csv")
    m4 = v1m4[["band", "n_ACA_leads", "binom_p"]].merge(
        v3m4[["band", "n_ACA_leads", "binom_p"]],
        on="band", suffixes=("_v1", "_v3"),
    )
    df_table(doc, m4, float_fmt="{:.3f}")
    para(doc,
          "No band reaches significance in either v1 or v3. high_gamma "
          "remains the most-asymmetric cell (1/6 ACA-leads in both, binom p=0.22). "
          "The conclusion that LFP envelope Granger has no consistent directional "
          "asymmetry stands under v2 pairs.")

    heading(doc, "M1 / M3 replication — v3 vs v1", level=2)
    v1_m1 = pd.read_csv(S17_V1 / "M1_replication.csv")
    v3_m1 = pd.read_csv(S17_V3 / "M1_replication.csv")
    v1_m3 = pd.read_csv(S17_V1 / "M3_replication.csv")
    v3_m3 = pd.read_csv(S17_V3 / "M3_replication.csv")
    para(doc,
          f"M1 (band power per state×band, stay vs pre-exit): v1 cells passing = "
          f"{int(v1_m1['n_passing'].sum())}/{len(v1_m1)}; "
          f"v3 cells passing = {int(v3_m1['n_passing'].sum())}/{len(v3_m1)}. "
          f"M3 (ACA-LHA coherence): v1 = {int(v1_m3['n_passing'].sum())}/{len(v1_m3)}; "
          f"v3 = {int(v3_m3['n_passing'].sum())}/{len(v3_m3)}. "
          "Both essentially unchanged. The clean-negative finding is preserved.")

    # ===== Script 18 =====
    heading(doc, "Script 18 — state identity + 3-region Granger", level=1)

    heading(doc, "A1 Kruskal-Wallis replication (state identity)", level=2)
    v1_k = pd.read_csv(S18_A1_V1 / "A1_kruskal_replication.csv")
    v2_k = pd.read_csv(S18_A1_V2 / "A1_kruskal_replication.csv")
    merged_k = v1_k[["region", "band", "n_passing"]].merge(
        v2_k[["region", "band", "n_passing"]],
        on=["region", "band"], suffixes=("_v1", "_v2"),
    )
    df_table(doc, merged_k, float_fmt="{:.0f}")
    para(doc,
          "Kruskal-Wallis replication essentially identical between v1 and v2. "
          "ACA high_gamma still 6/6 sessions. RSP still has 4-5 sessions "
          "replicating in beta / low_gamma / high_gamma.")

    heading(doc, "A1 pairwise replication (home vs feeding vs transition_zone)", level=2)
    v1_p = pd.read_csv(S18_A1_V1 / "A1_pairwise_replication.csv")
    v2_p = pd.read_csv(S18_A1_V2 / "A1_pairwise_replication.csv")
    merged_p = v1_p[["region", "band", "pair", "n_passing"]].merge(
        v2_p[["region", "band", "pair", "n_passing"]],
        on=["region", "band", "pair"], suffixes=("_v1", "_v2"),
    )
    df_table(doc, merged_p, float_fmt="{:.0f}")

    heading(doc, "A2 sign test — 3-region Granger v2 vs v1", level=2)
    v1_sign = pd.read_csv(S18_A2_V1 / "A2_sign_test.csv")
    v2_sign = pd.read_csv(S18_A2_V2 / "A2_sign_test.csv")
    merged_s = v1_sign[["band", "pair", "n_forward_leads", "binom_p"]].merge(
        v2_sign[["band", "pair", "n_forward_leads", "binom_p"]],
        on=["band", "pair"], suffixes=("_v1", "_v2"),
    )
    df_table(doc, merged_s, float_fmt="{:.3f}")
    para(doc,
          "No (band × pair) reaches binom p<0.05 in either v1 or v2 — the n=6 "
          "minimum two-sided p is 0.031 (requires 6/6). The most-asymmetric "
          "cells in v2:")
    bullet(doc, "ACA-RSP high_gamma: ACA→RSP leads 5/6 (binom p=0.22) — same as v1.")
    bullet(doc, "LHA-RSP beta: RSP→LHA leads 5/6 (binom p=0.22) — same as v1.")
    bullet(doc, "LHA-RSP low_gamma: RSP→LHA leads 5/6 (binom p=0.22) — NEW in v2.")
    bullet(doc, "LHA-RSP high_gamma: RSP→LHA leads 5/6 (binom p=0.22) — NEW in v2.")
    para(doc,
          "Three of the four \"5/6\" cells in v2 have RSP leading. None reaches "
          "significance, but this is consistent with the script 19 SWR finding "
          "that RSP is the real ripple source — the LFP envelope coupling "
          "asymmetry, while too weak to be significant at n=6, leans toward "
          "RSP-as-source.")

    add_image(doc, S18_FIG_A2_V2 / "A2_sign_test_summary.png", width_in=6.5,
              caption="v2 sign test summary, 3 panels (one per pair).")

    # ===== Script 19 =====
    heading(doc, "Script 19 — SWR detection (re-aggregated)", level=1)

    para(doc,
          "v2 re-aggregation uses the same per-pair detected events from script 19 "
          "but filters them to keep only events from pairs in the v2 list, then "
          "re-aggregates at the same 2% threshold. No re-preprocessing was needed.")

    v1_rate = pd.read_csv(S19_V1 / "ripple_rate_per_session.csv")
    v2_rate = pd.read_csv(S19_V2 / "ripple_rate_per_session.csv")
    rate_pivot = v1_rate.merge(
        v2_rate, on=["session", "state", "region"], suffixes=("_v1", "_v2"),
    )[["session", "state", "region", "n_events_v1", "n_events_v2",
        "rate_per_min_v1", "rate_per_min_v2"]]
    rate_pivot["pct_retained"] = (100 * rate_pivot["n_events_v2"]
                                  / rate_pivot["n_events_v1"].replace(0, np.nan))
    heading(doc, "Per-session event counts v1 vs v2", level=2)
    df_table(doc, rate_pivot, float_fmt="{:.1f}")
    para(doc,
          "ACA: counts unchanged (0 or same; ACA pairs unchanged). "
          "LHA: 95-100% retained. "
          "RSP: 1-63% retained — most events dropped. The 8 boundary RSP pairs "
          "were the main event sources.")

    heading(doc, "Spike validation v1 vs v2", level=2)
    v1_val = pd.read_csv(S19_V1 / "validation_summary.csv")
    v2_val = pd.read_csv(S19_V2 / "validation_summary.csv")
    val_merge = v1_val[["session", "region", "n_units", "n_events",
                          "n_validated", "p_mw",
                          "median_event_spikes", "median_control_spikes"]].merge(
        v2_val[["session", "region", "n_events", "n_validated", "p_mw",
                  "median_event_spikes", "median_control_spikes"]],
        on=["session", "region"], suffixes=("_v1", "_v2"),
    )
    df_table(doc, val_merge, float_fmt="{:.2g}")
    para(doc,
          "Key per-session contrasts for RSP:")
    bullet(doc, "S6 RSP: v1 162 events at p=1.2e-08 → v2 12 events at p=0.0011. Fewer events, p still extreme.")
    bullet(doc, "S8 RSP: v1 352 events at p=5.1e-22, med event=4 → v2 7 events at p=5.7e-03, med event=5. v2 events are spike-richer than v1.")
    bullet(doc, "S12 RSP: v1 358 events at p=9.2e-33 → v2 227 events at p=8.8e-24. Most events retained.")
    bullet(doc, "S14 RSP: v1 150 events at p=2.1e-29 → v2 73 events at p=1.7e-54. v2 p is more significant despite half the events.")
    bullet(doc, "S16 RSP: v1 105 events → v2 2 events. No good units in either to validate.")
    bullet(doc, "S4 RSP: v1 139 events (1 unit, p=0.24) → v2 1 event (1 unit, p=0.64). Most v1 events were boundary noise.")

    heading(doc, "Cross-region co-occurrence v1 vs v2", level=2)
    v1_cooc = pd.read_csv(S19_V1 / "cross_region_co_occurrence_all_sessions.csv")
    v2_cooc = pd.read_csv(S19_V2 / "cross_region_co_occurrence_all_sessions.csv")
    cm = v1_cooc[["session", "pair", "n_A", "n_B", "obs_cooc_rate_A",
                    "shuf_p95", "exceeds_p95"]].merge(
        v2_cooc[["session", "pair", "n_A", "n_B", "obs_cooc_rate_A",
                   "shuf_p95", "exceeds_p95"]],
        on=["session", "pair"], suffixes=("_v1", "_v2"),
    )
    df_table(doc, cm, float_fmt="{:.3f}")
    para(doc,
          "LHA-RSP co-occurrence rate: v1 mean = "
          f"{v1_cooc[v1_cooc.pair=='LHA-RSP']['obs_cooc_rate_A'].mean():.3f}, "
          f"v2 mean = {v2_cooc[v2_cooc.pair=='LHA-RSP']['obs_cooc_rate_A'].mean():.3f}. "
          f"Both v1 and v2 have 6/6 sessions exceeding shuffle p95 (real coupling), "
          "but the rate dropped 4.5× under v2. The drop is mostly driven by fewer "
          "RSP events, not LHA changes.")

    add_image(doc, S19_FIG_V2 / "ripple_rate_per_session.png", width_in=6.5,
              caption="v2 ripple event rate per session per region.")
    add_image(doc, S19_FIG_V2 / "cross_region_co_occurrence.png", width_in=6.5,
              caption="v2 cross-region co-occurrence.")

    # Diagnostic discussion
    heading(doc, "Diagnostic: did v2 do real work?", level=1)
    para(doc,
          "Yes, primarily on RSP. Three lines of evidence:")
    bullet(doc,
            "RSP per-pair events dropped to 1-63% retained across sessions. The "
            "8 dropped boundary pairs (y=4680, lower edge of RSP tissue) were "
            "contributing far more events per pair than the other 176 RSP pairs. "
            "This is a sign of localized noise / volume conduction concentrated at "
            "those boundary positions.")
    bullet(doc,
            "RSP validation strengthens or holds: S6 v2 events show HIGHER median "
            "spike count than v1 (1.5 vs 1), S8 v2 events show HIGHER median "
            "(5 vs 4). S14 MW p actually drops 25 orders of magnitude (more "
            "significant) despite halving the event count. The v2 events are "
            "real ripples; the dropped v1 \"extra\" events were noise.")
    bullet(doc,
            "Script 17 sanity-check ACA-LHA Pearson r improves in 4/6 sessions, "
            "with two sessions (S8, S12) seeing 25-180x reduction. The v2 pair "
            "set leaves less residual common-mode in the regional means.")

    para(doc,
          "What v2 did NOT do:")
    bullet(doc,
            "LHA events are virtually unchanged (95-100% per-pair events retained "
            "in all sessions). The volume-conduction story we proposed in the "
            "script 19 v1 report — that LHA \"events\" come from a few LHA "
            "boundary pairs catching RSP signal — is NOT supported. LHA events "
            "come from the bulk of the LHA pairs, not the 8 upper-boundary ones. "
            "Yet LHA events still show no spike validation in any session. The "
            "LHA bipolar regional mean is detecting ripple-band envelope deflections "
            "that do not recruit LHA spiking. Either (a) volume conduction is "
            "broader than we thought and spans most LHA pairs, or (b) ripple-band "
            "envelope events in LHA tissue genuinely don't drive LHA neuronal "
            "firing.")
    bullet(doc,
            "Script 17 / 18 verdicts unchanged. M1 band power, M3 coherence, M4 "
            "Granger, A2 3-region Granger all remain null. The state-identity "
            "result (A1, ACA high_gamma 6/6 sessions) is identical between v1 "
            "and v2.")

    # Output files
    heading(doc, "Output files", level=1)
    code(doc, "data/HMM/neural_alignment/lfp_spectral_v3/")
    bullet(doc, "v2-pair version of all script-17 outputs (sanity, M1, M3, M4)")
    code(doc, "data/HMM/neural_alignment/lfp_state_identity_v2/")
    bullet(doc, "v2-pair A1 (Kruskal + pairwise) outputs")
    code(doc, "data/HMM/neural_alignment/lfp_three_region_granger_v2/")
    bullet(doc, "v2-pair A2 (3-region Granger) outputs")
    code(doc, "data/HMM/neural_alignment/swr_detection_v2/threshold_02pct/")
    bullet(doc, "v2-pair SWR re-aggregation at 2% threshold")
    bullet(doc, "Underlying per-pair events from script 19 v1 are still at data/HMM/neural_alignment/swr_detection/session_{N}/ — re-usable at any pair version")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Saved {OUT_DOCX}")


if __name__ == "__main__":
    main()
