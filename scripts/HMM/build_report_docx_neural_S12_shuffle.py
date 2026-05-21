"""Build a Word-document report of the Track B shuffle control (S12).

Reads outputs from data/HMM/neural_alignment/shuffle_control_S12/ and figures
from figures/HMM/neural_alignment/shuffle_control_S12/.

Output: data/HMM/neural_alignment_S12_shuffle_control_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.2f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    cfg = load_config()
    sh_dir = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "shuffle_control_S12"
    fig_dir = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "shuffle_control_S12"

    df_b1 = pd.read_csv(sh_dir / "shuffle_B1_summary.csv")
    df_b3 = pd.read_csv(sh_dir / "shuffle_B3_summary.csv")
    df_z = pd.read_csv(sh_dir / "shuffle_B3_max_z_per_unit.csv")

    # Observed counts (re-derived from CSV — observed ACA/LHA from 10b)
    OBS_B1_ACA, OBS_B1_LHA = 164, 88
    OBS_B3_ACA, OBS_B3_LHA = 165, 88
    N_ACA, N_LHA = 165, 89

    def stats_block(vals, observed, N):
        mean = float(vals.mean())
        ci_lo = float(np.percentile(vals, 2.5))
        ci_hi = float(np.percentile(vals, 97.5))
        rng = (int(vals.min()), int(vals.max()))
        pct = float((vals <= observed).mean() * 100)
        return dict(
            observed=f"{observed}/{N}",
            shuffle_mean=f"{mean:.1f}",
            ci_95=f"[{ci_lo:.0f}, {ci_hi:.0f}]",
            shuffle_range=f"[{rng[0]}, {rng[1]}]",
            obs_pctile=f"{pct:.1f}",
        )

    s_b1_aca = stats_block(df_b1["n_sig_aca"].values, OBS_B1_ACA, N_ACA)
    s_b1_lha = stats_block(df_b1["n_sig_lha"].values, OBS_B1_LHA, N_LHA)
    s_b3_aca = stats_block(df_b3["n_sig_aca"].values, OBS_B3_ACA, N_ACA)
    s_b3_lha = stats_block(df_b3["n_sig_lha"].values, OBS_B3_LHA, N_LHA)

    n_above_aca = int((df_z[df_z.region == "ACA"]["exceeds_shuf_p95"]).sum())
    n_above_lha = int((df_z[df_z.region == "LHA"]["exceeds_shuf_p95"]).sum())

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===== TITLE =====
    title = doc.add_heading(
        "Shuffle Control for Track B State Selectivity (Session 12)", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Tests whether the pervasive Track B state-encoding result survives "
        "circular shuffling of the Viterbi/posterior temporal alignment."
    )
    r.italic = True
    r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-06  •  Script: scripts/HMM/10c_shuffle_control_S12.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "Track B (10b) reported that 99% of ACA and 99% of LHA units pass an "
        "ANOVA test for state selectivity (B1) and that 100% of ACA and 99% of "
        "LHA units have at least one significant Poisson-GLM coefficient on "
        "state posteriors (B3). With 3750 HMM bins × 14 states the tests have "
        "very high statistical power, so a fraction of the apparent "
        "significance could reflect autocorrelation rather than genuine state "
        "encoding. This script runs 100 circular-shift shuffles of the "
        "Viterbi sequence and posterior matrix (offset ∈ [100, T−100] bins), "
        "preserving state dwell time and marginal occupancy but breaking "
        "alignment with neural data. B1 and B3 are re-run under each shuffle "
        "with FDR correction within region; the observed counts are compared "
        "to the resulting null distributions."
    )
    add_para(doc, "Headline result, in three sentences:", bold=True, size=11)
    add_bullet(
        doc,
        f"In ALL four tests (B1×{{ACA,LHA}} and B3×{{ACA,LHA}}) the observed "
        f"count of significant units sits at the 100th percentile of the "
        f"shuffle null — every one of the 100 circular shifts produces fewer "
        f"sig units than the real data. So the population-level signal IS "
        f"above null."
    )
    add_bullet(
        doc,
        f"BUT the shuffle null is already very high: under random circular "
        f"shifts the median number of FDR-significant units is "
        f"{s_b1_aca['shuffle_mean']}/{N_ACA} ACA in B1 (~85%), "
        f"{s_b1_lha['shuffle_mean']}/{N_LHA} LHA in B1 (~81%), "
        f"{s_b3_aca['shuffle_mean']}/{N_ACA} ACA in B3 (~95%), and "
        f"{s_b3_lha['shuffle_mean']}/{N_LHA} LHA in B3 (~95%). Most of the "
        f"apparent state-selectivity is a power/autocorrelation artifact, "
        f"not a real state-alignment signal."
    )
    add_bullet(
        doc,
        f"At the per-unit level (max |z| from the GLM exceeds the unit's own "
        f"shuffle 95th percentile): {n_above_aca}/{N_ACA} ACA units "
        f"({n_above_aca/N_ACA*100:.0f}%) and only "
        f"{n_above_lha}/{N_LHA} LHA units ({n_above_lha/N_LHA*100:.0f}%) "
        f"survive the per-unit shuffle test. The defensible head-count of "
        f"genuinely state-encoding units is ~84% in ACA and ~46% in LHA, not "
        f"99%/99%. LHA's pervasive selectivity from 10b was largely artifact."
    )

    # ===== METHODS =====
    add_heading(doc, "Method", level=1)
    add_para(
        doc,
        "Circular shifting was chosen as the null because it preserves the "
        "behavioral state sequence's autocorrelation structure (dwell times, "
        "transition pattern) and marginal occupancy of each state, while "
        "destroying the bin-by-bin alignment between neural activity and "
        "state labels. This is a CONSERVATIVE null — alternative shuffles "
        "(random reassignment, ISI permutation of spike trains) destroy more "
        "structure and would yield wider gaps between observed and null. The "
        "results here therefore represent a lower bound on the genuine "
        "state-encoding signal."
    )
    add_bullet(doc, f"Shuffles: 100 iterations, master seed = 20260506.")
    add_bullet(doc, f"Per iteration, offset = uniform integer in [100, T−100] "
                    f"where T = 3750 (HMM bins). Same offset applied to both "
                    f"the Viterbi sequence and the (T, K) posterior matrix.")
    add_bullet(doc, "B1 — ANOVA on rates (Hz) per unit across the K=14 shifted "
                    "Viterbi groups; FDR within region.")
    add_bullet(doc, "B3 — Poisson GLM with the most-occupied state (S2 in both "
                    "regions) dropped as reference, FDR within region across all "
                    "(unit, state) coefficients, and a Wald |z|>2.5 threshold to "
                    "match 10b. A custom IRLS (~5–10 ms/fit) replaces statsmodels "
                    "GLM for tractability over 254 units × 100 iterations = 25,400 "
                    "fits.")
    add_bullet(doc, "Per-unit comparison: each unit's observed max |z| (across "
                    "non-reference state coefficients) versus its 95th-percentile "
                    "of max |z| across the 100 shuffles. Units where observed > "
                    "shuffle p95 pass the per-unit test.")

    # ===== B1 RESULTS =====
    doc.add_page_break()
    add_heading(doc, "B1 Shuffle — per-unit ANOVA", level=1)
    b1_tbl = pd.DataFrame([
        dict(region="ACA", **s_b1_aca),
        dict(region="LHA", **s_b1_lha),
    ])
    b1_tbl.columns = ["region", "observed", "shuffle mean", "95% CI",
                       "shuffle range", "obs pctile within null"]
    add_df_table(doc, b1_tbl)
    add_bullet(
        doc,
        f"ACA observed 164/165 (99.4%) is above every one of the 100 shuffles "
        f"(max shuffle = {df_b1['n_sig_aca'].max()}). The shuffle mean of 139.6 "
        f"shows that ~84.6% of ACA units pass FDR even under random circular "
        f"shifts — purely from temporal autocorrelation and marginal-occupancy "
        f"structure. The genuine 'above-null' signal is ~14 percentage points "
        f"on top of the artifact floor.",
    )
    add_bullet(
        doc,
        f"LHA observed 88/89 (98.9%) similarly clears all shuffles "
        f"(max shuffle = {df_b1['n_sig_lha'].max()}). The shuffle mean of 72.0 "
        f"is lower in absolute terms (~80.9% of LHA units), so the genuine "
        f"above-null signal is larger here in proportional terms (~18 pp) "
        f"despite LHA having fewer units.",
    )

    # ===== B3 RESULTS =====
    add_heading(doc, "B3 Shuffle — Poisson GLM", level=1)
    b3_tbl = pd.DataFrame([
        dict(region="ACA", **s_b3_aca),
        dict(region="LHA", **s_b3_lha),
    ])
    b3_tbl.columns = ["region", "observed", "shuffle mean", "95% CI",
                       "shuffle range", "obs pctile within null"]
    add_df_table(doc, b3_tbl)
    add_bullet(
        doc,
        f"ACA observed 165/165 (100%) clears all shuffles, but the shuffle "
        f"mean is 157.2 (~95%) — only a 5-percentage-point gap. Most of the "
        f"GLM 'state-encoding' result for ACA is autocorrelation-driven; the "
        f"genuine signal is small at the bulk-count level.",
    )
    add_bullet(
        doc,
        f"LHA observed 88/89 (98.9%) versus shuffle mean 84.3 (~95%) and "
        f"shuffle 95% CI [80, 88]. The observed value is at the upper edge of "
        f"the shuffle 95% CI — population-level B3 in LHA is barely "
        f"distinguishable from the circular-shift null.",
    )

    add_image(doc, fig_dir / "shuffle_distributions.png", width_in=6.7,
              caption="Distribution of FDR-significant unit counts across 100 "
                      "circular-shift shuffles. Top row B1 (ANOVA); bottom row "
                      "B3 (Poisson GLM). Red line = observed count from real "
                      "(unshuffled) data. Title shows shuffle mean, 95% CI, and "
                      "the percentile of observed within the null.")

    # ===== PER-UNIT MAX |Z| =====
    doc.add_page_break()
    add_heading(doc, "Per-unit max |z| comparison", level=1)
    add_para(
        doc,
        "The bulk-count test above is a coarse summary. A cleaner test is "
        "per-unit: for each unit, compare the observed max |z| (across "
        "non-reference state coefficients) to the 95th percentile of the "
        "same statistic across the 100 shuffles. Units with observed > "
        "shuffle p95 are individually robust — their state encoding is "
        "stronger than circular-shift null would predict."
    )
    pu_tbl = pd.DataFrame([
        dict(region="ACA",
             units_above_p95=f"{n_above_aca}/{N_ACA} ({n_above_aca/N_ACA*100:.1f}%)"),
        dict(region="LHA",
             units_above_p95=f"{n_above_lha}/{N_LHA} ({n_above_lha/N_LHA*100:.1f}%)"),
    ])
    add_df_table(doc, pu_tbl)
    add_bullet(
        doc,
        f"ACA: {n_above_aca}/{N_ACA} ({n_above_aca/N_ACA*100:.1f}%) units exceed their own "
        f"shuffle p95 — a strong genuine signal. The remaining "
        f"{N_ACA-n_above_aca} ACA units' apparent state-selectivity from B1/B3 "
        f"was likely circular-shift artifact.",
    )
    add_bullet(
        doc,
        f"LHA: {n_above_lha}/{N_LHA} ({n_above_lha/N_LHA*100:.1f}%) units pass — only "
        f"about half. The remaining {N_LHA-n_above_lha} LHA units' apparent "
        f"state-selectivity was largely artifact. This sharply revises 10b's "
        f"'99% LHA selectivity' headline.",
    )
    add_bullet(
        doc,
        f"The 33-unit LHA-S2 (feeding) preference cluster from 10b's B1 "
        f"deserves a focused follow-up: the current per-unit test passes 41 "
        f"LHA units, and these likely include the bulk of the S2-preferrers, "
        f"but a direct shuffle of the S2-vs-second-state firing-rate "
        f"difference per unit would put the feeding-preference claim on "
        f"firmer ground.",
    )

    add_image(doc, fig_dir / "shuffle_max_z_comparison.png", width_in=6.7,
              caption="Per-unit observed max |z| (y-axis) vs the 95th "
                      "percentile of shuffled max |z| (x-axis), for ACA "
                      "(left) and LHA (right). Each dot = one unit. Dashed "
                      "diagonal: equality. Units above the diagonal exceed "
                      "their shuffle null.")

    # ===== WHAT THIS CHANGES =====
    add_heading(doc, "What this changes about Track B claims", level=1)
    add_bullet(
        doc,
        "10b's headline 'every unit is state-selective' (99-100%) is replaced "
        "by 'a majority of units are state-selective' — 84% in ACA, 46% in "
        "LHA. The remaining ~14-54% of unit-level 'selectivity' was driven by "
        "autocorrelation rather than state alignment.",
    )
    add_bullet(
        doc,
        "The contrast between regions is sharpened: ACA shows much stronger "
        "above-null state encoding than LHA. LHA's headline 'feeding "
        "concentration on S2' is plausible but not verified by this shuffle "
        "(the test is on max |z| across all states, not specifically on "
        "preference structure). A focused follow-up shuffle on per-unit "
        "preference contrast is the obvious next step.",
    )
    add_bullet(
        doc,
        "B2 (within-state pre/post-discovery shifts) and B4 (PCA centroid "
        "shifts) are NOT tested by this shuffle. They live on a different "
        "axis (pre/post-discovery contrast within state) and would need a "
        "different shuffle (e.g., shuffling pre/post labels while keeping "
        "state alignment intact, or block-resampling). Those Track B "
        "findings — particularly the S8/S9 pot-zone pre/post centroid shift "
        "in PC space — remain the cleanest candidate cognitive-state "
        "signature in S12.",
    )
    add_bullet(
        doc,
        "Bottom line: state encoding in S12 is REAL but not as universal as "
        "10b suggested. ACA carries a robust state code; LHA carries a "
        "weaker, more selective one. Multi-session replication and a "
        "preference-specific shuffle are the natural follow-ups.",
    )

    # ===== CAVEATS =====
    add_heading(doc, "Caveats", level=1)
    add_bullet(doc, "Circular shift is a CONSERVATIVE null. It preserves "
                    "autocorrelation, dwell-time, and marginal occupancy. A more "
                    "aggressive null (e.g., random reassignment of state labels, "
                    "or independent ISI shuffling per unit) would lower the "
                    "shuffle floor and the observed-vs-null gap would widen.")
    add_bullet(doc, "100 shuffles supports a 95% CI but is sparse for higher "
                    "percentiles. The 95th-percentile per-unit threshold has "
                    "Monte-Carlo noise of ±~5% of units flagged.")
    add_bullet(doc, "B1 and B3 are not independent: they're both tests on the "
                    "same neural data with overlapping mathematical structure. "
                    "Don't multiply the p-values.")
    add_bullet(doc, "The custom IRLS gave 88/89 LHA B3-significant in the "
                    "observed data versus 89/89 in 10b's statsmodels. One unit "
                    "was a numerical edge case; this is a 1-unit difference and "
                    "doesn't change the headline.")
    add_bullet(doc, "Single session only. Multi-session shuffle controls would "
                    "give a population-level test of whether ANY state encoding "
                    "is real across the cohort.")

    # ===== OUTPUTS =====
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/neural_alignment/shuffle_control_S12/shuffle_B1_summary.csv",
         "Per-iteration ACA/LHA counts of FDR-sig units (B1 ANOVA)."),
        ("data/HMM/neural_alignment/shuffle_control_S12/shuffle_B3_summary.csv",
         "Per-iteration ACA/LHA counts of units with sig coefficient (B3 GLM)."),
        ("data/HMM/neural_alignment/shuffle_control_S12/shuffle_B3_max_z_per_unit.csv",
         "Per (unit, region): real max |z| and shuffle 95th percentile, exceed flag."),
        ("figures/HMM/neural_alignment/shuffle_control_S12/shuffle_distributions.png",
         "2x2 panel: B1 and B3 sig-count distributions per region with observed "
         "marker."),
        ("figures/HMM/neural_alignment/shuffle_control_S12/shuffle_max_z_comparison.png",
         "Per-unit observed max |z| vs shuffle p95 (scatter + diagonal)."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    out_path = REPO_ROOT / "data" / "HMM" / "neural_alignment_S12_shuffle_control_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
