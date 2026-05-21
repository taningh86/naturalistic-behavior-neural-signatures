"""Build a Word-document report of the six population-level metrics on stay
vs pre-exit (script 15).

Reads outputs from data/HMM/neural_alignment/state_transitions/population_metrics/
and figures/.../population_metrics/.

Output: data/HMM/neural_alignment_population_metrics_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.2f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def state_label(prof_row):
    zone_keys = ["home", "transition", "pot", "pot_zone", "arena", "other"]
    zone_p = {z: prof_row[f"zone_{z}_prob"] for z in zone_keys}
    top_zone = max(zone_p, key=zone_p.get)
    desc = f"{top_zone}({zone_p[top_zone]:.2f})"
    events = []
    for ev in ["digging_sand", "feeding", "rearing",
               "contemplation_at_transition", "exploration_at_transition"]:
        if prof_row[f"event_{ev}_prob"] > 0.30:
            short = ev.split("_")[0]
            events.append(f"{short}={prof_row[f'event_{ev}_prob']:.2f}")
    if events:
        desc += " + " + ",".join(events)
    return desc


def main():
    cfg = load_config()
    base_out = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "state_transitions" / "population_metrics"
    base_fig = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "state_transitions" / "population_metrics"

    master = pd.read_csv(base_out / "master_replication_table.csv")
    cross = pd.read_csv(base_out / "cross_session_pass.csv")

    prof = pd.read_csv(REPO_ROOT / cfg["merge_dirs"]["state_profiles_csv"])
    K = len(prof)
    state_descriptions = {k: state_label(prof.iloc[k]) for k in range(K)}

    # Reshape: (region, state) rows × metrics columns, cell = "n_pass/n_tested"
    metric_order = ["M1_n_sig_units", "M2_cv_isi_diff", "M3_pc_speed_diff",
                     "M4_pr_diff", "M5_corr_norm", "M6_cross_corr_diff"]
    metric_short = {"M1_n_sig_units": "M1 Fano",
                     "M2_cv_isi_diff": "M2 ISI CV",
                     "M3_pc_speed_diff": "M3 PC speed",
                     "M4_pr_diff": "M4 PR",
                     "M5_corr_norm": "M5 corr norm",
                     "M6_cross_corr_diff": "M6 cross-corr"}

    def reshape_region(reg):
        rows = []
        sub = master[master.region == reg]
        states = sorted(sub["state"].unique().astype(int))
        for st in states:
            row = {"state": f"S{st}", "behavior": state_descriptions.get(st, "")[:28]}
            for m in metric_order:
                sel = sub[(sub.state == st) & (sub.metric == m)]
                if len(sel):
                    n_p = int(sel.iloc[0]["n_sessions_passing"])
                    n_t = int(sel.iloc[0]["n_sessions_tested"])
                    row[metric_short[m]] = f"{n_p}/{n_t}"
                else:
                    row[metric_short[m]] = ""
            rows.append(row)
        return pd.DataFrame(rows)

    aca_tbl = reshape_region("ACA")
    lha_tbl = reshape_region("LHA")
    aca_lha_tbl = reshape_region("ACA-LHA")

    # ===== Build doc =====
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_heading(
        "Six Population-Level Metrics for Pre-Exit Neural Signatures", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Extends script 14's per-unit firing-rate test with five additional "
        "population metrics applied to the same stay-vs-pre-exit contrast."
    )
    r.italic = True
    r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-08  •  Script: scripts/HMM/15_pre_exit_population_metrics.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "Script 14 found a strong ACA generic pre-exit signal at the per-unit "
        "MEAN firing-rate level (4-6/6 sessions for most states tested). This "
        "script asks: do other population properties — variance/dispersion, "
        "ISI variability, trajectory geometry, dimensionality, correlation "
        "structure, cross-region coupling — also shift pre-exit, and at "
        "which states do they replicate across sessions?"
    )
    add_para(doc, "Three plain-language findings:", bold=True, size=11)
    add_bullet(
        doc,
        "**S3 (home) is the standout pre-exit state across multiple "
        "metrics.** ACA passes shuffle p95 in 5/5 sessions for M2 (ISI CV), "
        "M3 (PC trajectory speed), and M5 (correlation structure norm), and "
        "in 4/5 sessions for M4 (participation ratio). LHA passes M5 in "
        "5/5 sessions and M2 in 3/5. The pre-exit window from home base "
        "shows coordinated changes in ISI variability, population speed, "
        "dimensionality, and correlation structure — the cleanest "
        "multi-metric pre-exit signature in the dataset.",
    )
    add_bullet(
        doc,
        "**M5 (pairwise correlation structure norm) is the most replicating "
        "metric.** ACA-S3 5/5, LHA-S3 5/5, ACA-S2 3/3 (limited testing). "
        "Pairwise correlation patterns shift pre-exit more reliably than "
        "any single-unit property. The signal is collective, not in "
        "individual cells' dispersion or trajectory speed.",
    )
    add_bullet(
        doc,
        "**M1 (Fano factor / dispersion residuals) is the weakest** — "
        "0/6 sessions for nearly every state. **The strong ACA pre-exit "
        "signal from script 14 is purely a MEAN firing-rate effect; "
        "per-unit dispersion does NOT change pre-exit.** This is an "
        "important specification of the script 14 result: cells fire at "
        "different rates pre-exit, but their spike-count variance:mean "
        "ratio stays constant.",
    )

    # ===== METHODS =====
    add_heading(doc, "Methods", level=1)
    add_bullet(doc, "Same QC-filtered units, 480 ms HMM-aligned bins, and bin-"
                    "labelling scheme as script 14 (K_pre=3, runs ≥6 bins).")
    add_bullet(doc, "Six metrics, all applied to the same A1 contrast "
                    "(stay bins vs pre-exit bins, pooled across destinations).")
    add_para(doc, "Metrics:", bold=True, size=10)
    add_bullet(doc, "**M1 — Fano factor (per-unit, per-state):** Mann-Whitney "
                    "on per-bin |count − cond_mean|/√cond_mean (Pearson "
                    "residual magnitude) between stay and pre-exit. Per-state "
                    "FDR-sig unit count. Shuffle compares the count.")
    add_bullet(doc, "**M2 — ISI CV (per-state scalar):** CV(ISI_pre_exit) − "
                    "CV(ISI_stay), computed by extracting within-bin spikes "
                    "from time windows defined by stay/pre-exit bin labels, "
                    "concatenating, and computing CV(ISI). Shuffle compares "
                    "the |scalar| (two-tailed).")
    add_bullet(doc, "**M3 — PC trajectory speed (per-state scalar):** mean "
                    "Euclidean step length in PC1-3 between consecutive bins "
                    "of the same condition within the same run; pre-exit "
                    "minus stay. Shuffle |scalar|.")
    add_bullet(doc, "**M4 — Participation ratio (per-state scalar):** PR(C) = "
                    "(Σλ)² / Σ(λ²) of the n_units × n_units covariance; pre-"
                    "exit minus stay. Shuffle |scalar|.")
    add_bullet(doc, "**M5 — Pairwise correlation structure norm:** "
                    "‖C_pre − C_stay‖_F / (n_units × mean(|C_stay|+|C_pre|)/2), "
                    "where C is the unit×unit Pearson correlation matrix. "
                    "Shuffle |scalar|.")
    add_bullet(doc, "**M6 — Cross-region ACA-LHA correlation:** mean |r| of "
                    "all ACA × LHA unit pairs, pre-exit minus stay. Shuffle "
                    "|scalar|.")
    add_bullet(doc, "100 circular-shift Viterbi shuffles per session (offset ∈ "
                    "[200, T−200], same protocol as script 14). Pass = "
                    "observed exceeds shuffle p95 (one-tailed for M1 unit "
                    "count, two-tailed for M2-M6 scalars).")
    add_bullet(doc, "States analyzed: ACA {S2, S3, S4, S6, S8, S9, S12} per "
                    "script 14's ≥3-session A1 replication. LHA {S2, S3}. M6 "
                    "cross-region: union of those state lists. Per-session "
                    "skip if <30 stay or <30 pre-exit bins.")
    add_bullet(doc, "**Spec deviation**: M2 done at per-state-scalar level "
                    "(CV pooled across units & bins) rather than per-unit "
                    "FDR-sig unit count. Per-unit ISI shuffles (250 units × "
                    "100 shuffles) are tractable but slow; the per-state "
                    "scalar with shuffle null is the same kind of test as "
                    "M3-M6 and matches their cross-session replication test.")

    # ===== MASTER TABLE / HEATMAP =====
    doc.add_page_break()
    add_heading(doc, "Master replication table", level=1)
    add_para(
        doc,
        "Each cell shows sessions passing shuffle p95 / sessions tested. "
        "States are tested only when both stay and pre-exit have ≥30 bins; "
        "many states are skipped in S14 (very early discovery, few pre-"
        "exit bins) and S6 (rare digging runs)."
    )
    add_para(doc, "ACA — per-state replication:", bold=True, size=10)
    add_df_table(doc, aca_tbl)
    add_para(doc, "LHA — per-state replication:", bold=True, size=10)
    add_df_table(doc, lha_tbl)
    add_para(doc, "ACA-LHA cross-region (M6 only):", bold=True, size=10)
    add_df_table(doc, aca_lha_tbl)

    add_image(doc, base_fig / "master_replication_heatmap.png", width_in=6.7,
              caption="Master replication heatmap. Rows = (region, state). "
                      "Columns = M1-M6. Cell = # sessions passing shuffle p95 "
                      "/ # tested. Darker red = more replicating.")

    # ===== INTERPRETATION =====
    doc.add_page_break()
    add_heading(doc, "Interpretation by metric", level=1)

    add_heading(doc, "M5 — pairwise correlation structure (most replicating)", level=2)
    add_bullet(doc, "ACA-S3 home: 5/5 sessions; LHA-S3: 5/5; ACA-S2 pot+feed: 3/3.")
    add_bullet(doc, "Pre-exit changes the unit × unit correlation pattern more "
                    "reliably than any single-unit property. The signal lives in "
                    "the geometry of co-firing, not in individual cells' rates "
                    "or variances.")
    add_bullet(doc, "Other ACA states show partial M5 replication (S4 2/4, S9 "
                    "2/5, S6 1/6, S8 0/4, S12 1/3) — consistent with S3 and S2 "
                    "(home and feeding) being the cleanly-bounded states where "
                    "coordinated population transitions are most identifiable.")

    add_heading(doc, "M2 — ISI CV (second-most replicating)", level=2)
    add_bullet(doc, "ACA-S3: 5/5; LHA-S3: 3/5; other ACA states 1-2/4-6.")
    add_bullet(doc, "Inter-spike-interval distribution becomes more (or less) "
                    "regular pre-exit at S3 in both regions. The pooled-across-"
                    "units CV captures this collective change without needing "
                    "per-unit testing.")

    add_heading(doc, "M3 — PC trajectory speed (S3-specific)", level=2)
    add_bullet(doc, "Replicates strongly only at ACA-S3 (5/5). Other states "
                    "show 1-2/3-6 sessions passing.")
    add_bullet(doc, "The population trajectory in PC space accelerates (or "
                    "decelerates) reliably pre-exit only when leaving the home "
                    "state.")

    add_heading(doc, "M4 — Participation ratio (S3-only)", level=2)
    add_bullet(doc, "ACA-S3: 4/5; LHA-S2: 2/3; LHA-S3: 2/5. Other states "
                    "0-2/3-6.")
    add_bullet(doc, "Effective dimensionality of the activity changes pre-exit "
                    "primarily at home (ACA) and feeding (LHA).")

    add_heading(doc, "M1 — Fano factor (weakest)", level=2)
    add_bullet(doc, "0/6 sessions in nearly every state. Best result: ACA-S3 "
                    "2/6.")
    add_bullet(doc, "**Specifies the script 14 ACA mean-rate signal**: cells "
                    "change MEAN firing rate pre-exit but their spike-count "
                    "VARIANCE:MEAN ratio (Fano factor) stays constant. The "
                    "signal is in the rate, not the dispersion.")
    add_bullet(doc, "This is theoretically interesting: a uniform rate shift "
                    "without a Fano change is consistent with Poisson-like "
                    "firing where the mean rate moves but variance scales with "
                    "it. A non-Poisson shift (e.g., bursting modulation) would "
                    "change Fano. We see no evidence for the latter.")

    add_heading(doc, "M6 — Cross-region ACA-LHA correlation (null)", level=2)
    add_bullet(doc, "0-1/6 sessions for every state. ACA-LHA inter-regional "
                    "coupling does not reliably change pre-exit.")
    add_bullet(doc, "Combined with M5's positive within-region result: "
                    "**within-region correlation structure shifts pre-exit, "
                    "but cross-region coupling does not.** The pre-exit "
                    "computation appears local to each region.")

    add_heading(doc, "Pattern across metrics — S3 is the anchor", level=1)
    add_bullet(doc, "S3 (home) is the only state that replicates strongly on "
                    "M2, M3, M4, M5 simultaneously in ACA. LHA replicates on "
                    "M5 (5/5) and M2 (3/5) at S3. No other state shows this "
                    "multi-metric convergence.")
    add_bullet(doc, "Why home? S3 has long, well-bounded runs (the animal "
                    "spends extended periods at the home zone, see Track B "
                    "B1), giving the largest and cleanest stay-vs-pre-exit "
                    "samples. Pre-exit from home is also a high-stakes "
                    "decision (the animal is leaving safety to begin a "
                    "foraging bout) — biologically plausible that multiple "
                    "population properties shift coherently.")
    add_bullet(doc, "The ACA generic pre-exit signal (script 14) is therefore "
                    "decomposable: at most states it's ONLY a mean-rate shift; "
                    "at S3 it expands to include collective population "
                    "geometry changes (M3-M5) and ISI-pattern changes (M2).")
    add_bullet(doc, "**For downstream analyses**, S3 (home) is the cleanest "
                    "pre-exit anchor for computational modeling — multi-metric "
                    "convergence is a strong sign of a coordinated neural "
                    "transition.")

    # ===== CAVEATS =====
    add_heading(doc, "Caveats", level=1)
    add_bullet(doc, "S14 sample-limited: very early discovery (54 s) means "
                    "many states have <30 pre-exit bins. Consequently S14 only "
                    "contributes 23 pass-test rows vs 42-43 for other sessions.")
    add_bullet(doc, "M2 ISI CV deviates from spec (per-state scalar instead "
                    "of per-unit FDR-sig count). Per-unit ISI testing requires "
                    "100 shuffle CV resamples per unit per state — feasible "
                    "but ~3× slower. The per-state scalar test gives the same "
                    "kind of cross-session replication answer.")
    add_bullet(doc, "Metrics are not orthogonal. M3 (PC trajectory speed) and "
                    "M5 (corr structure) both depend on the population "
                    "covariance; M4 (PR) is also a function of the same "
                    "covariance eigenstructure. The fact that all three "
                    "replicate at S3 is partly a consistency check, partly "
                    "redundancy.")
    add_bullet(doc, "K_pre = 3 bins is unchanged from script 14. Sensitivity "
                    "to K_pre = 2 or 5 not tested in this script.")
    add_bullet(doc, "Cross-session unit matching not done. Per-unit replication "
                    "(M1) is at the population-of-units level, not the same-"
                    "cell level.")
    add_bullet(doc, "Single-mouse, single-paradigm dataset. The S3-home "
                    "convergence finding needs replication in additional mice "
                    "before any general claim about pre-exit population "
                    "geometry at home base.")

    # ===== OUTPUTS =====
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/neural_alignment/state_transitions/population_metrics/master_replication_table.csv",
         "Per (region, state, metric): n_sessions_tested, n_sessions_passing, sessions_passing."),
        ("data/HMM/neural_alignment/state_transitions/population_metrics/cross_session_pass.csv",
         "Per (session, region, state, metric): observed value, shuffle p95, exceeds_p95 flag."),
        ("data/HMM/neural_alignment/state_transitions/population_metrics/session_{N}/observed_metrics.csv",
         "Per-session observed metric values per (region, state)."),
        ("data/HMM/neural_alignment/state_transitions/population_metrics/session_{N}/shuffle_records.csv",
         "Per-session full shuffle records (100 iterations × all combos)."),
        ("data/HMM/neural_alignment/state_transitions/population_metrics/session_{N}/pass_summary.csv",
         "Per-session pass/fail summary."),
        ("figures/HMM/neural_alignment/state_transitions/population_metrics/master_replication_heatmap.png",
         "Master replication heatmap, all metrics × (region, state) combinations."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    out_path = REPO_ROOT / "data" / "HMM" / "neural_alignment_population_metrics_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
