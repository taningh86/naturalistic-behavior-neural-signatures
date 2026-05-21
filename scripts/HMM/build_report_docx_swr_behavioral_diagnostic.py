"""Build a detailed Word-document report for the SWR behavioral diagnostic.

Reads outputs at:
  data/HMM/neural_alignment/swr_behavioral_diagnostic/
  figures/HMM/neural_alignment/swr_behavioral_diagnostic/

Output: data/HMM/swr_behavioral_diagnostic_report.docx
"""
from pathlib import Path

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO = Path(__file__).resolve().parents[2]
DIAG = REPO / "data/HMM/neural_alignment/swr_behavioral_diagnostic"
FIG = REPO / "figures/HMM/neural_alignment/swr_behavioral_diagnostic"
OUT_DOCX = REPO / "data/HMM/swr_behavioral_diagnostic_report.docx"

SESSIONS = [4, 6, 8, 12, 14, 16]
SESSION_STATE = {4: "fed", 6: "fed", 8: "fed",
                  12: "fasted", 14: "fasted", 16: "fasted"}
REGIONS = ("ACA", "LHA", "RSP")


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True; run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def df_table(doc, df, float_fmt="{:.3f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        c = table.cell(0, j)
        c.text = str(col)
        for r in c.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, (bool, np.bool_)):
                txt = "True" if bool(v) else "False"
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            c = table.cell(i + 1, j)
            c.text = txt
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(8)


def code(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    d1 = pd.read_csv(DIAG / "D1_speed_distribution_per_session.csv")
    d2 = pd.read_csv(DIAG / "D2_lookup_comparison.csv")
    d3 = pd.read_csv(DIAG / "D3_ripple_behavior_reclassified.csv")
    d4 = pd.read_csv(DIAG / "D4_timestamp_alignment.csv")
    d5 = pd.read_csv(DIAG / "D5_waveform_comparison_summary.csv")

    doc = Document()

    # Title
    t = doc.add_heading("SWR Behavioral Diagnostic — Bug Audit & Locomotion Contamination",
                        level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sr = sub.add_run("Verification audit of script 19's behavioral context output. "
                       "Triggered by the observation that 57% of v2 RSP events were "
                       "labeled 'fast_locomotion' — biologically unusual for canonical "
                       "ripples.")
    sr.italic = True; sr.font.size = Pt(11)

    para(doc, "Script: scripts/HMM/19_swr_behavioral_diagnostic.py", size=9)
    para(doc, f"Outputs: {DIAG.as_posix()}", size=9)
    para(doc, "Generated 2026-05-13", size=9)
    doc.add_paragraph()

    # Executive Summary
    heading(doc, "Executive Summary", level=1)
    para(doc, "Four findings.", bold=True)
    bullet(doc,
            "Two confirmed bugs in script 19's behavioral_context() — "
            "(1) zone lookup is 100% broken: all 1799 events stored zone='-1'. "
            "Cause: `d.get(\"zone_int\", default)` while the HMM binned npz key "
            "is `zone`. "
            "(2) event-flag lookups are 100% broken: all dig/feed/rear stored as "
            "False. Cause: `d.get(\"digging_sand\", ...)` etc. while the npz "
            "stores all 7 flags in a single 2D `events` array. Speed lookup is "
            "correct (0/1799 mismatches). Timestamp alignment is clean "
            "(D4: offset ≈ 0 ms in all 6 sessions).")
    bullet(doc,
            "The 5 cm/s 'fast_locomotion' threshold is misleading. Median session "
            "speed is 1.83–4.11 cm/s; the 5 cm/s cutoff covers 20–41% of all "
            "behavioral bins — not the tail of the distribution. Animals spend "
            "a third of their time above the threshold during normal foraging.")
    bullet(doc,
            "BUT the RSP locomotion bias survives an animal-relative correction. "
            "Even when 'fast' is redefined as the top quartile of speeds within "
            "each session (only 25% of bins by definition), 56% of RSP events "
            "still land there — 2.2× enrichment relative to uniform. The pattern "
            "isn't a threshold artifact.")
    bullet(doc,
            "D5 waveform analysis identifies likely motion-artifact contamination "
            "in the fast-locomotion RSP events. The 100-250 Hz bandpass envelope "
            "amplitude during 'fast' events is 0.230 µV (n=179), versus 0.430 µV "
            "during stationary and slow events (n=74+69=143). RSP fast events are "
            "46% weaker than stationary ones. Same direction in LHA (-18%). "
            "Interpretation: about half of RSP \"ripples\" are motion-artifact "
            "energy passing the per-pair 4-SD detector; the stationary/slow subset "
            "is the candidate-real-ripple population.")

    # Method
    heading(doc, "Method", level=1)
    para(doc,
          "Five diagnostics on the v2 SWR event tables, no re-detection. Inputs: "
          "data/HMM/neural_alignment/swr_detection_v2/threshold_02pct/session_{N}/* "
          "and the HMM binned behavior npz files at data/HMM/binned/session_{N}.npz.")
    bullet(doc, "D1 — Per-session speed distribution (HMM 480 ms bins): percentiles + fraction in absolute speed bins.")
    bullet(doc, "D2 — v2-stored vs fresh behavioral lookup, with correct HMM npz keys. Checks zone, speed, dig, feed, rear.")
    bullet(doc, "D3 — Re-classify v2 events using animal-relative locomotion thresholds (session median, p75).")
    bullet(doc, "D4 — Timestamp alignment between behavior and LFP across all 6 sessions.")
    bullet(doc, "D5 — Re-read raw .lf.bin in ±100 ms windows around each event peak, bandpass 100-250 Hz, mean waveform and envelope amplitude per region × locomotion state.")
    bullet(doc, "Master seed reused from script 19 (20260512). No shuffles or null distributions; this is a diagnostic, not a statistical test.")

    # The two bugs
    heading(doc, "Confirmed bugs in script 19", level=1)

    heading(doc, "Bug 1: Zone lookup", level=2)
    code(doc, '# script 19, behavioral_context():\n'
                'zone = np.asarray(d.get("zone_int", np.full_like(trial_time, -1, dtype=int)))')
    para(doc,
          "The default fallback always fires because the HMM binned npz key is "
          "`zone`, not `zone_int`. Result: every event in every v2 "
          "session_{N}_{region}_event_behavior.csv has zone='-1'. "
          "D2 confirmed 1799/1799 = 100% of events affected. Fresh re-lookup "
          "with the correct key recovers the true zone distribution.")

    heading(doc, "Bug 2: Event-flag lookups", level=2)
    code(doc, '# script 19, behavioral_context():\n'
                'dig = np.asarray(d.get("digging_sand", np.zeros_like(trial_time)))\n'
                'feed = np.asarray(d.get("feeding", np.zeros_like(trial_time)))\n'
                'rear = np.asarray(d.get("rearing", np.zeros_like(trial_time)))')
    para(doc,
          "The HMM binned npz stores all 7 event flags in a single 2D array "
          "`events` (shape T × 7) with the column names in `event_names`. "
          "Looking up by individual key names like `digging_sand` always returns "
          "the default (all zeros). Result: every v2 event has dig=False, "
          "feed=False, rear=False. D2 confirms 0/1799 events have any flag True.")

    heading(doc, "Bug 3 (potential): speed cutoffs vs animal speed", level=2)
    para(doc,
          "Not a code bug but a parameter choice issue. The 5 cm/s 'fast_locomotion' "
          "cutoff is well within the body of the speed distribution rather than at "
          "the tail. D1 quantifies this; D3 tests whether the locomotion bias is an "
          "artifact of the choice.")

    # D1
    heading(doc, "D1 — Speed distribution per session", level=1)
    para(doc, "All 6 foraging sessions, HMM 480 ms bins.")
    d1_view = d1[["session", "state", "n_bins", "median", "percentile_75",
                    "percentile_95", "percentile_99",
                    "frac_below_1", "frac_1_to_5", "frac_5_to_10",
                    "frac_10_to_20", "frac_above_20"]].copy()
    df_table(doc, d1_view, float_fmt="{:.3f}")

    para(doc,
          f"Median speed range: {d1['median'].min():.2f} – "
          f"{d1['median'].max():.2f} cm/s. "
          f"Fraction of bins above the v2 5 cm/s 'fast' cutoff "
          f"(= frac_5_to_10 + frac_10_to_20 + frac_above_20): "
          f"{((d1['frac_5_to_10']+d1['frac_10_to_20']+d1['frac_above_20'])*100).min():.1f}% "
          f"to {((d1['frac_5_to_10']+d1['frac_10_to_20']+d1['frac_above_20'])*100).max():.1f}%.")
    bullet(doc,
            "Fed sessions (S4, S6, S8) are more active than fasted (S12, S14, S16) "
            "as expected — medians 3.09-4.11 cm/s vs 1.83-3.02 cm/s.")
    bullet(doc,
            "S14 is the slowest session with median 1.83 cm/s and 25% of bins "
            "below 1 cm/s. Likely a heavy feeding session.")

    add_image(doc, FIG / "D1_speed_distribution_per_session.png", width_in=6.5,
              caption="D1: Per-session speed histograms (log-log). Red dashed "
                       "= session median; black dotted = 5 cm/s (v2 fast cutoff). "
                       "The v2 cutoff sits in the body of the distribution, not "
                       "the tail.")

    # D2
    heading(doc, "D2 — v2 stored vs fresh behavioral lookup", level=1)
    para(doc,
          "Re-looked-up zone, speed, dig, feed, rear from the HMM binned npz "
          "using the correct keys, compared against the v2-stored values from "
          "session_{N}_{region}_event_behavior.csv.")

    n_total = len(d2)
    bad_zone = int((d2["v2_zone"].astype(str).isin(["-1", "-1.0"])).sum())
    speed_mm = int((~d2["match_speed_flag"]).sum())
    zone_mm = int((~d2["match_zone_flag"]).sum())
    summary = pd.DataFrame([
        dict(metric="total events compared", v1_stored=n_total, fresh_lookup=n_total),
        dict(metric="events with v2 zone='-1'",
              v1_stored=f"{bad_zone}/{n_total} (100%)", fresh_lookup="0 (all valid)"),
        dict(metric="speed mismatches", v1_stored=f"{speed_mm}/{n_total}",
              fresh_lookup="—"),
        dict(metric="zone mismatches", v1_stored=f"{zone_mm}/{n_total}",
              fresh_lookup="—"),
    ])
    df_table(doc, summary)

    # Per-region zone breakdowns
    heading(doc, "D2 — Fresh zone-lookup distribution per region", level=2)
    para(doc,
          "Once the correct key is used, the actual zones where events occur are "
          "recoverable. This is data that was completely lost in v2.")

    rows = []
    for region in REGIONS:
        sub = d2[d2.region == region]
        n_reg = len(sub)
        if not n_reg:
            continue
        zc = sub["lookup_zone"].value_counts()
        for zone, count in zc.items():
            rows.append(dict(region=region, zone=zone,
                              n_events=int(count),
                              pct=f"{100*count/n_reg:.1f}%"))
    df_table(doc, pd.DataFrame(rows))
    para(doc,
          "RSP events concentrate at pots (56%) and transition zones (23%). Only "
          "1% at home and 8% at pot_zone. This is the OPPOSITE of canonical "
          "sharp-wave ripples (which classically occur during quiet wakefulness "
          "at the home/rest location). The zonal distribution alone is a flag "
          "that something unusual is going on with RSP.")

    # Per-region event flags
    heading(doc, "D2 — Fresh event-flag counts per region", level=2)
    rows = []
    for region in REGIONS:
        sub = d2[d2.region == region]
        if not len(sub):
            continue
        rows.append(dict(region=region,
                          n_events=len(sub),
                          n_dig=int(sub["fresh_dig"].sum()),
                          n_feed=int(sub["fresh_feed"].sum()),
                          n_rear=int(sub["fresh_rear"].sum())))
    df_table(doc, pd.DataFrame(rows))
    para(doc,
          "33/322 RSP events (10%) co-occur with digging. 0 co-occur with feeding "
          "or rearing. Digging is a high-motor-engagement behavior; this is "
          "consistent with the locomotion bias observed in D3.")

    # D3
    heading(doc, "D3 — Animal-relative reclassification", level=1)
    para(doc,
          "Re-classified v2 events using session-specific percentile thresholds: "
          "stationary = speed < session median; slow = median ≤ speed < session "
          "p75; fast = speed ≥ session p75. This makes the thresholds "
          "animal-relative — 'fast' always means the top quartile of speeds for "
          "this session.")

    # Build comparison table
    rows = []
    for region in REGIONS:
        sub = d3[d3.region == region]
        if not len(sub):
            continue
        old_counts = sub["old_loc_state"].value_counts()
        new_counts = sub["new_loc_state"].value_counts()
        rows.append(dict(
            region=region,
            n_events=len(sub),
            old_stationary=int(old_counts.get("stationary", 0)),
            old_slow=int(old_counts.get("slow_locomotion", 0)),
            old_fast=int(old_counts.get("fast_locomotion", 0)),
            new_stationary=int(new_counts.get("stationary", 0)),
            new_slow=int(new_counts.get("slow", 0)),
            new_fast=int(new_counts.get("fast", 0)),
            new_fast_pct=f"{100*int(new_counts.get('fast', 0))/len(sub):.1f}%",
        ))
    df_table(doc, pd.DataFrame(rows))

    para(doc,
          "Reading: under absolute thresholds, RSP has 57% fast events. Under "
          "animal-relative (top 25%) thresholds, RSP STILL has 56% fast events — "
          "a 2.2× enrichment vs the 25% baseline expected by chance. The bias is "
          "robust. ACA shifts the opposite direction: under animal-relative "
          "thresholds, 73% of ACA's tiny 30-event set is in the stationary "
          "bottom-half (canonical pattern). LHA is intermediate.")

    add_image(doc, FIG / "D3_ripple_locomotion_reclassified.png", width_in=6.5,
              caption="D3: Stacked bar of ripple counts per session per region "
                       "under animal-relative locomotion thresholds. RSP fast "
                       "(pink) dominates in every session.")

    # D4
    heading(doc, "D4 — Timestamp alignment", level=1)
    para(doc,
          "Behavior-bin index 0 should correspond to LFP sample 0 (after "
          "foraging-phase truncation). Offsets here would mean the per-event "
          "behavioral lookup is hitting the wrong time bin systematically.")
    d4_view = d4[["session", "beh_bins", "beh_duration_s", "lfp_total_samples",
                    "lfp_total_duration_s", "truncated_lfp_duration_s",
                    "max_event_peak_time_s", "offset_ms", "alignment_ok_flag"]]
    df_table(doc, d4_view, float_fmt="{:.2f}")
    para(doc,
          "All 6 sessions: offset ≈ 0 ms. Maximum event peak time is within the "
          "behavioral duration in every session. No timestamp bug.")

    # D5
    heading(doc, "D5 — Waveform amplitude across locomotion states (smoking gun)", level=1)
    para(doc,
          "Re-read the raw .lf.bin in a ±100 ms window around each event peak, "
          "applied notch + bipolar + 100-250 Hz bandpass at the native 2500 Hz "
          "sampling, and computed the mean envelope amplitude. Per region × "
          "locomotion state (animal-relative, D3 classification).")

    d5_view = d5.copy()
    d5_view["mean_envelope_uV"] = d5_view["mean_envelope_uV"].round(3)
    df_table(doc, d5_view, float_fmt="{:.3f}")

    para(doc, "Key observations:")
    bullet(doc,
            "RSP fast events: mean envelope 0.230 µV (n=179). Stationary: 0.430 µV "
            "(n=74). Slow: 0.429 µV (n=69). RSP fast events are 46% weaker than "
            "stationary ones, despite outnumbering them 2.4-to-1.")
    bullet(doc,
            "LHA: same direction. Fast 0.126 µV (n=587) vs stationary 0.154 µV "
            "(n=424). 18% weaker during fast.")
    bullet(doc,
            "ACA: small n's (22/6/2). Mean envelope actually increases with speed, "
            "but that's noise — the 2 fast ACA events are outliers.")

    para(doc,
          "Interpretation: real sharp-wave ripples have a stereotyped amplitude "
          "envelope. Motion artifact passing through a 100-250 Hz bandpass "
          "produces a lower-amplitude envelope because the artifact energy is "
          "broadband, not concentrated in the band. Per-pair detection at 4 SD "
          "captures both: it can't distinguish a real ripple from a coincidental "
          "broadband-noise excursion. The amplitude split between stationary and "
          "fast events is direct evidence that the fast events have a different "
          "physical origin.")

    add_image(doc, FIG / "D5_mean_ripple_waveform_per_loc_state.png", width_in=6.5,
              caption="D5: Mean ±SD ripple waveform (100-250 Hz bandpass) per "
                       "region × locomotion state. Visibly higher amplitude in "
                       "RSP stationary and slow rows than RSP fast (rightmost "
                       "panel).")
    add_image(doc, FIG / "D5_mean_ripple_spectrum_per_loc_state.png", width_in=6.5,
              caption="D5: Mean power spectra of event-window LFP. Gray band = "
                       "100-250 Hz ripple band. RSP fast events have lower "
                       "power throughout the band than stationary/slow events.")

    # Revised interpretation
    heading(doc, "Revised interpretation", level=1)
    para(doc,
          "The original script-19 v2 SWR finding ('RSP carries ~7 events/min, "
          "validated by spikes at MW p << 1e-22, modal 181 Hz, ~99 ms duration') "
          "is partially overcounted. The diagnostic suggests:")
    bullet(doc,
            "About 56% of RSP events (n=179/322) occur during animal-relative-fast "
            "locomotion with 46% lower bandpass amplitude than the stationary "
            "subset. These are likely motion-artifact contamination.")
    bullet(doc,
            "The remaining ~44% (n=143 stationary + slow) are the candidate real "
            "ripples. At session-level rates this is closer to ~3 real ripples/min "
            "than the original 7/min.")
    bullet(doc,
            "Spike validation in script 19 used ALL events, mixing real ripples "
            "and artifacts. The MW p<<1e-22 result is still likely valid for the "
            "stationary subset, but should be re-tested on stationary-only events.")
    bullet(doc,
            "The fresh zone lookup (lost in v2 due to bug 1) shows RSP events "
            "concentrate at pots (56%) and transitions (23%) — at task-relevant "
            "zones. About 10% co-occur with digging. This is consistent with "
            "real motivated behavior, not background noise, lending some credence "
            "to the surviving stationary RSP events as biologically meaningful.")
    bullet(doc,
            "The original interpretation that 'RSP shows real ripples at canonical "
            "~181 Hz' likely stands for the stationary subset, but the headline "
            "rate and total count should be revised downward by ~50%.")

    # Action items
    heading(doc, "Action items for downstream pipelines", level=1)
    bullet(doc, "Fix script 19's behavioral_context(): replace `d.get(\"zone_int\", ...)` with `d[\"zone\"]`; replace per-flag lookups with `d[\"events\"][:, j]` indexed by `d[\"event_names\"]`. Re-run on the v2 events to produce corrected event_behavior CSVs.")
    bullet(doc, "Filter v2 SWR events to the stationary/slow subset (animal-relative bottom-75%) before reporting RSP rates and validations. Re-do spike validation on filtered subset.")
    bullet(doc, "Consider adding a motion-artifact prefilter to the per-pair ripple detector: e.g., require event amplitude > some minimum AND co-occurrence with low broadband-power baseline.")
    bullet(doc, "Update memory: original v2 SWR report headline numbers need revision.")
    bullet(doc, "Anyone reading the original script-19 v1 or v2 SWR Word reports should be referred to this diagnostic for context.")

    # Caveats
    heading(doc, "Caveats", level=1)
    bullet(doc,
            "Behavior is at 480 ms resolution (HMM binned npz). A ripple of "
            "~99 ms duration falls inside one behavioral bin, so the speed "
            "value we look up is an average over the 480 ms surrounding the "
            "ripple. A 40 ms raw-behavior xlsx lookup would give finer temporal "
            "resolution. The current speed values are a slight smoothed proxy.")
    bullet(doc,
            "D5 used a representative bipolar pair (median pair_index within "
            "each region) to extract the LFP window. For events where the "
            "ripple was concentrated at a different shank, this representative "
            "pair may underestimate the amplitude. A more thorough analysis "
            "would extract per-event from the contributing pairs.")
    bullet(doc,
            "Animal-relative thresholds (D3) use session-internal percentiles. "
            "If a session is unusually quiet overall, the 'fast' bucket "
            "captures lower absolute speeds than in an active session. The "
            "biological meaning of 'fast' varies session-to-session.")
    bullet(doc,
            "Spike validation was not re-run in this diagnostic. The original "
            "MW p<<1e-22 result for RSP combined real and artifactual events. "
            "Re-running validation on the stationary subset alone is the obvious "
            "next step.")

    # Output files
    heading(doc, "Output files", level=1)
    code(doc, "data/HMM/neural_alignment/swr_behavioral_diagnostic/")
    bullet(doc, "D1_speed_distribution_per_session.csv — per-session percentiles and speed-bin fractions")
    bullet(doc, "D2_lookup_comparison.csv — 1799 events × {v2-stored vs fresh} for speed and zone, plus fresh dig/feed/rear")
    bullet(doc, "D3_ripple_behavior_reclassified.csv — every event with old_loc_state and new_loc_state (animal-relative)")
    bullet(doc, "D4_timestamp_alignment.csv — per-session LFP/behavior duration comparison")
    bullet(doc, "D5_waveform_comparison_summary.csv — mean envelope amplitude per (region, loc_state)")
    code(doc, "figures/HMM/neural_alignment/swr_behavioral_diagnostic/")
    bullet(doc, "D1_speed_distribution_per_session.png — log-log histograms with v2 cutoff and session medians marked")
    bullet(doc, "D3_ripple_locomotion_reclassified.png — stacked bars per region per session under new classification")
    bullet(doc, "D5_mean_ripple_waveform_per_loc_state.png — 3×3 grid, region × loc state, mean ±SD bandpass waveform")
    bullet(doc, "D5_mean_ripple_spectrum_per_loc_state.png — same grid but for multitaper power spectra")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Saved {OUT_DOCX}")


if __name__ == "__main__":
    main()
