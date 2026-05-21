"""Build a Word-document report for script 18:
LFP state-identity (A1) + three-region Granger (A2).

Reads:
  data/HMM/neural_alignment/lfp_state_identity/
  data/HMM/neural_alignment/lfp_three_region_granger/
  figures/HMM/neural_alignment/lfp_state_identity/
  figures/HMM/neural_alignment/lfp_three_region_granger/

Output: data/HMM/lfp_state_identity_and_3region_granger_report.docx
"""
from pathlib import Path

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO_ROOT = Path(__file__).resolve().parents[2]

A1_OUT = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "lfp_state_identity"
A2_OUT = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "lfp_three_region_granger"
A1_FIG = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "lfp_state_identity"
A2_FIG = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "lfp_three_region_granger"

SESSIONS = [4, 6, 8, 12, 14, 16]
SESSION_STATE = {4: "fed", 6: "fed", 8: "fed",
                  12: "fasted", 14: "fasted", 16: "fasted"}
REGIONS = ("ACA", "LHA", "RSP")
BAND_NAMES = ["delta", "theta", "beta", "low_gamma", "high_gamma"]

OUT_DOCX = REPO_ROOT / "data" / "HMM" / "lfp_state_identity_and_3region_granger_report.docx"


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def df_table(doc, df, float_fmt="{:.3f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, (bool, np.bool_)):
                txt = "True" if bool(v) else "False"
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    # Load all CSVs
    a1_k_rep = pd.read_csv(A1_OUT / "A1_kruskal_replication.csv")
    a1_pw_rep = pd.read_csv(A1_OUT / "A1_pairwise_replication.csv")
    a1_bp_all = pd.read_csv(A1_OUT / "A1_band_power_per_category_all_sessions.csv")
    a2_F = pd.read_csv(A2_OUT / "A2_granger_F_table.csv")
    a2_sign = pd.read_csv(A2_OUT / "A2_sign_test.csv")
    a2_rep = pd.read_csv(A2_OUT / "A2_replication.csv")

    # Per-session A1 Kruskal data
    per_sess_k = []
    per_sess_pw = []
    for sn in SESSIONS:
        per_sess_k.append(pd.read_csv(A1_OUT / f"A1_session_{sn}_kruskal.csv").assign(session=sn))
        per_sess_pw.append(pd.read_csv(A1_OUT / f"A1_session_{sn}_pairwise.csv").assign(session=sn))
    a1_k_all = pd.concat(per_sess_k, ignore_index=True)
    a1_pw_all = pd.concat(per_sess_pw, ignore_index=True)

    # Per-session bin counts
    bin_counts_rows = []
    for sn in SESSIONS:
        bp = pd.read_csv(A1_OUT / f"A1_session_{sn}_band_power_per_category.csv")
        # n_bins is constant per category across regions; take from one region
        sub = bp[bp.region == "ACA"][["category", "n_bins"]].drop_duplicates(
            subset=["category"]
        )
        row = {"session": sn, "state": SESSION_STATE[sn]}
        for _, r in sub.iterrows():
            row[r["category"]] = int(r["n_bins"])
        bin_counts_rows.append(row)
    bin_counts_df = pd.DataFrame(bin_counts_rows)[
        ["session", "state", "home", "feeding", "transition_zone"]
    ]

    doc = Document()

    # Title
    t = doc.add_heading("LFP State-Identity & Three-Region Granger", level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sr = sub.add_run("Script 18 — LFP power spectra at HMM behavioral categories "
                       "(home / feeding / transition_zone) and 3-region Granger "
                       "extension (ACA-LHA-RSP)")
    sr.italic = True
    sr.font.size = Pt(11)

    para(doc, "Script: scripts/HMM/18_lfp_state_identity_and_three_region_granger.py", size=9)
    para(doc, f"Outputs: {A1_OUT.as_posix()}", size=9)
    para(doc, f"         {A2_OUT.as_posix()}", size=9)
    para(doc, "Generated 2026-05-11", size=9)
    doc.add_paragraph()

    # Executive Summary
    heading(doc, "Executive Summary", level=1)
    para(doc,
          "Two analyses on the corrected (script 17 v2) bipolar regional LFP traces "
          "(ACA, LHA, RSP at 500 Hz). 6 foraging sessions (S4/6/8 fed, S12/14/16 fasted).",
          bold=False)

    bullet(doc,
            "A1 LFP state-identity: feeding is spectrally distinct from the other "
            "two categories; home and transition_zone are not. Strongest single "
            "signals are ACA high_gamma (Kruskal-Wallis 6/6 sessions ≥ shuffle "
            "p95; home_vs_feeding 6/6; feeding_vs_transition 6/6) and RSP beta "
            "(Kruskal 5/6; feeding_vs_transition 5/6; home_vs_feeding 4/6). "
            "Across all 15 region×band cells in home_vs_transition_zone, 0/15 "
            "replicate ≥4/6 — the two categories look the same at LFP power "
            "level. LHA only weakly carries category identity (high_gamma 4/6 "
            "Kruskal, no pairwise contrast ≥4/6).")
    bullet(doc,
            "A2 three-region LFP Granger: no band × pair combination reaches "
            "binomial p < 0.05 in the sign test across sessions. At n=6 the "
            "minimum two-sided p is 0.031 (requires 6/6). Best cells are 5/6: "
            "ACA-LHA theta (LHA→ACA), ACA-RSP beta (RSP→ACA), ACA-RSP "
            "high_gamma (ACA→RSP), LHA-RSP beta (RSP→LHA) — all p=0.219. "
            "ACA-LHA reproduces script 17 M4's null; extending to RSP shows "
            "the same pattern. Strong bidirectional F in raw values, no "
            "consistent directional asymmetry at LFP envelope level for any pair.")
    bullet(doc,
            "Reconciles with script 16 (spike-PC1 Granger 6/6 ACA→LHA leads, "
            "binom p=0.031): the ACA→LHA directional signal lives at the "
            "spike-population level, not the LFP envelope level, and does not "
            "generalize to ACA-RSP or LHA-RSP at LFP either.")

    # Method
    heading(doc, "Method", level=1)
    bullet(doc, "Sessions: 6 dual-probe foraging — S4, S6, S8 (fed), S12, S14, S16 (fasted).")
    bullet(doc, "Inputs: regional bipolar LFP at 500 Hz from script 17 v2 (preprocessed_v2/).")
    bullet(doc, "HMM bin: 480 ms (= 240 samples at 500 Hz). Multitaper PSD via DPSS NW=3, K=5 tapers. Bands: delta 1-4, theta 4-12, beta 15-30, low_gamma 30-60, high_gamma 60-100 Hz.")
    bullet(doc, "A1 categories: home=S3, feeding=S2, transition_zone=S4 (merged Viterbi from script 08).")
    bullet(doc, "A1 statistics per session: Kruskal-Wallis across 3 categories per (region, band); pairwise Mann-Whitney with BH-FDR within session.")
    bullet(doc, "A1 null: 100 circular-shift shuffles of Viterbi; per (region, band) report whether observed H exceeds shuffle p95.")
    bullet(doc, "A2 segments: S3 stay + 5 s post-exit, runs ≥ 2 × K_PRE bins and S3 portion ≥ 5 s.")
    bullet(doc, "A2 statistics: bandpass + Hilbert envelope per band, z-score, bivariate VAR with BIC lag selection 1-20 samples at 500 Hz (2-40 ms), analytical F and 100 circular-shift shuffles of the source envelope. 3 region pairs × 2 directions × 5 bands × 6 sessions.")
    bullet(doc, "A2 sign test: binomial test of (n_sessions forward_F > reverse_F) vs 0.5 per (band, pair). Replication threshold ≥4/6 sessions.")
    bullet(doc, "Master seed: 20260510. Min bins per category = 30.")

    # A1
    heading(doc, "A1. LFP State-Identity", level=1)

    heading(doc, "A1.a Bin counts per category per session", level=2)
    para(doc,
          "Number of valid (artifact-free) HMM bins assigned to each category. "
          "Fasted sessions (S12/14/16) are feeding-dominant by design; fed "
          "sessions show more balanced distributions.")
    df_table(doc, bin_counts_df, float_fmt="{:.0f}")

    heading(doc, "A1.b Kruskal-Wallis omnibus replication", level=2)
    para(doc,
          "Per (region, band): number of sessions where the observed Kruskal-Wallis "
          "H statistic exceeds the 95th percentile of the circular-shift shuffle "
          "null distribution. Replication threshold ≥4/6.")
    df_table(doc, a1_k_rep, float_fmt="{:.0f}")

    n_rep_aca = int(a1_k_rep[(a1_k_rep.region == "ACA") & a1_k_rep.replicates].shape[0])
    n_rep_lha = int(a1_k_rep[(a1_k_rep.region == "LHA") & a1_k_rep.replicates].shape[0])
    n_rep_rsp = int(a1_k_rep[(a1_k_rep.region == "RSP") & a1_k_rep.replicates].shape[0])
    para(doc,
          f"Summary: ACA {n_rep_aca}/5 bands replicate, LHA {n_rep_lha}/5, "
          f"RSP {n_rep_rsp}/5. RSP carries the most band-level structure. "
          f"ACA high_gamma is the cleanest single cell at 6/6 sessions.")

    add_image(doc, A1_FIG / "A1_replication_heatmap.png", width_in=5.5,
              caption="A1 replication heatmap — n_sessions passing per (region × band, pair).")

    heading(doc, "A1.c Pairwise replication (home vs feeding vs transition_zone)", level=2)
    para(doc,
          "Pairwise Mann-Whitney U per (region, band, pair). Replication = number "
          "of sessions where observed |U − shuffle_median| exceeds shuffle p95. "
          "Threshold ≥4/6.")

    heading(doc, "ACA pairwise", level=3)
    df_table(doc, a1_pw_rep[a1_pw_rep.region == "ACA"][["band", "pair", "n_passing", "replicates"]],
              float_fmt="{:.0f}")
    heading(doc, "LHA pairwise", level=3)
    df_table(doc, a1_pw_rep[a1_pw_rep.region == "LHA"][["band", "pair", "n_passing", "replicates"]],
              float_fmt="{:.0f}")
    heading(doc, "RSP pairwise", level=3)
    df_table(doc, a1_pw_rep[a1_pw_rep.region == "RSP"][["band", "pair", "n_passing", "replicates"]],
              float_fmt="{:.0f}")

    top = a1_pw_rep[a1_pw_rep.n_passing >= 4].sort_values(
        "n_passing", ascending=False
    )[["region", "band", "pair", "n_passing"]].reset_index(drop=True)
    heading(doc, "A1.c Cells reaching ≥4/6 pairwise replication", level=3)
    df_table(doc, top, float_fmt="{:.0f}")

    home_vs_trans = a1_pw_rep[a1_pw_rep.pair == "home_vs_transition_zone"]
    para(doc,
          f"Cells in home_vs_transition_zone with n_passing ≥4/6: "
          f"{int((home_vs_trans.n_passing >= 4).sum())}/{len(home_vs_trans)}. "
          "Home and transition_zone are not separable at the LFP-power level — "
          "the LFP state-identity signal is feeding-vs-the-rest.")

    heading(doc, "A1.d Cross-session band power per category", level=2)
    para(doc,
          "Mean band power per category, averaged across sessions per region. "
          "Error bars = SEM across sessions. log scale on y-axis.")
    add_image(doc, A1_FIG / "A1_spectral_fingerprints_aggregate.png", width_in=6.5,
              caption="A1 aggregate spectral fingerprints — 3 regions × 3 categories")

    # Detail table: cross-session mean power per (region, band, category)
    bp_summary = a1_bp_all.groupby(["region", "band", "category"]).agg(
        cross_session_mean=("mean_power", "mean"),
        cross_session_sem=("mean_power", "sem"),
    ).reset_index()
    heading(doc, "A1.d Cross-session band power table", level=3)
    df_table(doc, bp_summary, float_fmt="{:.4g}")

    heading(doc, "A1.e Per-session spectral fingerprints", level=2)
    para(doc,
          "Per-session spectral fingerprint plots: mean band power per category "
          "with within-bin SEM, log scale. One figure per region per session.")
    for sn in SESSIONS:
        for region in REGIONS:
            fp = A1_FIG / f"A1_session_{sn}_spectral_fingerprint_{region}.png"
            if fp.exists():
                add_image(doc, fp, width_in=4.5,
                          caption=f"S{sn} ({SESSION_STATE[sn]}) — {region}")

    # A2
    heading(doc, "A2. Three-Region LFP Granger Extension", level=1)

    heading(doc, "A2.a S3 segments per session", level=2)
    # Count from F table — one entry per session × band × pair × direction
    segs_per_session = a2_F.groupby("session").size() // (5 * 3 * 2)
    s3_counts = pd.DataFrame({
        "session": segs_per_session.index.astype(int),
        "state": [SESSION_STATE[int(s)] for s in segs_per_session.index],
        "implied_n_pairs_tested": segs_per_session.values,
    })
    # Actual S3 segments come from individual session run logs; report n_sessions tested
    para(doc, f"All {len(SESSIONS)} sessions ran A2 successfully. S3 segment counts "
                 f"(home runs with stay ≥ 5 s + 5 s post-exit) per session "
                 "(from per-session run logs): S4=15, S6=13, S8=12, S12=13, S14=6, S16=12.")

    heading(doc, "A2.b Cross-session F table (one row per session × band × pair × direction)",
              level=2)
    para(doc,
          "Observed Granger F per session × band × pair × direction. Note F values "
          "are large in both directions for most cells — coupling is strong but "
          "bidirectional. Per cell, exceeds_p95 flag indicates whether observed F "
          "exceeds the 95th percentile of 100 source-shuffled F values.")

    # Pivot to make it readable: one section per band, rows = (pair, direction), cols = session
    for band in BAND_NAMES:
        heading(doc, f"A2.b Granger F — {band}", level=3)
        sub = a2_F[a2_F.band == band]
        pivot = sub.pivot_table(
            index=["pair", "direction"], columns="session",
            values="observed_F", aggfunc="first",
        ).reset_index()
        df_table(doc, pivot, float_fmt="{:.2f}")

    heading(doc, "A2.c Sign test across sessions", level=2)
    para(doc,
          "Per (band, pair): count of sessions where forward direction has larger F "
          "than reverse, and the two-sided binomial test against 0.5. Minimum "
          "achievable p at n=6 is 0.031 (requires 6/6 same direction).")
    df_table(doc, a2_sign[["band", "pair", "forward", "n_sessions",
                              "n_forward_leads", "n_reverse_leads", "binom_p"]],
              float_fmt="{:.3f}")

    n_sig = int((a2_sign.binom_p < 0.05).sum())
    para(doc,
          f"Cells with binom p < 0.05: {n_sig}/{len(a2_sign)}. No band × pair "
          "combination reaches significance. Best cells are 5/6 (binom p=0.219): "
          "ACA-LHA theta LHA→ACA, ACA-RSP beta RSP→ACA, ACA-RSP high_gamma "
          "ACA→RSP, LHA-RSP beta RSP→LHA. Notably ACA-RSP and LHA-RSP "
          "high-frequency cells lean toward RSP-input directions (RSP-leading "
          "in beta, ACA-leading only in high_gamma) — suggests RSP acts more "
          "as input than output of the cortical-hypothalamic axis at the LFP "
          "envelope level, but the asymmetry does not survive the sign test.")

    add_image(doc, A2_FIG / "A2_sign_test_summary.png", width_in=6.5,
              caption="A2 sign test — three panels, one per pair. Stacked bar = "
                       "n_sessions forward vs reverse leading. Annotation = binom p.")

    heading(doc, "A2.d Replication: shuffle-pass counts per direction × band × pair", level=2)
    para(doc,
          "Number of sessions where observed F > shuffle p95 in each direction. "
          "Replicates flag = n_passing ≥ 4/6. This captures whether the F is "
          "non-trivially above the source-shuffle null at the per-session level.")
    df_table(doc, a2_rep, float_fmt="{:.0f}")

    add_image(doc, A2_FIG / "A2_replication_heatmap.png", width_in=6.5,
              caption="A2 replication heatmap — rows = (pair × direction), columns = band.")

    heading(doc, "A2.e Per-session F values, all pairs", level=2)
    add_image(doc, A2_FIG / "A2_F_per_session_per_pair.png", width_in=6.5,
              caption="A2 F values per session × band × direction. Bars: solid = "
                       "forward, hatched = reverse. Three panels by region pair.")

    # Interpretation
    heading(doc, "Interpretation", level=1)
    bullet(doc,
            "LFP power discriminates feeding from non-feeding behavioral categories, "
            "but home and transition_zone are spectrally indistinguishable. The 0/15 "
            "cells in home_vs_transition_zone with ≥4/6 replication confirms this is "
            "robust, not a power-of-test artifact.")
    bullet(doc,
            "ACA high_gamma is the single cleanest band-level state-identity signal. "
            "Likely reflects increased local population spiking and synaptic activity "
            "during feeding bouts (high gamma tracks population firing rate).")
    bullet(doc,
            "RSP beta and gamma show feeding selectivity. The RSP regional mean here "
            "includes channels at y ≥ 2500 µm (the static-geometry split), which "
            "is broader than RSP proper (spike-defined y > 4680 µm). Some of this "
            "signal may originate from white matter or transitional tissue between "
            "LHA and RSP.")
    bullet(doc,
            "LHA high_gamma 4/6 Kruskal but no pairwise contrast ≥4/6 means the LHA "
            "omnibus difference is real but does not localize to a single category "
            "contrast — likely a mix of feeding>others in some sessions and "
            "transition>others in others. Not a clean state-identity signal at LFP "
            "level for LHA.")
    bullet(doc,
            "Three-region Granger reproduces script 17 M4's null for ACA-LHA and "
            "extends the same null to ACA-RSP and LHA-RSP. LFP envelope coupling is "
            "strong in both directions but symmetric — no consistent directional "
            "lead. This is independent of the spike-PC1 result (script 16, 6/6 "
            "ACA→LHA, p=0.031), which therefore represents a spike-population-level "
            "signal not captured by LFP power envelopes.")
    bullet(doc,
            "Combining A1 + A2: the LFP at the ACA / LHA / RSP triplet carries "
            "category-level information (mostly feeding-related) but not "
            "directional information at the envelope level. State identity is "
            "decodable; state transition causality at LFP level is not.")

    # Caveats
    heading(doc, "Caveats", level=1)
    bullet(doc,
            "Bin-level power values within a single Viterbi run are temporally "
            "autocorrelated, partially violating Kruskal-Wallis independence. "
            "The 100 circular-shift shuffle null is the primary safeguard; "
            "shuffle preserves the autocorrelation structure of the LFP itself "
            "(only the Viterbi label is rolled), so the null distribution "
            "captures the chance level under that exact autocorrelation.")
    bullet(doc,
            "LHA regional LFP includes bipolar pairs at y < 2500 µm, spanning LHA "
            "tissue (y < 345 µm by spike depth) plus white matter / non-LHA "
            "tissue between LHA proper and RSP. RSP regional LFP includes pairs "
            "at y ≥ 2500 µm, spanning white matter + RSP tissue (y > 4680 µm by "
            "spike depth). The y=2500 µm boundary is the IMRO physical midpoint, "
            "not a histologically-defined tissue boundary.")
    bullet(doc,
            "Sample size: n=6 sessions per category contrast and per pair. Sign "
            "test minimum two-sided p with n=6 is 2/64 = 0.031 (requires 6/6 same "
            "direction). Replication threshold 4/6 corresponds to two-sided binom "
            "p = 0.69 — descriptive, not inferential.")
    bullet(doc,
            "All within-session pairwise comparisons use BH-FDR. No cross-analysis "
            "or cross-region multiple-comparison correction at the cell level. "
            "Replication across sessions is the protection against false positives.")
    bullet(doc,
            "BIC lag selection in A2 produced different selected lags per session × "
            "band × direction (range 1-20 samples = 2-40 ms). The per-cell F is "
            "comparable to its own shuffle but not directly across cells with "
            "different lag depths.")
    bullet(doc,
            "Common-input confound: bidirectional Granger may reflect a shared "
            "third driver (e.g. global state, arousal) rather than direct "
            "ACA-LHA-RSP coupling. This is unavoidable in a 3-region recording.")

    # Output Files
    heading(doc, "Output Files", level=1)
    code(doc, "data/HMM/neural_alignment/lfp_state_identity/")
    bullet(doc, "A1_session_{N}_kruskal.csv, A1_session_{N}_pairwise.csv — per-session detail")
    bullet(doc, "A1_session_{N}_band_power_per_category.csv — per-region per-band per-category mean/sd/n")
    bullet(doc, "A1_kruskal_replication.csv — cross-session Kruskal replication count")
    bullet(doc, "A1_pairwise_replication.csv — cross-session pairwise replication count")
    bullet(doc, "A1_band_power_per_category_all_sessions.csv — long-form band power")
    code(doc, "figures/HMM/neural_alignment/lfp_state_identity/")
    bullet(doc, "A1_replication_heatmap.png — cell-wise replication heatmap")
    bullet(doc, "A1_spectral_fingerprints_aggregate.png — cross-session fingerprints")
    bullet(doc, "A1_session_{N}_spectral_fingerprint_{ACA,LHA,RSP}.png — per-session fingerprints")
    code(doc, "data/HMM/neural_alignment/lfp_three_region_granger/")
    bullet(doc, "A2_session_{N}_granger.csv — per-session per-band per-pair-direction F + shuffle p95")
    bullet(doc, "A2_granger_F_table.csv — long form across sessions")
    bullet(doc, "A2_sign_test.csv — binomial test per (band, pair)")
    bullet(doc, "A2_replication.csv — shuffle-pass counts per direction × band × pair")
    code(doc, "figures/HMM/neural_alignment/lfp_three_region_granger/")
    bullet(doc, "A2_F_per_session_per_pair.png — 3-panel bar chart per region pair")
    bullet(doc, "A2_replication_heatmap.png — replication heatmap")
    bullet(doc, "A2_sign_test_summary.png — sign test summary, 3 panels")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Saved {OUT_DOCX}")


if __name__ == "__main__":
    main()
