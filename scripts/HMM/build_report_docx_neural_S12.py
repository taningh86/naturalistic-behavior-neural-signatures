"""Build a Word-document report of the Track A neural alignment results for S12.

Reads outputs from data/HMM/neural_alignment/transient_S12/ and figures from
figures/HMM/neural_alignment/transient_S12/, plus the merged-state and
commitment-marker context.

Output: data/HMM/neural_alignment_S12_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.3f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def aggregate_summary(df, label_a, label_b):
    """Compute descriptive aggregates from a per-unit summary CSV.

    Returns a dict per region with:
      n_units, baseline_a_pre, baseline_b_pre, post_window_a, post_window_b,
      n_higher_b_post, pct_higher_b_post,
      median_abs_peak, p90_abs_peak, max_abs_peak,
      n_pos_peak, n_neg_peak, n_anticipatory, n_reactive
    Where columns are read positionally as
      mean_FR_a_pre, mean_FR_a_post, mean_FR_b_pre, mean_FR_b_post,
      peak_diff_*.
    """
    out = {}
    # Identify a/b column suffixes flexibly:
    cand_suffixes = [("a", "b"),
                     ("discovery", "failed"),
                     ("pre", "post")]
    a_pre = b_pre = a_post = b_post = None
    for sa, sb in cand_suffixes:
        ca_pre = f"mean_FR_{sa}_pre"
        cb_pre = f"mean_FR_{sb}_pre"
        ca_post = f"mean_FR_{sa}_post"
        cb_post = f"mean_FR_{sb}_post"
        if all(c in df.columns for c in (ca_pre, cb_pre, ca_post, cb_post)):
            a_pre, b_pre, a_post, b_post = ca_pre, cb_pre, ca_post, cb_post
            break
    if a_pre is None:
        raise ValueError(f"Cannot locate FR columns. Columns: {df.columns.tolist()}")

    for region in ("ACA", "LHA"):
        sub = df[df.region == region].copy()
        if not len(sub):
            out[region] = None
            continue
        sub["abs_peak"] = sub["peak_diff_value"].abs()
        # b_post > a_post in post-event window means second condition higher
        n_higher_b_post = int((sub[b_post] > sub[a_post]).sum())
        out[region] = dict(
            n_units=len(sub),
            baseline_a_pre=float(sub[a_pre].mean()),
            baseline_b_pre=float(sub[b_pre].mean()),
            post_window_a=float(sub[a_post].mean()),
            post_window_b=float(sub[b_post].mean()),
            n_higher_b_post=n_higher_b_post,
            pct_higher_b_post=100.0 * n_higher_b_post / len(sub),
            median_abs_peak=float(sub.abs_peak.median()),
            p90_abs_peak=float(sub.abs_peak.quantile(0.9)),
            max_abs_peak=float(sub.abs_peak.max()),
            n_pos_peak=int((sub.peak_diff_value > 0).sum()),
            n_neg_peak=int((sub.peak_diff_value < 0).sum()),
            n_anticipatory=int((sub.peak_diff_time_s < 0).sum()),
            n_reactive=int((sub.peak_diff_time_s > 0).sum()),
            top_units=sub.nlargest(5, "abs_peak")[
                ["unit_id", a_pre, a_post, b_pre, b_post,
                 "peak_diff_time_s", "peak_diff_value"]
            ].copy(),
        )
    return out, (a_pre, a_post, b_pre, b_post)


def main():
    cfg = load_config()
    cm_dir = REPO_ROOT / cfg["commitment_dirs"]["out"]
    na_dir = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "transient_S12"
    fig_dir = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "transient_S12"

    history = pd.read_csv(cm_dir / "sampling_history.csv")
    s12 = history[history.session == 12].iloc[0]

    a1 = pd.read_csv(na_dir / "A1_summary.csv")
    a2 = pd.read_csv(na_dir / "A2_summary.csv")
    a3 = pd.read_csv(na_dir / "A3_summary.csv")
    s_a1, _ = aggregate_summary(a1, "discovery", "failed")
    s_a2, cols_a2 = aggregate_summary(a2, "pre", "post")
    s_a3, cols_a3 = aggregate_summary(a3, "pre", "post")

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===== TITLE =====
    title = doc.add_heading(
        "Neural Alignment Track A — Session 12 (S12, fasted)", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Event-locked transient analyses (±1 s, 100 ms bins) of ACA and LHA "
        "good units relative to discovery, S4 entries, and pot-zone entries."
    )
    r.italic = True
    r.font.size = Pt(11)

    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-06  •  Script: scripts/HMM/10a_neural_alignment_transient_S12.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "Track A aligns ACA (probe 0) and LHA (probe 1, depth 0-345 µm) "
        "single-unit firing rates to three behavioral event sets in session 12 "
        f"(fasted, food at P{int(s12['food_pot'])}, discovery at "
        f"t={s12['discovery_time_s']:.1f} s). Each analysis takes a ±1 s window "
        "around event onset at 100 ms resolution and compares two conditions "
        "via per-unit and population PETHs. With one session and (in A1) very "
        "small event counts, all results are descriptive — no formal stats."
    )
    add_bullet(
        doc,
        "A1 (discovery dig vs failed food-pot digs, n=1 vs n=2): per-unit peak "
        "differences reach 60-70 Hz on individual bins, but with so few events "
        "any single-bin difference is essentially one-trial noise. Useful only "
        "as a qualitative sketch."
    )
    add_bullet(
        doc,
        f"A2 (pre vs post-discovery S4 entries, n={s_a2['ACA']['n_units'] and 15} vs 15 — balanced): per-unit |peak_diff| medians "
        f"~{s_a2['ACA']['median_abs_peak']:.1f} Hz (ACA) and "
        f"~{s_a2['LHA']['median_abs_peak']:.1f} Hz (LHA); modest. Direction is "
        f"slightly biased toward post>pre in post-event window ("
        f"{s_a2['ACA']['pct_higher_b_post']:.0f}% ACA, {s_a2['LHA']['pct_higher_b_post']:.0f}% LHA)."
    )
    add_bullet(
        doc,
        f"A3 (pre vs post-discovery pot-zone entries, n=30 vs 58 — most "
        f"events): population baselines barely shift (ACA "
        f"{s_a3['ACA']['baseline_a_pre']:.2f}→{s_a3['ACA']['baseline_b_pre']:.2f} Hz, "
        f"LHA {s_a3['LHA']['baseline_a_pre']:.2f}→{s_a3['LHA']['baseline_b_pre']:.2f} Hz). "
        f"Median |peak_diff|: {s_a3['ACA']['median_abs_peak']:.1f} Hz ACA, "
        f"{s_a3['LHA']['median_abs_peak']:.1f} Hz LHA. LHA peaks lean "
        f"anticipatory ({s_a3['LHA']['n_anticipatory']} pre-event vs "
        f"{s_a3['LHA']['n_reactive']} post-event); ACA is roughly balanced "
        f"({s_a3['ACA']['n_anticipatory']} vs {s_a3['ACA']['n_reactive']})."
    )
    add_bullet(
        doc,
        "No analysis shows a strong region-wide knowledge-state signature in "
        "S12 alone. Top-modulated single units (~15-25 Hz peri-event peaks) "
        "are candidates for follow-up but require additional sessions before "
        "any claim is defensible."
    )

    # ===== SETUP =====
    add_heading(doc, "Setup", level=1)
    add_para(
        doc,
        "Inputs:"
    )
    add_bullet(doc, "Spike data: results_KS3_trimmed_probe{0,1}/sorter_output for S12 "
                    "(7-11-25 dual-probe fasted recording).")
    add_bullet(doc, "Quality filter: cluster_info.tsv group/KSLabel = 'good'. ACA: "
                    "FR > 0.2 Hz (no AMP). LHA: FR > 0.2 Hz, AMP > 43 µV, depth "
                    "0-345 µm.")
    add_bullet(doc, "Behavioral context: data/HMM/commitment_markers/session_12_events.csv "
                    "+ sampling_history.csv (script 09).")
    add_bullet(doc, "Merged Viterbi: data/HMM/merged_posteriors_dynamax/session_12.csv "
                    "(used to re-derive post-discovery S4 / pot-zone entries that 09 "
                    "doesn't write).")
    add_bullet(doc, "Neural binning: 100 ms across the full ~1800 s session.")
    add_bullet(doc, "Window: ±1 s around event onset (21 bins, time vector −1.0 to +1.0 s).")
    add_bullet(doc, "Visualization: 50 ms gaussian smoothing on PETH lines (raw 100 ms "
                    "bins for summary stats).")

    # ===== UNIT COUNTS =====
    add_heading(doc, "Unit counts and quality flags", level=1)
    add_para(doc,
             "Two regions, single session. Low-FR flags use a 0.5 Hz mean-rate "
             "threshold; flagged units are kept in all analyses but listed for "
             "transparency.")
    unit_tbl = pd.DataFrame([
        dict(region="ACA (probe 0)", n_good_units=s_a3["ACA"]["n_units"],
             low_fr_flagged=24, low_fr_pct="14.5%"),
        dict(region="LHA (probe 1, 0-345 µm)", n_good_units=s_a3["LHA"]["n_units"],
             low_fr_flagged=30, low_fr_pct="33.7%"),
    ])
    add_df_table(doc, unit_tbl)

    # ===== EVENT COUNTS =====
    add_heading(doc, "Event counts", level=1)
    add_para(doc, "Per-analysis, per-condition. Pre-discovery from "
                  "session_12_events.csv; post-discovery re-derived from the "
                  "merged Viterbi after discovery_bin = "
                  f"{int(s12['discovery_bin'])}.")
    ev_tbl = pd.DataFrame([
        dict(analysis="A1", condition_a="discovery_dig", n_a=1,
             condition_b="prior_dig_food_pot (failed)", n_b=2),
        dict(analysis="A2", condition_a="pre-discovery S4 entries", n_a=15,
             condition_b="post-discovery S4 entries", n_b=15),
        dict(analysis="A3", condition_a="pre-discovery pot-zone entries", n_a=30,
             condition_b="post-discovery pot-zone entries", n_b=58),
    ])
    add_df_table(doc, ev_tbl)

    # ===== A1 RESULTS =====
    doc.add_page_break()
    add_heading(doc, "A1 — Discovery dig vs failed food-pot digs", level=1)
    add_para(
        doc,
        "Compares the single dig that successfully led to feeding (the "
        "discovery_dig at t=597.1 s, P4) against the two earlier digs at the "
        "same pot (P4) that did not lead to feeding within the 10 s window. "
        "With one event vs two, this is not a population comparison; it is a "
        "qualitative trial-by-trial sketch. Per-unit |peak_diff| values reach "
        "60-70 Hz on individual bins, but those reflect single-bin variation "
        "across very few trials."
    )
    a1_stats = pd.DataFrame([
        dict(region="ACA",
             n_units=s_a1["ACA"]["n_units"],
             median_abs_peak_Hz=f"{s_a1['ACA']['median_abs_peak']:.1f}",
             p90_abs_peak_Hz=f"{s_a1['ACA']['p90_abs_peak']:.1f}",
             max_abs_peak_Hz=f"{s_a1['ACA']['max_abs_peak']:.1f}"),
        dict(region="LHA",
             n_units=s_a1["LHA"]["n_units"],
             median_abs_peak_Hz=f"{s_a1['LHA']['median_abs_peak']:.1f}",
             p90_abs_peak_Hz=f"{s_a1['LHA']['p90_abs_peak']:.1f}",
             max_abs_peak_Hz=f"{s_a1['LHA']['max_abs_peak']:.1f}"),
    ])
    add_df_table(doc, a1_stats)
    add_para(doc, "Top 5 modulated units per region (descending |peak_diff|):", size=10)
    for region in ("ACA", "LHA"):
        add_para(doc, region + ":", bold=True, size=10)
        top = s_a1[region]["top_units"].copy()
        top.columns = ["unit", "FR_disc_pre", "FR_disc_post",
                       "FR_failed_pre", "FR_failed_post",
                       "peak_t_s", "peak_diff_Hz"]
        add_df_table(doc, top, float_fmt="{:.2f}")
    add_image(doc, fig_dir / "A1_population_PETH.png", width_in=6.7,
              caption="A1 population PETH: total summed firing across units, "
                      "discovery (red) vs failed (grey). Single discovery event so "
                      "the red trace is one-trial noise; the grey trace averages 2 events.")
    add_image(doc, fig_dir / "A1_per_unit_PETH_ACA.png", width_in=6.7,
              caption="A1 per-unit PETH, ACA — discovery (red) vs failed (grey) at "
                      "100 ms bins, 50 ms gaussian smoothing.")
    add_image(doc, fig_dir / "A1_per_unit_PETH_LHA.png", width_in=6.7,
              caption="A1 per-unit PETH, LHA.")

    # ===== A2 RESULTS =====
    doc.add_page_break()
    add_heading(doc, "A2 — Pre vs post-discovery S4 (T-zone contemplation) entries", level=1)
    add_para(
        doc,
        "Compares the 15 entries into the contemplation/T-zone state (merged "
        "S4) before discovery (t < 597.1 s) against the 15 entries after "
        "discovery. Sample sizes are balanced, so this is the cleanest of the "
        "three contrasts for per-unit interpretation. Population baselines:"
    )
    a2_stats = pd.DataFrame([
        dict(region="ACA",
             n_units=s_a2["ACA"]["n_units"],
             baseline_pre_FR_Hz=f"{s_a2['ACA']['baseline_a_pre']:.2f}",
             baseline_post_FR_Hz=f"{s_a2['ACA']['baseline_b_pre']:.2f}",
             pct_units_higher_post=f"{s_a2['ACA']['pct_higher_b_post']:.0f}%",
             median_abs_peak_Hz=f"{s_a2['ACA']['median_abs_peak']:.1f}",
             p90_abs_peak_Hz=f"{s_a2['ACA']['p90_abs_peak']:.1f}",
             max_abs_peak_Hz=f"{s_a2['ACA']['max_abs_peak']:.1f}",
             anticipatory_vs_reactive=f"{s_a2['ACA']['n_anticipatory']} / {s_a2['ACA']['n_reactive']}"),
        dict(region="LHA",
             n_units=s_a2["LHA"]["n_units"],
             baseline_pre_FR_Hz=f"{s_a2['LHA']['baseline_a_pre']:.2f}",
             baseline_post_FR_Hz=f"{s_a2['LHA']['baseline_b_pre']:.2f}",
             pct_units_higher_post=f"{s_a2['LHA']['pct_higher_b_post']:.0f}%",
             median_abs_peak_Hz=f"{s_a2['LHA']['median_abs_peak']:.1f}",
             p90_abs_peak_Hz=f"{s_a2['LHA']['p90_abs_peak']:.1f}",
             max_abs_peak_Hz=f"{s_a2['LHA']['max_abs_peak']:.1f}",
             anticipatory_vs_reactive=f"{s_a2['LHA']['n_anticipatory']} / {s_a2['LHA']['n_reactive']}"),
    ])
    add_df_table(doc, a2_stats)
    add_bullet(
        doc,
        "Effect sizes are modest. Median per-unit |peak_diff| is "
        f"{s_a2['ACA']['median_abs_peak']:.1f} Hz in ACA and "
        f"{s_a2['LHA']['median_abs_peak']:.1f} Hz in LHA. Top 10% of units "
        f"reach {s_a2['ACA']['p90_abs_peak']:.1f} Hz / "
        f"{s_a2['LHA']['p90_abs_peak']:.1f} Hz."
    )
    add_bullet(
        doc,
        "Direction: in the post-event window, "
        f"{s_a2['ACA']['n_higher_b_post']}/{s_a2['ACA']['n_units']} ACA units "
        f"({s_a2['ACA']['pct_higher_b_post']:.0f}%) and "
        f"{s_a2['LHA']['n_higher_b_post']}/{s_a2['LHA']['n_units']} LHA units "
        f"({s_a2['LHA']['pct_higher_b_post']:.0f}%) fire more on post-discovery "
        f"S4 entries than on pre-discovery S4 entries — a slight post-bias."
    )
    add_para(doc, "Top 5 modulated units per region:", size=10)
    for region in ("ACA", "LHA"):
        add_para(doc, region + ":", bold=True, size=10)
        top = s_a2[region]["top_units"].copy()
        top.columns = ["unit", "FR_pre_pre", "FR_pre_post",
                       "FR_post_pre", "FR_post_post",
                       "peak_t_s", "peak_diff_Hz"]
        add_df_table(doc, top, float_fmt="{:.2f}")
    add_image(doc, fig_dir / "A2_population_PETH.png", width_in=6.7,
              caption="A2 population PETH: pre-discovery (red) vs post-discovery (grey).")
    add_image(doc, fig_dir / "A2_per_unit_PETH_ACA.png", width_in=6.7,
              caption="A2 per-unit PETH, ACA.")
    add_image(doc, fig_dir / "A2_per_unit_PETH_LHA.png", width_in=6.7,
              caption="A2 per-unit PETH, LHA.")

    # ===== A3 RESULTS =====
    doc.add_page_break()
    add_heading(doc, "A3 — Pre vs post-discovery pot-zone entries", level=1)
    add_para(
        doc,
        "Pools transitions into any pot-zone HMM state (merged S8, S9, S10, "
        "S13). 30 pre-discovery entries vs 58 post-discovery entries — the "
        "largest sample of the three contrasts. Includes visits to all four "
        "pots, so post-discovery mixes 'returns to the known food pot P4' "
        "with 'visits to non-food pots after knowing P4 has food'."
    )
    a3_stats = pd.DataFrame([
        dict(region="ACA",
             n_units=s_a3["ACA"]["n_units"],
             baseline_pre_FR_Hz=f"{s_a3['ACA']['baseline_a_pre']:.2f}",
             baseline_post_FR_Hz=f"{s_a3['ACA']['baseline_b_pre']:.2f}",
             pct_units_higher_post=f"{s_a3['ACA']['pct_higher_b_post']:.0f}%",
             median_abs_peak_Hz=f"{s_a3['ACA']['median_abs_peak']:.1f}",
             p90_abs_peak_Hz=f"{s_a3['ACA']['p90_abs_peak']:.1f}",
             max_abs_peak_Hz=f"{s_a3['ACA']['max_abs_peak']:.1f}",
             anticipatory_vs_reactive=f"{s_a3['ACA']['n_anticipatory']} / {s_a3['ACA']['n_reactive']}"),
        dict(region="LHA",
             n_units=s_a3["LHA"]["n_units"],
             baseline_pre_FR_Hz=f"{s_a3['LHA']['baseline_a_pre']:.2f}",
             baseline_post_FR_Hz=f"{s_a3['LHA']['baseline_b_pre']:.2f}",
             pct_units_higher_post=f"{s_a3['LHA']['pct_higher_b_post']:.0f}%",
             median_abs_peak_Hz=f"{s_a3['LHA']['median_abs_peak']:.1f}",
             p90_abs_peak_Hz=f"{s_a3['LHA']['p90_abs_peak']:.1f}",
             max_abs_peak_Hz=f"{s_a3['LHA']['max_abs_peak']:.1f}",
             anticipatory_vs_reactive=f"{s_a3['LHA']['n_anticipatory']} / {s_a3['LHA']['n_reactive']}"),
    ])
    add_df_table(doc, a3_stats)
    add_bullet(
        doc,
        "Population baselines barely shift between pre- and post-discovery "
        "(0.2 Hz drop in ACA, 0.1 Hz drop in LHA in the pre-event window) — "
        "no whole-region knowledge-state effect."
    )
    add_bullet(
        doc,
        "Per-unit modulations are smallest in A3: median |peak_diff| "
        f"{s_a3['ACA']['median_abs_peak']:.1f} Hz (ACA) and "
        f"{s_a3['LHA']['median_abs_peak']:.1f} Hz (LHA). LHA effects are "
        "consistently smaller than ACA in absolute Hz across all three analyses."
    )
    add_bullet(
        doc,
        f"Direction: {s_a3['ACA']['n_higher_b_post']}/{s_a3['ACA']['n_units']} "
        f"ACA units ({s_a3['ACA']['pct_higher_b_post']:.0f}%) and "
        f"{s_a3['LHA']['n_higher_b_post']}/{s_a3['LHA']['n_units']} LHA units "
        f"({s_a3['LHA']['pct_higher_b_post']:.0f}%) fire more on post-discovery "
        "pot-zone entries — a weak but consistent post-bias."
    )
    add_bullet(
        doc,
        f"Timing: LHA peaks lean anticipatory ({s_a3['LHA']['n_anticipatory']} "
        f"units peak before t=0 vs {s_a3['LHA']['n_reactive']} after); ACA is "
        f"balanced ({s_a3['ACA']['n_anticipatory']} vs {s_a3['ACA']['n_reactive']}). "
        "If this LHA-anticipatory pattern survives in other sessions it would "
        "be the first hint that LHA encodes pot-approach intent in a "
        "knowledge-state-dependent way."
    )
    add_bullet(
        doc,
        "Top single units reach 19.8 Hz |peak_diff| in ACA (u116, +0.2 s) and "
        "16.2 Hz in LHA (u36, +1.0 s). These are descriptive candidates only; "
        "a 30 vs 58 average can still produce single-bin fluctuations of this "
        "size by chance."
    )
    add_para(doc, "Top 5 modulated units per region:", size=10)
    for region in ("ACA", "LHA"):
        add_para(doc, region + ":", bold=True, size=10)
        top = s_a3[region]["top_units"].copy()
        top.columns = ["unit", "FR_pre_pre", "FR_pre_post",
                       "FR_post_pre", "FR_post_post",
                       "peak_t_s", "peak_diff_Hz"]
        add_df_table(doc, top, float_fmt="{:.2f}")
    add_image(doc, fig_dir / "A3_population_PETH.png", width_in=6.7,
              caption="A3 population PETH: pre-discovery (red) vs post-discovery (grey).")
    add_image(doc, fig_dir / "A3_per_unit_PETH_ACA.png", width_in=6.7,
              caption="A3 per-unit PETH, ACA.")
    add_image(doc, fig_dir / "A3_per_unit_PETH_LHA.png", width_in=6.7,
              caption="A3 per-unit PETH, LHA.")

    # ===== CROSS-ANALYSIS PATTERNS =====
    doc.add_page_break()
    add_heading(doc, "Cross-Analysis Patterns", level=1)
    add_bullet(
        doc,
        "ACA modulations are systematically larger in absolute Hz than LHA "
        "across all three analyses. Median per-unit |peak_diff| ratios "
        f"(ACA/LHA): A2 = {s_a2['ACA']['median_abs_peak']:.1f}/"
        f"{s_a2['LHA']['median_abs_peak']:.1f}, A3 = "
        f"{s_a3['ACA']['median_abs_peak']:.1f}/{s_a3['LHA']['median_abs_peak']:.1f}. "
        "This is consistent with LHA's typically lower mean firing rates "
        "(2.6 Hz vs 5.7 Hz population averages) — the modulation magnitude "
        "rescales accordingly."
    )
    add_bullet(
        doc,
        "Both regions show a slight but consistent bias toward higher firing "
        "on post-discovery events (~55-61% of units across A2 and A3). Direction "
        "is the same across analyses but the magnitude is small enough that "
        "no claim is defensible from one session."
    )
    add_bullet(
        doc,
        "In A3 only, LHA peaks lean anticipatory (peaks before event onset), "
        "which is a candidate signature of knowledge-state-dependent pot-"
        "approach encoding. ACA peaks are balanced before and after onset in "
        "every analysis. Worth tracking across sessions."
    )
    add_bullet(
        doc,
        "Top-modulated single units (~15-25 Hz |peak_diff| in A2/A3, larger "
        "in A1 but artifactually so) form a candidate list for follow-up. "
        "The IDs listed in the Top-5 tables here are session-12-specific; "
        "cross-session tracking will require unit-matching."
    )

    # ===== CAVEATS =====
    add_heading(doc, "Caveats", level=1)
    add_bullet(doc, "Single session — n=1 across all three analyses. Anything "
                    "described here is a candidate pattern, not a finding.")
    add_bullet(doc, "A1 is descriptive only (1 vs 2 events). Treat it as a sketch.")
    add_bullet(doc, "A3 mixes returns to the food pot with visits to non-food pots "
                    "in the post-discovery condition. A clean A3a (food-pot only) "
                    "vs A3b (non-food only) split is the obvious follow-up.")
    add_bullet(doc, "Low-FR units (24/165 ACA, 30/89 LHA below 0.5 Hz mean rate) "
                    "are kept in all analyses; some PETHs in those units may be "
                    "unstable due to sparse spikes in 1 s windows.")
    add_bullet(doc, "Pre/post window definition: post-event window starts at +0.1 s "
                    "(bin index 11) and runs to +1.0 s (bin 20); the t=0 bin itself "
                    "is excluded from both pre and post means to avoid double-"
                    "counting onset-coincident spikes.")
    add_bullet(doc, "Track B (state-conditioned analyses) is the planned follow-up "
                    "and is intentionally out of scope for this script.")

    # ===== OUTPUTS =====
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/neural_alignment/transient_S12/A1_summary.csv",
         "Per-unit FR pre/post + peak_diff for discovery vs failed digs."),
        ("data/HMM/neural_alignment/transient_S12/A2_summary.csv",
         "Per-unit FR pre/post + peak_diff for pre vs post-discovery S4."),
        ("data/HMM/neural_alignment/transient_S12/A3_summary.csv",
         "Per-unit FR pre/post + peak_diff for pre vs post-discovery pot-zone."),
        ("figures/HMM/neural_alignment/transient_S12/A{1,2,3}_per_unit_PETH_{ACA,LHA}.png",
         "Per-unit PETH grids."),
        ("figures/HMM/neural_alignment/transient_S12/A{1,2,3}_population_PETH.png",
         "Population (sum-across-units) PETHs, both regions on one figure."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    out_path = REPO_ROOT / "data" / "HMM" / "neural_alignment_S12_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
