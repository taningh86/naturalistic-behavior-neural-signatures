"""Build a Word-document report of the 10d B2/B4 fake-discovery shuffle (S12).

Reads outputs from data/HMM/neural_alignment/shuffle_control_B2_B4_S12/ and
figures from figures/HMM/neural_alignment/shuffle_control_B2_B4_S12/.

Output: data/HMM/neural_alignment_S12_shuffle_B2_B4_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.2f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    cfg = load_config()
    sh_dir = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "shuffle_control_B2_B4_S12"
    fig_dir = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "shuffle_control_B2_B4_S12"

    df_b2 = pd.read_csv(sh_dir / "shuffle_B2_summary.csv")
    df_b2_per_state = pd.read_csv(sh_dir / "shuffle_B2_per_state.csv")
    df_b4 = pd.read_csv(sh_dir / "shuffle_B4_summary.csv")
    df_sig = pd.read_csv(sh_dir / "shuffle_B4_state_significance.csv")

    OBS_B2_ACA, OBS_B2_LHA = 142, 74
    N_ACA, N_LHA = 165, 89
    REAL_DISCOVERY_BIN = 1244

    def stats_block(vals, observed, N):
        mean = float(vals.mean())
        ci_lo = float(np.percentile(vals, 2.5))
        ci_hi = float(np.percentile(vals, 97.5))
        rng = (int(vals.min()), int(vals.max()))
        pct = float((vals <= observed).mean() * 100)
        return dict(
            observed=f"{observed}/{N} ({observed/N*100:.0f}%)",
            shuffle_mean=f"{mean:.1f}",
            ci_95=f"[{ci_lo:.0f}, {ci_hi:.0f}]",
            shuffle_range=f"[{rng[0]}, {rng[1]}]",
            obs_pctile=f"{pct:.0f}",
        )

    s_b2_aca = stats_block(df_b2["n_sig_aca"].values, OBS_B2_ACA, N_ACA)
    s_b2_lha = stats_block(df_b2["n_sig_lha"].values, OBS_B2_LHA, N_LHA)

    aca_sig = df_sig[df_sig.region == "ACA"].sort_values("observed_shift",
                                                          ascending=False)
    lha_sig = df_sig[df_sig.region == "LHA"].sort_values("observed_shift",
                                                          ascending=False)
    aca_pass = aca_sig[aca_sig["exceeds_p95"]]
    lha_pass = lha_sig[lha_sig["exceeds_p95"]]
    aca_below_median = aca_sig[aca_sig.obs_pctile < 50.0]
    lha_below_median = lha_sig[lha_sig.obs_pctile < 50.0]

    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    # ===== TITLE =====
    title = doc.add_heading(
        "B2/B4 Shuffle Control — Pre/Post-Discovery Axis (Session 12)", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Tests Track B's pre/post-discovery findings (B2 within-state firing-"
        "rate shift, B4 PC centroid shift) against a fake-discovery shuffle null."
    )
    r.italic = True
    r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-06  •  Script: scripts/HMM/10d_shuffle_control_B2_B4_S12.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "Track B (10b) reported strong pre/post-discovery effects: 86% ACA / "
        "83% LHA units with FDR-significant within-state firing-rate change "
        "(B2), and per-state PC1-3 centroid shifts of 1.5-1.6 in the pot-zone "
        "states S8 and S9 in BOTH regions (B4). Neither was shuffle-controlled "
        "until now. The previous shuffle (10c) tested the state-encoding axis; "
        "this script (10d) tests the orthogonal pre/post-discovery axis by "
        "replacing the real discovery boundary (bin 1244, t=597.1 s) with "
        "100 random 'fake discoveries' uniformly drawn from [500, 3250] (with "
        "±20 bins around 1244 excluded). The Viterbi sequence and unit firing "
        "rates are NOT shuffled — only the pre/post label is randomized."
    )
    add_para(doc, "Headline result, in three sentences:", bold=True, size=11)
    add_bullet(
        doc,
        f"B2 ACA: observed {s_b2_aca['observed']} sits at the "
        f"{s_b2_aca['obs_pctile']}th percentile of the shuffle null (shuffle "
        f"mean {s_b2_aca['shuffle_mean']}). The real ACA result is BELOW the "
        "typical fake-discovery result. Track B's ACA B2 finding is artifact.",
    )
    add_bullet(
        doc,
        f"B2 LHA: observed {s_b2_lha['observed']} at the "
        f"{s_b2_lha['obs_pctile']}th percentile (shuffle mean "
        f"{s_b2_lha['shuffle_mean']}). Above the mean but not above the 95th "
        "percentile — marginal, does not survive the standard threshold.",
    )
    add_bullet(
        doc,
        f"B4: NO ACA state exceeds shuffle 95th percentile. In LHA, "
        f"{len(lha_pass)}/{len(lha_sig)} states pass — including the headline "
        "S8 (1.61, 99th pctile) and S9 (1.48, 100th pctile). The LHA pot-zone "
        "claim survives; the ACA equivalent does not.",
    )
    add_para(
        doc,
        "Net revision: the only Track B pre/post-discovery finding that "
        "survives in S12 is LHA pot-zone (S8, S9, plus S7 and S10) PC "
        "centroid shifts. ACA pre/post-discovery effects across both B2 and "
        "B4 are essentially null — they were autocorrelation/chance. The "
        "10b 'S8/S9 shift in BOTH regions' headline was half right.",
        bold=False, size=11,
    )

    # ===== METHOD =====
    add_heading(doc, "Method", level=1)
    add_bullet(doc, "100 fake-discovery iterations, master seed = 20260507.")
    add_bullet(doc, "Per iteration: fake_boundary_bin uniform in [500, 3250], "
                    "rejection-sampling out ±20 bins around the real "
                    f"discovery (bin {REAL_DISCOVERY_BIN}). The exclusion zone "
                    "is small enough that the random-draw distribution is "
                    "essentially uniform over the eligible range.")
    add_bullet(doc, "B2: per (unit, state) Mann-Whitney U on bin-level firing "
                    "rates, fake-pre vs fake-post, with ≥30 bins required on "
                    "each side. FDR within region across all (unit, state) "
                    "p-values. Count units with ≥1 FDR-significant state.")
    add_bullet(doc, "B4: PCA loadings are FIXED — computed once on the real "
                    "z-scored rates, no refit per shuffle. Each fake boundary "
                    "splits the (n_bins, N_PC) projection into fake-pre and "
                    "fake-post; per-state centroids are taken in PC1-3 with "
                    "≥5 bins required on each side; centroid shift = "
                    "Euclidean distance.")
    add_bullet(doc, "Observed pass uses the real discovery bin (1244) "
                    "through the same code path, so observed and shuffle "
                    "values are directly comparable.")

    # ===== B2 RESULTS =====
    doc.add_page_break()
    add_heading(doc, "B2 — Within-state pre/post Mann-Whitney", level=1)
    b2_tbl = pd.DataFrame([
        dict(region="ACA", **s_b2_aca),
        dict(region="LHA", **s_b2_lha),
    ])
    b2_tbl.columns = ["region", "observed", "shuffle mean", "95% CI",
                       "shuffle range", "obs pctile within null"]
    add_df_table(doc, b2_tbl)
    add_bullet(
        doc,
        "ACA observed (142/165) is BELOW the shuffle mean (146.7), at the "
        "25th percentile of the null distribution. **Three quarters of the "
        "fake-discovery iterations produce MORE significant units than the "
        "real discovery.** In the language of typical permutation testing, "
        "this is the opposite of significance: the real boundary appears to "
        "do worse than chance. The ACA B2 finding from 10b is an artifact "
        "of the within-region power × state-occupancy structure, not a "
        "genuine pre/post-discovery effect.",
    )
    add_bullet(
        doc,
        "LHA observed (74/89) sits at the 90th percentile — above the "
        "shuffle mean but below the conventional 95th-percentile threshold. "
        "There is a hint of a real LHA effect, but it does not clear the "
        "shuffle null. This contrasts with the strong LHA result for B4 "
        "(below); the discrepancy suggests that LHA's pre/post-discovery "
        "signature is structural (PC-space geometry) rather than a uniform "
        "bin-level firing-rate shift across many cells.",
    )

    add_image(doc, fig_dir / "shuffle_B2_distributions.png", width_in=6.7,
              caption="B2 unit-count distributions across 100 fake-discovery "
                      "iterations (purple histograms). Red line = observed "
                      "count from real (real-discovery) data. ACA observed "
                      "(left) is in the lower tail; LHA observed (right) sits "
                      "in the upper tail but inside the 95% CI of the null.")

    # ===== B4 RESULTS =====
    doc.add_page_break()
    add_heading(doc, "B4 — Per-state PC centroid pre/post shift", level=1)
    add_para(
        doc,
        "Each state's pre/post centroid shift is compared against its own "
        "shuffle null distribution (centroids computed from fake-pre and "
        "fake-post bins of that state). 'Exceeds p95' = observed shift "
        "above the 95th percentile of the shuffle distribution for that "
        "state. The shuffle preserves state encoding intact (Viterbi labels "
        "and PC loadings unchanged), so this isolates the pre/post-discovery "
        "axis as the source of any centroid difference."
    )
    add_para(doc, "ACA per-state results (top 6 by observed shift):", bold=True, size=10)
    aca_disp = aca_sig.head(6)[["state", "observed_shift", "shuffle_mean",
                                  "shuffle_p95", "obs_pctile",
                                  "exceeds_p95"]].copy()
    aca_disp.columns = ["state", "observed", "shuffle mean", "shuffle p95",
                         "obs pctile", "exceeds p95"]
    aca_disp["state"] = aca_disp["state"].apply(lambda k: f"S{int(k)}")
    add_df_table(doc, aca_disp)
    add_bullet(
        doc,
        f"NO ACA state exceeds its shuffle 95th percentile "
        f"({len(aca_pass)}/{len(aca_sig)} pass). The headline 10b states "
        "(S0, S8, S9, S10) are all in the 43-61st percentile of their own "
        "fake-discovery nulls — completely typical for circular-shift-style "
        "drift in PC space.",
    )
    if len(aca_below_median):
        states_below = ", ".join(f"S{int(s)} ({p:.0f}%)"
                                  for s, p in zip(aca_below_median["state"],
                                                   aca_below_median["obs_pctile"]))
        add_bullet(
            doc,
            f"In fact, {len(aca_below_median)} ACA states fall BELOW the shuffle "
            f"median (50th percentile): {states_below}. Strong indication that "
            "the ACA B4 result is essentially random.",
        )

    add_para(doc, "LHA per-state results (top 6 by observed shift):", bold=True, size=10)
    lha_disp = lha_sig.head(6)[["state", "observed_shift", "shuffle_mean",
                                  "shuffle_p95", "obs_pctile",
                                  "exceeds_p95"]].copy()
    lha_disp.columns = ["state", "observed", "shuffle mean", "shuffle p95",
                         "obs pctile", "exceeds p95"]
    lha_disp["state"] = lha_disp["state"].apply(lambda k: f"S{int(k)}")
    add_df_table(doc, lha_disp)
    pass_states = ", ".join(f"S{int(s)}" for s in lha_pass["state"])
    add_bullet(
        doc,
        f"{len(lha_pass)}/{len(lha_sig)} LHA states exceed their shuffle 95th "
        f"percentile: {pass_states}. **The headline pot-zone states S8 and S9 "
        "BOTH survive** — S8 at the 99th percentile (observed 1.61, p95 "
        "1.56), S9 at the 100th percentile (observed 1.48, p95 1.44). S7 and "
        "S10 also pass. This is the cleanest defensible finding from Track B "
        "in S12.",
    )
    add_bullet(
        doc,
        "The LHA-pot-zone result holds with intact state encoding (Viterbi "
        "labels are not shuffled). What changes between fake and real "
        "boundaries is whether the boundary aligns with the real discovery. "
        "The fact that the real discovery produces a larger PC1-3 centroid "
        "shift for these LHA states than 95% of fake boundaries indicates "
        "that the same pot-zone behavior carries different LHA neural "
        "content before vs after the mouse knows the food location.",
    )

    add_image(doc, fig_dir / "shuffle_B4_per_state.png", width_in=6.7,
              caption="Per-state pre/post centroid shift in PC1-3 space. "
                      "Red bars = observed (real discovery). Grey bars = "
                      "shuffle 95th percentile. White dots = shuffle mean. "
                      "Red asterisks = states where observed exceeds shuffle "
                      "p95. Top: ACA (no asterisks). Bottom: LHA (S7, S8, S9, "
                      "S10 marked).")

    # ===== WHAT THIS CHANGES =====
    doc.add_page_break()
    add_heading(doc, "Revised Track B picture (after 10c + 10d)", level=1)
    add_para(
        doc,
        "Combining this 10d shuffle with the earlier 10c shuffle gives a "
        "comprehensive null-controlled view of every Track B claim:"
    )
    rows = [
        ["Track B claim (from 10b)", "Status after shuffle"],
        ["B1: 99% units state-selective (ANOVA)",
         "Inflated. 10c per-unit test: 84% ACA, 46% LHA actually exceed shuffle null."],
        ["B3: 100% units have ≥1 sig GLM coef on state",
         "Inflated. Shuffle floor already 95%; observed barely above. ACA 84% / LHA 46% per-unit defensible."],
        ["B2: 86% ACA / 83% LHA show pre/post within-state shift",
         "ACA at 25th pctile of fake-discovery null (BELOW mean — ARTIFACT). LHA at 90th, marginal."],
        ["B4 ACA: S0/S8/S9 PC centroid shifts 1.6-1.9",
         "NO ACA state passes. S0 / S1 / S2 / S4 fall below shuffle median."],
        ["B4 LHA: S8/S9 PC centroid shifts ~1.5-1.6",
         "✓ S7, S8, S9, S10 all exceed shuffle 95th percentile. SURVIVES."],
    ]
    rev_tbl = pd.DataFrame(rows[1:], columns=rows[0])
    add_df_table(doc, rev_tbl)
    doc.add_paragraph()
    add_bullet(
        doc,
        "**The single defensible cognitive-state finding from S12 Track B**: "
        "LHA pot-zone states (S8 and S9, plus S7 and S10) carry a pre/post-"
        "discovery PC centroid shift that exceeds the fake-discovery null. "
        "The same behavior (being in a pot zone) corresponds to different "
        "LHA population neural states before vs after the mouse knows where "
        "the food is.",
    )
    add_bullet(
        doc,
        "ACA carries a robust state code (10c per-unit test, 84%) but NO "
        "robust pre/post-discovery effect (10d B2 below null, B4 within "
        "null). This is a meaningful asymmetry: ACA encodes 'what behavioral "
        "state am I in' generically; LHA encodes 'what does this behavior "
        "mean given my knowledge of food location'.",
    )
    add_bullet(
        doc,
        "Multi-session replication is the natural next step: a focused "
        "test would check whether the same LHA pot-zone states show >shuffle "
        "centroid shifts in S6, S8, S14, S16 (the other foraging sessions "
        "with successful discoveries). S4 should be excluded from that test "
        "given its manual-override discovery time.",
    )
    add_bullet(
        doc,
        "Caveat: this fake-discovery null is tighter than the 10c circular-"
        "shift null because the Viterbi labels remain aligned with neural "
        "data. So passing 10d is a STRONGER claim than passing 10c. The LHA "
        "pot-zone result clears both nulls, which is the strongest evidence "
        "available from a single session.",
    )

    # ===== CAVEATS =====
    add_heading(doc, "Caveats", level=1)
    add_bullet(doc, "100 shuffles supports a 95% CI but is sparse for higher "
                    "percentiles. The 95th-percentile threshold has Monte-Carlo "
                    "noise; an additional run with more iterations would tighten "
                    "the marginal cases (LHA S6 was at the 83rd percentile, just "
                    "below threshold).")
    add_bullet(doc, "B2 and B4 are not independent — both depend on the same "
                    "neural rates and Viterbi labels, only the boundary is "
                    "shuffled. Don't multiply the p-values.")
    add_bullet(doc, "PCA loadings are fixed at the real-data fit. If the goal "
                    "were to test whether the PCA structure ITSELF differs "
                    "pre/post, a different test (block-resample PCA + measure "
                    "subspace alignment) would be needed.")
    add_bullet(doc, "States with insufficient bins on either side of the fake "
                    "boundary are skipped per iteration. For a fake boundary "
                    "very close to either end, multiple states drop out — the "
                    "[500, 3250] range was chosen to keep most states testable.")
    add_bullet(doc, "Single session. The headline LHA pot-zone result needs "
                    "multi-session replication before it can be claimed as a "
                    "biological signature.")

    # ===== OUTPUTS =====
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/neural_alignment/shuffle_control_B2_B4_S12/shuffle_B2_summary.csv",
         "Per-iteration B2 unit counts per region, fake-boundary bin."),
        ("data/HMM/neural_alignment/shuffle_control_B2_B4_S12/shuffle_B2_per_state.csv",
         "Per-iteration per-region per-state count of FDR-sig (unit, state) pairs."),
        ("data/HMM/neural_alignment/shuffle_control_B2_B4_S12/shuffle_B4_summary.csv",
         "Per-iteration per-state PC1-3 centroid shift (long format)."),
        ("data/HMM/neural_alignment/shuffle_control_B2_B4_S12/shuffle_B4_state_significance.csv",
         "Per state per region: observed shift, shuffle mean, p95, observed "
         "percentile, exceeds_p95 flag."),
        ("figures/HMM/neural_alignment/shuffle_control_B2_B4_S12/shuffle_B2_distributions.png",
         "B2 unit-count histograms with observed marker."),
        ("figures/HMM/neural_alignment/shuffle_control_B2_B4_S12/shuffle_B4_per_state.png",
         "B4 per-state observed vs shuffle p95 bar charts (ACA top, LHA bottom)."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    out_path = REPO_ROOT / "data" / "HMM" / "neural_alignment_S12_shuffle_B2_B4_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
