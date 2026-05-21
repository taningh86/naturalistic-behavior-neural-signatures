"""Build a Word-document report of the multi-session Track B + shuffle controls.

Reads outputs from data/HMM/neural_alignment/track_B_all_sessions/ + S12 from
the original 10b/10c/10d output dirs.

Output: data/HMM/neural_alignment_track_B_all_sessions_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.1f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


# ---- Compute per-session preferred-state aggregates ----
def state_label(prof_row):
    """Short behavioral gloss for a merged-state profile row."""
    zone_keys = ["home", "transition", "pot", "pot_zone", "arena", "other"]
    zone_p = {z: prof_row[f"zone_{z}_prob"] for z in zone_keys}
    top_zone = max(zone_p, key=zone_p.get)
    desc = f"{top_zone}({zone_p[top_zone]:.2f})"
    events = []
    for ev in ["digging_sand", "feeding", "rearing",
               "contemplation_at_transition", "exploration_at_transition"]:
        if prof_row[f"event_{ev}_prob"] > 0.30:
            short = ev.split("_")[0]
            events.append(f"{short}={prof_row[f'event_{ev}_prob']:.2f}")
    if events:
        desc += " + " + ",".join(events)
    return desc


def main():
    cfg = load_config()
    base_out = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "track_B_all_sessions"
    base_fig = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "track_B_all_sessions"
    s12_state_cond = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "state_conditioned_S12"
    cm_dir = REPO_ROOT / "data" / "HMM" / "commitment_markers"

    cross = pd.read_csv(base_out / "cross_session_summary.csv")
    rep = pd.read_csv(base_out / "replication_count_per_state.csv")

    # Merged state behavioral content
    prof = pd.read_csv(REPO_ROOT / cfg["merge_dirs"]["state_profiles_csv"])
    K = len(prof)
    state_descriptions = [state_label(prof.iloc[k]) for k in range(K)]

    # Per-session preferred-state distributions for ACA + LHA
    session_pref_aca = {}
    session_pref_lha = {}
    new_sessions = [4, 6, 8, 14, 16]
    for sn in new_sessions:
        df_a = pd.read_csv(base_out / f"session_{sn}" / "B1_selectivity_summary_ACA.csv")
        df_l = pd.read_csv(base_out / f"session_{sn}" / "B1_selectivity_summary_LHA.csv")
        session_pref_aca[sn] = np.bincount(df_a["preferred_state"].values, minlength=K)
        session_pref_lha[sn] = np.bincount(df_l["preferred_state"].values, minlength=K)
    s12_sel = pd.read_csv(s12_state_cond / "B1_selectivity_summary.csv")
    session_pref_aca[12] = np.bincount(
        s12_sel[s12_sel.region == "ACA"]["preferred_state"].values, minlength=K)
    session_pref_lha[12] = np.bincount(
        s12_sel[s12_sel.region == "LHA"]["preferred_state"].values, minlength=K)

    sess_order = [4, 6, 8, 12, 14, 16]
    session_state = {4: "fed", 6: "fed", 8: "fed",
                       12: "fasted", 14: "fasted", 16: "fasted"}
    session_food = {4: "P4", 6: "P4", 8: "P3", 12: "P4", 14: "P3", 16: "P3"}

    aca_totals = sum(session_pref_aca[sn] for sn in sess_order)
    lha_totals = sum(session_pref_lha[sn] for sn in sess_order)

    # ---- Build doc ----
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_heading(
        "Track B Multi-Session Replication — All 6 Foraging Sessions", level=0
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "State-conditioned analyses + circular-shift + fake-discovery shuffle "
        "controls applied to S4, S6, S8, S12, S14, S16."
    )
    r.italic = True
    r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-06  •  Script: scripts/HMM/11_track_B_all_sessions.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "Track B (state-conditioned analyses) and its two shuffle controls "
        "(10c circular-shift on Viterbi+posteriors, 10d fake-discovery on "
        "pre/post boundary) were extended from session 12 alone to all six "
        "foraging sessions: S4, S6, S8 (fed) and S12, S14, S16 (fasted). "
        "S10 was excluded (no-food extinction; no discovery event). S4 used "
        "its manual-override discovery time. S12 was loaded from the "
        "existing 10b/10c/10d outputs without re-running. Across all 6 "
        "sessions the picture has three clean strands:"
    )
    add_para(doc, "Three findings in plain language:", bold=True, size=11)
    add_bullet(
        doc,
        "**State encoding is real and replicates per-unit, with a robust "
        "ACA > LHA gradient.** Per-unit shuffle pass rate (10c) is 63-84% "
        "in ACA and 23-61% in LHA across all 6 sessions. The 99-100% "
        "bulk-count results from the original 10b are autocorrelation "
        "artifact in every session — confirmed.",
    )
    add_bullet(
        doc,
        "**Pre/post-discovery effects DO NOT replicate at conventional "
        "thresholds.** The S12 LHA pot-zone result (S7/S8/S9/S10 all "
        "exceeding shuffle p95 in B4) appears in NO other session — fed "
        "or fasted. Across all 6 sessions, no state passes B4 shuffle p95 "
        "in more than 1 session. B2 percentile never exceeds the 95th in "
        "any session. The S12 'cognitive-state signature in pot-zone "
        "behavior' headline was session-specific.",
    )
    add_bullet(
        doc,
        "**The distribution of preferred states IS replicable and "
        "biologically coherent.** Across 1047 ACA units in 6 sessions, S6 "
        "(pot + digging) is the #1 preferred state in 5/6 sessions; ACA's "
        "top-3 across-session totals are S6 (digging, n=168), S4 (transition + "
        "contemplation, n=134), S12 (transition, n=98). Across 578 LHA "
        "units, the dominant preferences are S2 (pot+feeding, n=136), S11 "
        "(arena+feeding, n=85), S6 (pot+digging, n=80) — together 52% of "
        "all LHA units. ACA encodes deliberate-action states; LHA encodes "
        "consummatory states. This matches expected anatomy and is the "
        "headline finding to take forward.",
    )

    # ===== SESSION INVENTORY =====
    add_heading(doc, "Session inventory", level=1)
    inv_rows = []
    for sn in sess_order:
        row = cross[cross.sn == sn].iloc[0]
        history = pd.read_csv(cm_dir / "sampling_history.csv")
        h = history[history.session == sn].iloc[0]
        inv_rows.append(dict(
            session=sn,
            state=session_state[sn],
            food_pot=session_food[sn],
            discovery_time_s=f"{h['discovery_time_s']:.1f}" if pd.notna(h["discovery_time_s"]) else "—",
            discovery_method=h["discovery_method"],
            n_ACA=int(row["n_aca"]),
            n_LHA=int(row["n_lha"]),
        ))
    add_df_table(doc, pd.DataFrame(inv_rows))
    add_bullet(doc, "S4 uses manual-override discovery (raw EthoVision feeding onset, "
                    "not HMM-state-based; 14.9 s dig→feed lag).")
    add_bullet(doc, "S14 has very early discovery (54 s); pre-discovery has only 112 "
                    "HMM bins, so B2 sample sizes per state are reduced.")
    add_bullet(doc, "S16 has the smallest LHA cohort (54 units). All unit counts are "
                    "after KSLabel='good' + region-specific FR/AMP/depth filters.")

    # ===== METHOD =====
    add_heading(doc, "Method", level=1)
    add_bullet(doc, "Spike data binned at 480 ms (matched to HMM bin grid), "
                    "aggregated from the 100 ms grid by averaging 100 ms counts "
                    "within each 480 ms HMM bin.")
    add_bullet(doc, "B1 — one-way ANOVA per unit across the K=14 merged HMM "
                    "states; FDR within region.")
    add_bullet(doc, "B2 — Mann-Whitney U per (unit, state), bin-level firing rates "
                    "pre vs post-discovery; ≥30 bins/side; FDR within region.")
    add_bullet(doc, "B3 — Poisson GLM with state posteriors as predictors and "
                    "the most-occupied state per session dropped as reference; "
                    "custom IRLS (~5-10 ms/fit). FDR within region; |z|>2.5 "
                    "threshold.")
    add_bullet(doc, "B4 — PCA on z-scored rates (top 5 PCs); per-state pre/post "
                    "centroid in PC1-3 space; Euclidean distance.")
    add_bullet(doc, "10c shuffle — 100 circular shifts of Viterbi+posteriors "
                    "(offset ∈ [100, T−100]). Re-run B1+B3; per-unit max |z| "
                    "comparison.")
    add_bullet(doc, "10d shuffle — 100 fake-discovery boundaries (uniform in "
                    "[500, T−500] excluding ±20 around real). Re-run B2+B4 "
                    "with PCA loadings fixed.")
    add_bullet(doc, "S12 results loaded from the existing 10b/10c/10d output "
                    "directories rather than recomputed (matches the original "
                    "single-session pipeline exactly).")

    # ===== STATE ENCODING (cross-session B1) =====
    doc.add_page_break()
    add_heading(doc, "Cross-session state encoding (B1 + per-unit shuffle)", level=1)
    add_para(
        doc,
        "The 10b/10c shuffle test asks: per-unit, does the unit's max |z| "
        "from the GLM exceed the unit's own shuffle 95th percentile? This "
        "is the cleanest per-unit test of state encoding because it is "
        "robust to autocorrelation and to the bulk-count power inflation."
    )
    rows = []
    for sn in sess_order:
        row = cross[cross.sn == sn].iloc[0]
        rows.append(dict(
            session=f"S{sn}",
            state=session_state[sn],
            n_ACA=int(row["n_aca"]),
            ACA_above_p95=int(row["b1_per_unit_above_aca"]),
            ACA_pct=f"{row['b1_per_unit_above_aca']/row['n_aca']*100:.0f}%",
            n_LHA=int(row["n_lha"]),
            LHA_above_p95=int(row["b1_per_unit_above_lha"]),
            LHA_pct=f"{row['b1_per_unit_above_lha']/row['n_lha']*100:.0f}%",
        ))
    add_df_table(doc, pd.DataFrame(rows))
    add_bullet(doc, "ACA per-unit pass rate is consistently 63-84% across all 6 "
                    "sessions. This is a robust, replicable per-unit signal of "
                    "state encoding above autocorrelation null.")
    add_bullet(doc, "LHA per-unit pass rate is more variable (23-61%) and "
                    "consistently lower than ACA in every session. The "
                    "cross-session asymmetry ACA > LHA on per-unit state "
                    "encoding is one of the clearest replicable findings.")
    add_bullet(doc, "S14 has the lowest LHA per-unit pass (23%) — likely partly "
                    "due to the unbalanced pre/post split affecting the GLM, but "
                    "also probably reflecting a real session-level difference.")
    add_bullet(doc, "The bulk-count headline from 10b ('99% of units state-"
                    "selective') reproduces in every session — and is artifact in "
                    "every session. Without the per-unit shuffle test the result "
                    "is misleading.")

    # ===== PREFERRED-STATE DISTRIBUTIONS =====
    doc.add_page_break()
    add_heading(doc, "Preferred-state distributions across sessions", level=1)
    add_para(
        doc,
        "For each session and each region, every unit's 'preferred state' is "
        "the merged HMM state with the unit's highest mean firing rate. The "
        "table below shows per-session counts plus the across-session total "
        "for each merged state, with the behavioral content inferred from "
        "the merged-state emission profiles."
    )

    # ACA preferred-state table
    add_para(doc, "ACA — preferred-state counts per session (1047 units across 6 sessions):",
             bold=True, size=10)
    rows_a = []
    for k in range(K):
        row = dict(
            state=f"S{k}",
            S4=int(session_pref_aca[4][k]),
            S6=int(session_pref_aca[6][k]),
            S8=int(session_pref_aca[8][k]),
            S12=int(session_pref_aca[12][k]),
            S14=int(session_pref_aca[14][k]),
            S16=int(session_pref_aca[16][k]),
            total=int(aca_totals[k]),
            behavior=state_descriptions[k],
        )
        rows_a.append(row)
    add_df_table(doc, pd.DataFrame(rows_a), float_fmt="{:.0f}")

    # LHA preferred-state table
    add_para(doc, "LHA — preferred-state counts per session (578 units across 6 sessions):",
             bold=True, size=10)
    rows_l = []
    for k in range(K):
        row = dict(
            state=f"S{k}",
            S4=int(session_pref_lha[4][k]),
            S6=int(session_pref_lha[6][k]),
            S8=int(session_pref_lha[8][k]),
            S12=int(session_pref_lha[12][k]),
            S14=int(session_pref_lha[14][k]),
            S16=int(session_pref_lha[16][k]),
            total=int(lha_totals[k]),
            behavior=state_descriptions[k],
        )
        rows_l.append(row)
    add_df_table(doc, pd.DataFrame(rows_l), float_fmt="{:.0f}")

    add_para(doc, "Top across-session preferred states (with biological content):",
             bold=True, size=10)
    aca_order = np.argsort(-aca_totals)
    lha_order = np.argsort(-lha_totals)
    top_rows = []
    for rnk in range(5):
        ka = int(aca_order[rnk])
        kl = int(lha_order[rnk])
        top_rows.append(dict(
            rank=rnk + 1,
            ACA_state=f"S{ka}",
            ACA_total=int(aca_totals[ka]),
            ACA_pct_of_units=f"{aca_totals[ka]/aca_totals.sum()*100:.0f}%",
            ACA_behavior=state_descriptions[ka],
            LHA_state=f"S{kl}",
            LHA_total=int(lha_totals[kl]),
            LHA_pct_of_units=f"{lha_totals[kl]/lha_totals.sum()*100:.0f}%",
            LHA_behavior=state_descriptions[kl],
        ))
    add_df_table(doc, pd.DataFrame(top_rows))

    add_bullet(
        doc,
        "ACA — S6 (pot + digging) is the #1 preferred state in 5/6 sessions "
        f"(only S16 has S12-transition tied/above). Across-session total "
        f"168/{int(aca_totals.sum())} units (16%) prefer S6. The ACA top-3 are "
        "all 'deliberate behavior' states: digging, transition+contemplation, "
        "and a generic transition state.",
    )
    add_bullet(
        doc,
        "LHA — S2 (pot + feeding) is #1 in 4/6 sessions; S11 (other arena + "
        f"feeding) is #1 in S6 and S16. Together S2+S11 absorb {int(lha_totals[2]+lha_totals[11])}"
        f"/{int(lha_totals.sum())} ({(lha_totals[2]+lha_totals[11])/lha_totals.sum()*100:.0f}%) "
        "of all LHA preferences. Adding S6 (pot+digging) brings this to "
        f"{int(lha_totals[2]+lha_totals[11]+lha_totals[6])}/{int(lha_totals.sum())} "
        f"({(lha_totals[2]+lha_totals[11]+lha_totals[6])/lha_totals.sum()*100:.0f}%) — "
        "consummatory and food-acquisition states dominate.",
    )
    add_bullet(
        doc,
        "Functional anatomy story: ACA neurons preferentially encode "
        "attention/action-selection states (digging, deliberation at the "
        "T-zone). LHA neurons preferentially encode feeding and food-pot "
        "engagement. Both match expected regional roles. This is the "
        "**replicable, biologically interpretable** Track B signal.",
    )
    add_image(doc, base_fig / "preferred_state_counts_all_sessions.png",
              width_in=6.7,
              caption="Per-session B1 preferred-state distribution. Blue title = "
                      "fed, red title = fasted. ACA in blue, LHA in red within "
                      "each panel.")

    # ===== PRE/POST DISCOVERY (B2 + B4 NULL RESULT) =====
    doc.add_page_break()
    add_heading(doc, "Pre/post-discovery effects (B2 + B4) — REPLICATION FAILED", level=1)
    add_para(
        doc,
        "The 10d fake-discovery shuffle randomizes only the pre/post boundary "
        "while keeping Viterbi labels and PCA loadings intact. It is the "
        "appropriate null for any pre/post-discovery contrast. In S12 alone, "
        "B4 LHA pot-zone states S7/S8/S9/S10 cleared the 95th percentile of "
        "this null. The expectation under the multi-session pipeline was that "
        "fasted sessions (S12, S14, S16) would replicate this LHA pot-zone "
        "result and fed sessions would not. Reality is different."
    )

    # B4 cross-session table
    add_para(doc, "Per-session B4 pre/post centroid shift passes (states with "
                  "observed shift > shuffle 95th percentile):", bold=True, size=10)
    rows_b4 = []
    for sn in sess_order:
        row = cross[cross.sn == sn].iloc[0]
        aca_p = str(row["b4_aca_pass"]) if pd.notna(row["b4_aca_pass"]) else ""
        lha_p = str(row["b4_lha_pass"]) if pd.notna(row["b4_lha_pass"]) else ""
        rows_b4.append(dict(
            session=f"S{sn}",
            state=session_state[sn],
            ACA_pass=aca_p if aca_p and aca_p != "nan" else "—",
            LHA_pass=lha_p if lha_p and lha_p != "nan" else "—",
        ))
    add_df_table(doc, pd.DataFrame(rows_b4))
    add_bullet(
        doc,
        "**The S12 LHA pot-zone result (S7/S8/S9/S10) does not replicate in "
        "any other session.** The other 5 sessions show 0 LHA states passing "
        "B4 shuffle p95 (S4, S6, S8, S14) or only S6 (digging) passing in "
        "S16. Across all 6 sessions, no state passes the B4 shuffle in "
        "more than 1 session.",
    )
    add_bullet(
        doc,
        "ACA B4 also does not replicate: only S6 (digging) passes in 2/6 "
        "sessions (S8 fed, S14 fasted), and these are different sessions — "
        "no fed-vs-fasted pattern. The S0 / S8 / S9 candidates from 10b's "
        "ACA results are completely null in every session.",
    )
    add_bullet(
        doc,
        "Replication count per state is **0 or 1** for every (region, state) "
        "pair. There is no pre/post-discovery state effect that survives "
        "across multiple sessions at the conventional shuffle threshold. "
        "S12's B4 LHA pot-zone result was an outlier, not a generalisable "
        "biological signature.",
    )

    # B2 cross-session percentile
    add_para(doc, "Per-session B2 percentile in fake-discovery null (95th "
                  "percentile = significance threshold):", bold=True, size=10)
    rows_b2 = []
    for sn in sess_order:
        row = cross[cross.sn == sn].iloc[0]
        rows_b2.append(dict(
            session=f"S{sn}",
            state=session_state[sn],
            ACA_obs_pctile=f"{row['b2_aca_pctile']:.0f}",
            LHA_obs_pctile=f"{row['b2_lha_pctile']:.0f}",
        ))
    add_df_table(doc, pd.DataFrame(rows_b2))
    add_bullet(
        doc,
        "B2 percentiles range 0-94 across sessions; **none reach 95th in any "
        "session**. The B2 ACA value of 25th in S12 (real result *below* "
        "shuffle mean) is not a fluke — S4 ACA is 64th, S8 ACA is 40th, S16 "
        "ACA is 36th. No replicable pre/post within-state firing-rate "
        "signature in either region.",
    )
    add_bullet(
        doc,
        "S14 has B2 pctile = 0 in both regions, partly because of its "
        "unbalanced pre/post split (only 112 pre-discovery bins). The fake "
        "shuffles produce more balanced splits and find more sig units "
        "with their additional statistical power. This caveats S14 but the "
        "pattern (no replication) holds even excluding S14.",
    )
    add_image(doc, base_fig / "replication_heatmap_LHA.png", width_in=6.0,
              caption="LHA: per-state B4 shuffle pass across sessions (red ✓ = "
                      "passes shuffle p95). Note S12 is the only session with "
                      "any pot-zone-state passes (S7/S8/S9/S10).")
    add_image(doc, base_fig / "replication_heatmap_ACA.png", width_in=6.0,
              caption="ACA: per-state B4 shuffle pass across sessions. S6 "
                      "(digging) passes in 2 sessions; nothing else passes.")

    # ===== FED VS FASTED =====
    add_heading(doc, "Fed vs fasted aggregate", level=1)
    add_para(
        doc,
        "n=3 fed vs n=3 fasted. No formal statistical test (Mann-Whitney "
        "p_min ≈ 0.10 at this size and distributions overlap), so this is "
        "purely descriptive."
    )
    add_image(doc, base_fig / "fed_vs_fasted_aggregate_metrics.png",
              width_in=6.7,
              caption="Per-session metric values plotted by metabolic state. "
                      "No clear fed-vs-fasted dissociation visible in any of "
                      "the 8 metrics shown.")
    add_bullet(doc, "B1 per-unit pass rates: fed mean ACA 74%, LHA 43%. Fasted "
                    "mean ACA 75%, LHA 43%. Indistinguishable.")
    add_bullet(doc, "B4 LHA pass: fed = 0 states across 3 sessions; fasted = "
                    "S12 alone has 4 (S7/S8/S9/S10), S14 has none, S16 has 1 "
                    "(S6). The 'fasted has more LHA passes' narrative we "
                    "considered is driven entirely by S12 being an outlier.")
    add_bullet(doc, "Preferred-state distribution: ACA top-3 (S6, S4, S12) "
                    "consistent across both states. LHA top-3 (S2, S11, S6) "
                    "consistent across both states. The biological story "
                    "doesn't require fed-vs-fasted to be true.")

    # ===== CAVEATS =====
    add_heading(doc, "Caveats", level=1)
    add_bullet(doc, "Cross-session unit matching is NOT done. The 'preferred-"
                    "state distribution replicates' finding is at the population-"
                    "of-units level, not at the same-unit level. Tracking "
                    "individual cells across sessions would require UnitMatch.")
    add_bullet(doc, "Merged HMM state IDs are session-independent in the K=14 "
                    "fit (the dynamax model was fit on pooled data across all "
                    "7 foraging sessions), so directly comparing 'S2 in S12' "
                    "to 'S2 in S6' is valid.")
    add_bullet(doc, "Reference state for B3 GLM differs across sessions (S2 in "
                    "fasted, S3 in fed S6/S8, S11 in S4) because each session's "
                    "most-occupied state differs. This affects coefficient "
                    "interpretation but not the pass-count comparison.")
    add_bullet(doc, "S14 sample-size warning: only 112 pre-discovery bins. B2 "
                    "and B4 results in S14 should not be over-interpreted. The "
                    "headline 'no replication of S12 effect in S14' holds even "
                    "if S14 is excluded — S6 fed and S8 fed are both well-"
                    "balanced and also fail to replicate.")
    add_bullet(doc, "The 100-iteration shuffle null has Monte-Carlo noise on the "
                    "95th percentile. Marginal cases (e.g., S6 LHA at 60% "
                    "per-unit pass) might shift slightly with more iterations, "
                    "but the headline pass/fail per state is robust.")

    # ===== OUTPUTS =====
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/neural_alignment/track_B_all_sessions/cross_session_summary.csv",
         "One row per session, key metrics for B1/B2/B3/B4 + shuffle pctiles + B4 pass states."),
        ("data/HMM/neural_alignment/track_B_all_sessions/replication_count_per_state.csv",
         "Per (region, state): count of sessions where it passes B4 shuffle p95."),
        ("data/HMM/neural_alignment/track_B_all_sessions/session_{N}/...",
         "Per-session full Track B + shuffle CSVs + figures (5 sessions: S4, S6, S8, S14, S16)."),
        ("figures/HMM/neural_alignment/track_B_all_sessions/replication_heatmap_{ACA,LHA}.png",
         "Per-region states × sessions B4 shuffle pass heatmap."),
        ("figures/HMM/neural_alignment/track_B_all_sessions/preferred_state_counts_all_sessions.png",
         "Per-session B1 preferred-state distribution grid."),
        ("figures/HMM/neural_alignment/track_B_all_sessions/fed_vs_fasted_aggregate_metrics.png",
         "Per-session metric values by metabolic state."),
        ("data/HMM/neural_alignment/state_conditioned_S12/...",
         "S12 source files (loaded by 11 rather than recomputed)."),
        ("data/HMM/neural_alignment/shuffle_control_S12/...",
         "S12 10c shuffle outputs."),
        ("data/HMM/neural_alignment/shuffle_control_B2_B4_S12/...",
         "S12 10d shuffle outputs."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    out_path = REPO_ROOT / "data" / "HMM" / "neural_alignment_track_B_all_sessions_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
