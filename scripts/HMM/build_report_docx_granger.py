"""Build a Word-document report of the Granger causality analysis (script 16).

Reads outputs from data/HMM/neural_alignment/granger/ and figures.

Output: data/HMM/neural_alignment_granger_report.docx
"""
from pathlib import Path
import sys

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

sys.path.insert(0, str(Path(__file__).resolve().parent))
from _utils import load_config, REPO_ROOT


def add_heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for run in h.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def add_para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def add_bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        add_para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True
        run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def add_df_table(doc, df, float_fmt="{:.2f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        cell = table.cell(0, j)
        cell.text = str(col)
        for run in cell.paragraphs[0].runs:
            run.bold = True
            run.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            cell = table.cell(i + 1, j)
            cell.text = txt
            for run in cell.paragraphs[0].runs:
                run.font.size = Pt(8)


def code_run(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def main():
    cfg = load_config()
    base_out = REPO_ROOT / "data" / "HMM" / "neural_alignment" / "granger"
    base_fig = REPO_ROOT / "figures" / "HMM" / "neural_alignment" / "granger"

    cross = pd.read_csv(base_out / "cross_session_summary.csv")
    asym = pd.read_csv(base_out / "cross_session_asymmetry.csv")
    rep = pd.read_csv(base_out / "replication_summary.csv")
    sign = pd.read_csv(base_out / "sign_test.csv")

    sess_order = sorted(cross["session"].unique().astype(int))
    sess_state_map = {4: "fed", 6: "fed", 8: "fed",
                       12: "fasted", 14: "fasted", 16: "fasted"}

    # Per-session lag table (one row per session, both signals)
    lag_rows = []
    for sn in sess_order:
        for sig in ("pop_sum", "pc1"):
            sub = cross[(cross.session == sn) & (cross.signal == sig)
                          & (cross.direction == "ACA->LHA")]
            if len(sub):
                lag_rows.append(dict(session=sn,
                                       state=sess_state_map.get(sn, "?"),
                                       signal=sig,
                                       lag_p_bins=int(sub.iloc[0]["lag_p"]),
                                       lag_ms=int(sub.iloc[0]["lag_p"]) * 50))
    lag_df = pd.DataFrame(lag_rows)

    # ===== Build doc =====
    doc = Document()
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(1.0)
        section.right_margin = Inches(1.0)

    title = doc.add_heading(
        "Granger Causality — ACA-LHA Directionality at S3 Home-Exit Transitions",
        level=0,
    )
    for run in title.runs:
        run.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub.add_run(
        "Tests which region's activity Granger-predicts the other's at S3 "
        "(home) → exit transitions, where script 15 found multi-metric pre-"
        "exit convergence."
    )
    r.italic = True
    r.font.size = Pt(11)
    sub2 = doc.add_paragraph()
    sub2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = sub2.add_run("Report generated 2026-05-08  •  Script: scripts/HMM/16_granger_aca_lha.py")
    r.font.size = Pt(9)
    r.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    doc.add_paragraph()

    # ===== EXECUTIVE SUMMARY =====
    add_heading(doc, "Executive Summary", level=1)
    add_para(
        doc,
        "Bivariate Granger causality between ACA and LHA on 50 ms-binned "
        "population-level signals, computed over per-S3-run analysis "
        "segments (S3 stay + 5 s post-exit). Two signal types: pop_sum "
        "(z-scored summed firing rate) and PC1 (top PC of session-z-scored "
        "firing rates within region). Lag p selected by BIC (100-350 ms "
        "past). 100 circular-shift Viterbi shuffles of the predictor "
        "region's full-session signal generate a session-specific F-null."
    )
    add_para(doc, "Three plain-language findings:", bold=True, size=11)
    add_bullet(
        doc,
        "**ACA Granger-leads LHA on the population PC1 axis in 6/6 "
        "sessions.** F values for ACA→LHA reach up to 35 (S8, S12) vs "
        "LHA→ACA F up to 10. Sign test across sessions: 6/6 ACA-leading on "
        "PC1, binomial p = 0.031.",
    )
    add_bullet(
        doc,
        "**The bulk-rate (pop_sum) test is weaker.** 5/6 ACA→LHA passes "
        "shuffle p95, 4/6 LHA→ACA passes; sign test 4/6 ACA-leading "
        "(binomial p = 0.69, NS). Two fed sessions (S4, S6) show LHA-"
        "leading or symmetric on bulk rate. The bulk-rate scalar collapses "
        "the population-structure asymmetry that PC1 captures cleanly.",
    )
    add_bullet(
        doc,
        "**Lag scale ~100-350 ms.** BIC selects p = 2-7 lags at 50 ms each, "
        "implying ACA's pre-exit signal precedes LHA's on the order of "
        "0.1-0.35 s on the population-mode axis. Combined with script 14's "
        "ACA-broad-state pre-exit signal and script 15's S3-multi-metric "
        "convergence: **ACA initiates the home→foraging decision and LHA "
        "echoes it in 100-350 ms.**",
    )

    # ===== METHOD =====
    add_heading(doc, "Method", level=1)
    add_bullet(doc, "50 ms neural binning (re-binned from spike trains via "
                    "histogram on 50 ms edges).")
    add_bullet(doc, "Per-session signals computed:")
    add_bullet(doc, "  • pop_sum_aca/lha = ∑ unit counts per 50 ms bin, "
                    "z-scored within session.")
    add_bullet(doc, "  • PC1 = first principal component of session-z-scored "
                    "firing rates per region (full-session PCA, top component).")
    add_bullet(doc, "Analysis segments: each S3 (home) run from merged Viterbi "
                    "+ 5 s post-exit window. Skip runs with S3 portion < 5 s. "
                    "Per session: 6-15 segments, S3 lengths 8.7-22.3 s.")
    add_bullet(doc, "Granger F-test on pooled per-segment design matrices (no "
                    "lags spanning segment boundaries). Restricted model: y_t "
                    "regressed on its own past lags; unrestricted: + predictor "
                    "region's past lags. F = ((RSS_r − RSS_u) / p) / (RSS_u / "
                    "(n − 2p − 1)).")
    add_bullet(doc, "Lag selection: BIC across p ∈ [1, 20] on the bivariate "
                    "unrestricted model. The selected p is used for both "
                    "directions to keep them comparable.")
    add_bullet(doc, "Shuffle: 100 iterations per session × signal × "
                    "direction. Each iteration circularly shifts the "
                    "predictor region's FULL-session time series (offset ∈ "
                    "[200, T−200] bins, 50 ms each), re-extracts segments, "
                    "and recomputes F. Pass = observed F exceeds shuffle 95th "
                    "percentile.")
    add_bullet(doc, "Two replication tests reported: (a) shuffle p95 pass, "
                    "(b) analytical F-distribution p < 0.05. Both agree in "
                    "almost every session × signal × direction.")

    # ===== SEGMENTS =====
    add_heading(doc, "Analysis segments per session", level=1)
    seg_rows = []
    for sn in sess_order:
        seg_path = base_out / f"session_{sn}" / f"session_{sn}_segments.csv"
        if seg_path.exists():
            df_seg = pd.read_csv(seg_path)
            seg_rows.append(dict(
                session=sn,
                state=sess_state_map.get(sn, "?"),
                n_segments=len(df_seg),
                mean_S3_s=f"{df_seg['s3_len_bins'].mean()*0.05:.1f}",
                mean_segment_s=f"{df_seg['segment_len'].mean()*0.05:.1f}",
            ))
    seg_disp = pd.DataFrame(seg_rows)
    add_df_table(doc, seg_disp)
    add_para(doc, "S14 has the fewest segments (6) due to its very early "
                  "discovery time pulling most of the session out of the home "
                  "state into post-discovery feeding.", size=10)

    # ===== LAG SELECTION =====
    add_heading(doc, "Lag selection (BIC)", level=1)
    add_para(doc, "BIC-selected lag p per session × signal:", size=10)
    add_df_table(doc, lag_df)
    add_bullet(doc, "Lag values are uniformly small (p = 2-7). At 50 ms bin "
                    "size, this means the relevant past for predicting one "
                    "region from the other is the most recent 100-350 ms.")
    add_bullet(doc, "PC1 lags tend to be slightly longer than pop_sum lags "
                    "(S4 fed PC1 = 7 = 350 ms; S12 fasted PC1 = 2 = 100 ms), "
                    "consistent with PC1 capturing slower, more sustained "
                    "population dynamics.")

    # ===== CROSS-SESSION RESULTS =====
    doc.add_page_break()
    add_heading(doc, "Cross-session replication", level=1)
    add_para(doc, "Sessions passing shuffle p95 and F-distribution p < 0.05:",
             size=10)
    rep_disp = rep.copy()
    rep_disp.columns = ["signal", "direction", "n_sessions",
                          "passing shuf p95", "passing F p<0.05"]
    add_df_table(doc, rep_disp)
    add_image(doc, base_fig / "replication_summary.png", width_in=5.5,
              caption="Per signal × direction count of sessions passing each "
                      "test (out of 6). Red bars = shuffle p95 pass; blue bars "
                      "= F-distribution p < 0.05.")

    add_heading(doc, "Per-session F values and asymmetry", level=2)
    asym_disp = asym.copy()
    asym_disp["F_ACA_to_LHA"] = asym_disp["F_ACA_to_LHA"].round(2)
    asym_disp["F_LHA_to_ACA"] = asym_disp["F_LHA_to_ACA"].round(2)
    asym_disp["asymmetry_index"] = asym_disp["asymmetry_index"].round(2)
    asym_disp = asym_disp[["session", "state", "signal",
                             "F_ACA_to_LHA", "F_LHA_to_ACA",
                             "asymmetry_index"]]
    asym_disp.columns = ["session", "state", "signal",
                           "F_ACA→LHA", "F_LHA→ACA", "asymmetry"]
    add_df_table(doc, asym_disp)
    add_bullet(
        doc,
        "Asymmetry index = (F_ACA→LHA − F_LHA→ACA) / sum. Positive = ACA "
        "leads. PC1 asymmetry is positive in **6/6 sessions** (range +0.19 "
        "to +0.78). pop_sum asymmetry is positive in 4/6 (range −0.23 to "
        "+0.71); the two fed sessions S4 and S6 show negative or near-zero "
        "pop_sum asymmetry but strongly positive PC1 asymmetry — the bulk-"
        "rate test does not capture the population-mode directionality "
        "that PC1 does.",
    )
    add_bullet(
        doc,
        "F values for ACA→LHA on PC1 reach 35 (S8 fed, S12 fasted), 10.8 "
        "(S6 fed), 4-9 (S4, S14, S16). LHA→ACA F values are systematically "
        "smaller (max 10.7 in S8, mostly 1-7).",
    )
    add_image(doc, base_fig / "asymmetry_per_session.png", width_in=6.7,
              caption="Asymmetry index per session, fed (blue) vs fasted "
                      "(red). Left panel = pop_sum, right = PC1. PC1 has all "
                      "6 sessions positive; pop_sum has 4/6 positive.")

    add_heading(doc, "Sign test", level=2)
    add_para(doc, "Two-sided binomial test on count of ACA-leading sessions:")
    sign_disp = sign.copy()
    sign_disp.columns = ["signal", "n_sessions", "n_ACA_leads", "n_LHA_leads",
                           "binomial p"]
    add_df_table(doc, sign_disp, float_fmt="{:.4f}")
    add_bullet(doc, "**PC1: 6/6 sessions ACA leads, binomial p = 0.031** — "
                    "the cleanest cross-session-replicating directional "
                    "signal in the dataset.")
    add_bullet(doc, "pop_sum: 4/6, binom p = 0.69 — not significant. The two "
                    "fed sessions S4 and S6 with reversed pop_sum asymmetry "
                    "drag this down.")

    add_image(doc, base_fig / "F_comparison_pop_vs_pc1.png", width_in=6.7,
              caption="Granger F: pop_sum vs PC1, per session × direction. "
                      "PC1 systematically gives larger F values for "
                      "ACA→LHA, consistent with PC1 capturing population-"
                      "mode dynamics that the bulk-rate scalar averages away.")

    # ===== INTERPRETATION =====
    doc.add_page_break()
    add_heading(doc, "Interpretation", level=1)
    add_bullet(
        doc,
        "**ACA→LHA on PC1 is the cleanest directional signal.** All 6 "
        "sessions show positive asymmetry (ACA leads). Combined with the "
        "consistent ACA pre-exit firing-rate signal (script 14, 4-6/6 "
        "sessions across most states) and the multi-metric convergence at "
        "S3 (script 15: M2/M3/M5 all 5/5 in ACA), this is the strongest "
        "evidence so far that ACA is upstream of LHA on the home → "
        "foraging-bout decision axis.",
    )
    add_bullet(
        doc,
        "**Lag is fast: 100-350 ms.** The PC1 asymmetry suggests ACA's "
        "population mode shifts ~100-350 ms before LHA's matching shift. "
        "This is on the order of typical cortico-subcortical relay times "
        "and is consistent with ACA driving an LHA response (rather than "
        "the two regions reading a common upstream input at zero lag).",
    )
    add_bullet(
        doc,
        "**Why does pop_sum disagree on S4 and S6?** Bulk firing rate "
        "averages across the population, so a small subset of cells "
        "leading in one direction can be masked by a larger subset "
        "leading in the other. PC1 isolates the dominant covariation "
        "mode, which is what carries the directional information. The "
        "reversal on S4 and S6 is also consistent with the metabolic-"
        "state observation that LHA's overall feeding-state encoding "
        "dominates fed sessions (Track B B1 — 33% of LHA units prefer "
        "S2 in S12); the bulk LHA rate is high and changes more, "
        "creating apparent LHA-leading-pop_sum.",
    )
    add_bullet(
        doc,
        "**Three caveats limit the strength of the causal inference:** "
        "(1) Granger assumes within-segment stationarity, which is "
        "violated when segments span the home → foraging transition. "
        "(2) Linear Granger only — nonlinear coupling that operates "
        "through changes in correlation structure (script 15 M5) "
        "would be missed. (3) Common-input confound: a third upstream "
        "region driving both ACA and LHA at slightly different lags "
        "would produce the same Granger asymmetry. The result is best "
        "stated as 'ACA's PC1 dynamics consistently precede LHA's at "
        "S3 home transitions' rather than 'ACA causes LHA'.",
    )

    # ===== Per-session figures =====
    doc.add_page_break()
    add_heading(doc, "Per-session F distributions vs shuffle null", level=1)
    add_para(doc, "Each panel shows one session: 4-panel grid of pop_sum and "
                  "PC1 × ACA→LHA and LHA→ACA. Red line = observed F, "
                  "dashed black = shuffle 95th percentile, histogram = 100 "
                  "circular-shift shuffle null.", size=10)
    for sn in sess_order:
        path = base_fig / f"session_{sn}" / "observed_F_per_direction.png"
        if path.exists():
            add_image(doc, path, width_in=6.7,
                      caption=f"Session {sn} ({sess_state_map.get(sn,'?')}).")

    add_heading(doc, "Per-session segment overlay on Viterbi timeline",
                 level=1)
    add_para(doc, "Viterbi state sequence (color) with green-shaded S3 stay "
                  "windows and orange-shaded post-exit 5 s windows used as "
                  "Granger analysis segments.", size=10)
    for sn in sess_order:
        path = base_fig / f"session_{sn}" / "segments_overlay.png"
        if path.exists():
            add_image(doc, path, width_in=6.7,
                      caption=f"Session {sn} segments.")

    # ===== CAVEATS =====
    add_heading(doc, "Caveats", level=1)
    add_bullet(doc, "Granger assumes within-segment stationarity. The home "
                    "→ post-exit segment explicitly spans a behavioral "
                    "transition where the dynamics are NON-stationary. "
                    "Result is best interpreted as 'directional "
                    "predictability' rather than mechanistic causation.")
    add_bullet(doc, "Linear Granger only. Nonlinear coupling via "
                    "correlation-structure or higher-order joint statistics "
                    "(c.f. script 15 M5) would not be detected.")
    add_bullet(doc, "Common-input confound: a third region driving both ACA "
                    "and LHA at unequal lags can mimic ACA→LHA Granger. With "
                    "only ACA-LHA recordings we cannot rule this out.")
    add_bullet(doc, "BIC selects small lags (p = 2-7) consistently. A "
                    "longer-lag analysis (p = 10, 15, 20) would test "
                    "robustness of the asymmetry to model order; not done in "
                    "this pass.")
    add_bullet(doc, "S14 has only 6 segments due to early discovery; its F "
                    "values may be noisier than other sessions. The result "
                    "holds without S14 (5/5 ACA-leading on PC1, binomial p = "
                    "0.062 — marginal).")
    add_bullet(doc, "50 ms binning is fine but bursting events may be "
                    "coarse-grained. Sub-bin spike-time Granger (e.g., point-"
                    "process generalizations) could be a follow-up.")
    add_bullet(doc, "Single-mouse, 6 sessions, 1 paradigm. The headline ACA-"
                    "leads-LHA result needs replication in additional mice.")

    # ===== OUTPUTS =====
    add_heading(doc, "Output Files", level=1)
    outputs = [
        ("data/HMM/neural_alignment/granger/cross_session_summary.csv",
         "Per (session, signal, direction): observed F, F p-value, lag, shuffle stats."),
        ("data/HMM/neural_alignment/granger/cross_session_asymmetry.csv",
         "Per (session, signal): F_ACA→LHA, F_LHA→ACA, asymmetry index."),
        ("data/HMM/neural_alignment/granger/replication_summary.csv",
         "Per (signal, direction): n sessions passing shuffle p95 and analytical F."),
        ("data/HMM/neural_alignment/granger/sign_test.csv",
         "Sign test on ACA-leads-LHA across sessions per signal."),
        ("data/HMM/neural_alignment/granger/session_{N}/session_{N}_segments.csv",
         "Per-session S3+post-exit segment definitions."),
        ("data/HMM/neural_alignment/granger/session_{N}/session_{N}_signals.npz",
         "Per-session 50 ms binned signals (pop_sum, PC1-3) per region."),
        ("data/HMM/neural_alignment/granger/session_{N}/session_{N}_observed.csv",
         "Per (signal, direction) observed F + F p-value + shuffle stats."),
        ("data/HMM/neural_alignment/granger/session_{N}/session_{N}_shuffle.csv",
         "Per-iteration shuffle F values (100 per direction × signal)."),
        ("figures/HMM/neural_alignment/granger/asymmetry_per_session.png",
         "Asymmetry index per session, fed vs fasted color-coded."),
        ("figures/HMM/neural_alignment/granger/replication_summary.png",
         "Bar chart of replication counts per signal × direction."),
        ("figures/HMM/neural_alignment/granger/F_comparison_pop_vs_pc1.png",
         "Per-session pop_sum F vs PC1 F scatter, for both directions."),
        ("figures/HMM/neural_alignment/granger/session_{N}/observed_F_per_direction.png",
         "Per-session 2x2 panel of observed F vs shuffle null."),
        ("figures/HMM/neural_alignment/granger/session_{N}/segments_overlay.png",
         "Per-session Viterbi timeline with analysis segments highlighted."),
    ]
    for path, descr in outputs:
        p = doc.add_paragraph(style="List Bullet")
        r = p.add_run(path)
        r.font.name = "Consolas"
        r.font.size = Pt(9)
        p.add_run(" — " + descr).font.size = Pt(10)

    out_path = REPO_ROOT / "data" / "HMM" / "neural_alignment_granger_report.docx"
    doc.save(out_path)
    print(f"Saved {out_path}")


if __name__ == "__main__":
    main()
