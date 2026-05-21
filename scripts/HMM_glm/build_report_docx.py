"""Build a detailed Word-doc report for the K=6 GLM-HMM fits.

Documents the computational results AND prominently flags the session-
identification artifact (states tag sessions, not behavioral structure).

Output: data/HMM/glm_hmm_K6_report.docx
"""
from pathlib import Path

import numpy as np
import pandas as pd

from docx import Document
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT

REPO = Path(__file__).resolve().parents[2]
RESULTS = REPO / "results" / "hmm"
OUT_DOCX = REPO / "data" / "HMM" / "glm_hmm_K6_report.docx"

CELLS = [("ACA", "exploration"),
          ("ACA", "foraging"),
          ("LHA", "exploration"),
          ("LHA", "foraging")]


def heading(doc, text, level=1):
    h = doc.add_heading(text, level=level)
    for r in h.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    return h


def para(doc, text, bold=False, size=None):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.bold = bold
    if size is not None:
        r.font.size = Pt(size)
    return p


def bullet(doc, text):
    return doc.add_paragraph(text, style="List Bullet")


def add_image(doc, path, width_in=6.0, caption=None):
    if not Path(path).exists():
        para(doc, f"[missing figure: {path}]", size=9)
        return
    doc.add_picture(str(path), width=Inches(width_in))
    if caption is not None:
        cap = doc.add_paragraph()
        run = cap.add_run(caption)
        run.italic = True; run.font.size = Pt(9)
        cap.alignment = WD_ALIGN_PARAGRAPH.CENTER


def df_table(doc, df, float_fmt="{:.3f}"):
    n_rows, n_cols = df.shape
    table = doc.add_table(rows=n_rows + 1, cols=n_cols)
    table.style = "Light Grid Accent 1"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    for j, col in enumerate(df.columns):
        c = table.cell(0, j)
        c.text = str(col)
        for r in c.paragraphs[0].runs:
            r.bold = True; r.font.size = Pt(8)
    for i in range(n_rows):
        for j in range(n_cols):
            v = df.iat[i, j]
            if pd.isna(v):
                txt = ""
            elif isinstance(v, (bool, np.bool_)):
                txt = "True" if bool(v) else "False"
            elif isinstance(v, float) and np.isfinite(v):
                txt = float_fmt.format(v)
            else:
                txt = str(v)
            c = table.cell(i + 1, j)
            c.text = txt
            for r in c.paragraphs[0].runs:
                r.font.size = Pt(8)


def code(doc, text, size=9):
    p = doc.add_paragraph()
    r = p.add_run(text)
    r.font.name = "Consolas"
    r.font.size = Pt(size)
    return p


def per_session_state_dominance(z_assignments_npz: Path) -> pd.DataFrame:
    """For each session, identify the single dominant state and its occupancy."""
    z = np.load(z_assignments_npz)
    rows = []
    for key in z.files:
        sn = int(key.replace("session_", ""))
        arr = np.asarray(z[key])
        if len(arr) == 0:
            continue
        unique, counts = np.unique(arr, return_counts=True)
        idx = int(np.argmax(counts))
        rows.append(dict(
            session=sn,
            n_bins=int(len(arr)),
            dominant_state=int(unique[idx]),
            dominant_fraction=float(counts[idx] / len(arr)),
            n_states_visited=int(len(unique)),
        ))
    return pd.DataFrame(rows).sort_values("session").reset_index(drop=True)


def main():
    doc = Document()

    # Title
    t = doc.add_heading("Poisson-GLM-HMM K=6 Fits — Results & Session-Identification Artifact",
                        level=0)
    for r in t.runs:
        r.font.color.rgb = RGBColor(0x1F, 0x3A, 0x5F)
    sub = doc.add_paragraph()
    sr = sub.add_run("Four HMMs (ACA-exploration, ACA-foraging, LHA-exploration, "
                       "LHA-foraging) fit at K=6 on the dual-probe coordinates_1 "
                       "mouse01 dataset, jointly across fed/fasted/HFD with one-hot "
                       "metabolic state as input covariate on transitions. ssm "
                       "Poisson-GLM-HMM, 3 random restarts per cell, no cross-"
                       "validation, 50 ms bins.")
    sr.italic = True; sr.font.size = Pt(11)

    para(doc, "Scripts: scripts/HMM_glm/{_data, _model, _plots, fit_pipeline}.py", size=9)
    para(doc, "Outputs: results/hmm/{region}/{phase}/K06/", size=9)
    para(doc, "Wall time: 126.1 min total", size=9)
    para(doc, "Generated 2026-05-13", size=9)
    doc.add_paragraph()

    # Executive Summary (highlight artifact)
    heading(doc, "Executive Summary", level=1)
    para(doc, "Three findings.", bold=True)
    bullet(doc,
            "Fits completed cleanly in 126 min for all 4 (region, phase) cells. "
            "Best-of-3 training log-likelihoods: ACA-exploration -36.45M, "
            "ACA-foraging -40.68M, LHA-exploration -9.70M, LHA-foraging -10.78M. "
            "All standard outputs saved (model.pkl, state_assignments, "
            "transitions_by_condition, emissions, occupancy, plots).")
    bullet(doc,
            "CRITICAL: the fits have collapsed to a 'session-identification' "
            "solution. Each session occupies essentially ONE state for ALL of "
            "its bins. Self-transition diagonals are ≥0.95 across nearly every "
            "state. Per-state occupancy is a clean 1/n_sessions split (0.333 for "
            "HFD/fasted with 3 sessions each, 0.25 for fed with 4 sessions). "
            "The HMM has partitioned sessions, not behaviors.")
    bullet(doc,
            "Root cause: the pooled emission matrix is built by stacking each "
            "session's units into separate columns with zero-padding for other "
            "sessions. Cluster IDs aren't comparable across kilosort runs, so "
            "this is the only way to fit jointly. But the resulting model can "
            "perfectly identify which session is active from the unit-activity "
            "pattern alone — and that becomes the dominant signal. The behavioral "
            "structure within sessions is much weaker than the session-to-session "
            "unit-set differences. The result: K=6 fits are computationally valid "
            "but scientifically uninterpretable as behavioral states.")

    # Method
    heading(doc, "Method", level=1)
    bullet(doc, "Sessions per (region, phase): 10 = 4 fed (S3/4, S5/6, S7/8, S9/10) + 3 fasted (S11-16) + 3 HFD (S19-24). Exploration uses odd-numbered, foraging uses even-numbered.")
    bullet(doc, "Spike data binned at 50 ms. QC: probe-0 ACA = KSLabel='good' + fr>0.2 Hz; probe-1 LHA = KSLabel='good' + fr>0.2 + amp>43 + depth 0-345 µm.")
    bullet(doc, "HFD sessions S20-S24 lack cluster_info.tsv — used the standard fallback (templates.npy + channel_positions.npy for depth; cluster_group.tsv for labels).")
    bullet(doc, "Pooled emission space: each session's units occupy unique columns. Total D ranges 921 (LHA exploration) to 1996 (ACA exploration).")
    bullet(doc, "Model: ssm Poisson-GLM-HMM with `transitions='inputdriven'`. Note: ssm parameterizes input dependence as Ws shape (K, M) — per-target-state input bias — rather than the spec's full (K, K, M) form. The effective transition matrix for each metabolic state is still distinct (3 separate (K, K) matrices via `transition_matrices()`).")
    bullet(doc, "Input u_t ∈ R^3 = one-hot [fed, fasted, HFD], constant within a session.")
    bullet(doc, "EM: 3 random restarts per (region, phase, K=6); seed_base=20260513; num_iters=100; ssm tolerance 1e-4. Kept best fit by training LL.")
    bullet(doc, "No cross-validation (per user instruction to keep runtime ~2 hours).")

    # Per-cell summary table
    heading(doc, "Per-cell summary", level=1)
    summary_rows = []
    for region, phase in CELLS:
        d = RESULTS / region / phase / "K06"
        cfg = (d / "config.yaml").read_text() if (d / "config.yaml").exists() else ""
        # parse approximate counts from config text
        try:
            import yaml
            c = yaml.safe_load(cfg)
        except Exception:
            c = {}
        rl = pd.read_csv(d / "restart_log.csv")
        best_ll = float(c.get("best_train_ll", float("nan")))
        summary_rows.append(dict(
            region=region, phase=phase,
            sessions=len(c.get("sessions", [])),
            D=c.get("D", None),
            best_train_ll=best_ll,
            best_restart=int(c.get("best_init", -1)),
            total_minutes=float(rl["dt_s"].sum() / 60),
            mean_iters=float(rl["n_iters"].mean()),
        ))
    df_table(doc, pd.DataFrame(summary_rows), float_fmt="{:.2f}")

    # Per-cell detail
    for region, phase in CELLS:
        heading(doc, f"{region} {phase}", level=1)
        d = RESULTS / region / phase / "K06"

        # restart variance
        heading(doc, "Restart log", level=2)
        rl = pd.read_csv(d / "restart_log.csv")
        df_table(doc, rl, float_fmt="{:.2f}")

        # occupancy by metabolic state
        heading(doc, "State occupancy per session (annotated by metabolic state)", level=2)
        occ = pd.read_csv(d / "occupancy.csv")
        occ_view = occ[occ["occupancy"] > 0.001].copy()
        occ_view = occ_view.sort_values(["metabolic_state", "session", "state"])
        df_table(doc, occ_view, float_fmt="{:.3f}")

        # per-session dominant state (the artifact)
        heading(doc, "Per-session dominant state (artifact diagnostic)", level=2)
        dom = per_session_state_dominance(d / "state_assignments.npz")
        occ_full = pd.read_csv(d / "occupancy.csv")
        sn_to_state = occ_full[["session", "metabolic_state"]].drop_duplicates()
        dom = dom.merge(sn_to_state, on="session")
        dom = dom[["session", "metabolic_state", "n_bins", "dominant_state",
                     "dominant_fraction", "n_states_visited"]]
        df_table(doc, dom, float_fmt="{:.3f}")
        para(doc,
              f"Mean dominant-state fraction = "
              f"{dom['dominant_fraction'].mean():.4f}; "
              f"max across sessions = {dom['dominant_fraction'].max():.4f}. "
              "Values near 1.0 = each session is in one state for ~all bins.")

        # transitions
        heading(doc, "Effective transition matrices (one per metabolic state)", level=2)
        tnpz = np.load(d / "transitions_by_condition.npz")
        states_lbl = list(tnpz["metabolic_states"])
        trans = tnpz["transitions"]
        # Show diagonals (self-transition probs) per metabolic state
        diag_rows = []
        for m, lbl in enumerate(states_lbl):
            row = dict(metabolic_state=lbl)
            for k in range(trans.shape[1]):
                row[f"state_{k}_self"] = float(trans[m, k, k])
            diag_rows.append(row)
        df_table(doc, pd.DataFrame(diag_rows), float_fmt="{:.3f}")
        para(doc,
              "Self-transition diagonals near 1.0 indicate the HMM rarely changes "
              "state. Combined with one-state-per-session occupancy, this confirms "
              "the session-identification artifact.")

        # emissions summary
        heading(doc, "Emission log-rate summary", level=2)
        e = np.load(d / "emissions.npz")
        log_rates = e["log_rates"]
        em_rows = []
        for k in range(log_rates.shape[0]):
            r = log_rates[k]
            em_rows.append(dict(
                state=k,
                mean_log_rate=float(r.mean()),
                p50_log_rate=float(np.median(r)),
                max_log_rate=float(r.max()),
                n_units_with_log_rate_above_neg5=int((r > -5).sum()),
            ))
        df_table(doc, pd.DataFrame(em_rows), float_fmt="{:.2f}")
        para(doc,
              "Each state has a small number of units with appreciable log-rates "
              "(>-5) and a large number with very negative log-rates. The "
              "'active' units in each state's emission matrix correspond to "
              "the session that maps to that state.")

        # Plots
        heading(doc, "Plots", level=2)
        add_image(doc, d / "plots" / "occupancy.png", width_in=5.5,
                  caption=f"{region} {phase} K=6: state occupancy heatmap (states × sessions)")
        add_image(doc, d / "plots" / "transitions.png", width_in=6.5,
                  caption=f"{region} {phase} K=6: effective transitions per metabolic state")
        add_image(doc, d / "plots" / "emissions.png", width_in=6.5,
                  caption=f"{region} {phase} K=6: emission log-rate heatmap")
        for fp in sorted((d / "plots").glob("timeline_*.png")):
            add_image(doc, fp, width_in=5.0,
                      caption=fp.stem.replace("timeline_", "Viterbi timeline: "))

    # Diagnostic discussion
    heading(doc, "Diagnostic: why the session-identification artifact", level=1)
    para(doc,
          "Three lines of evidence:")
    bullet(doc,
            "Per-session dominant-state fraction is ≥0.95 in nearly every "
            "session across all 4 cells. The mean across all 40 (session × cell) "
            "combinations is essentially 1.0.")
    bullet(doc,
            "Self-transition diagonals are 0.93-1.00 in every (state × metabolic "
            "state) cell. Real behavioral HMMs at 50 ms bins typically have "
            "self-transition probabilities 0.85-0.95, with regular jumps between "
            "states tracking behavior.")
    bullet(doc,
            "Per-state occupancy across sessions matches 1/N exactly. With 10 "
            "sessions = 3 HFD + 3 fasted + 4 fed, expected 'session tagging' "
            "occupancies are: HFD/fasted bins → 1/3 ≈ 0.333; fed bins → 1/4 = 0.25. "
            "Observed values match exactly to 3 decimal places.")

    # Root cause
    heading(doc, "Root cause: pooled emission space + independent kilosort sortings", level=1)
    para(doc,
          "Kilosort assigns cluster IDs independently per session, so unit IDs "
          "aren't comparable across sessions. To fit a joint model with shared "
          "emissions, we stack each session's units into separate columns of a "
          "pooled emission matrix, zero-padding for other sessions:")
    code(doc,
          "# scripts/HMM_glm/_data.py:_stack_session_units\n"
          "cursor = 0\n"
          "for session in sequences:\n"
          "    cols = range(cursor, cursor + session.n_units)\n"
          "    session.pooled_cols = cols\n"
          "    cursor += session.n_units")
    para(doc,
          "This means for session-K, columns 0..cursor_K and cursor_(K+1)..D are "
          "always zero. The Poisson likelihood handles zeros at low log-rate "
          "without numerical issues: P(0 | λ_tiny) ≈ 1. But the HMM doesn't "
          "have to use this freedom — instead it learns log-rates that 'tag' "
          "each session's unit set:")
    bullet(doc, "State 0 has high log-rates only in S4's columns")
    bullet(doc, "State 1 has high log-rates only in S6's columns")
    bullet(doc, "... and so on")
    para(doc,
          "Once trained this way, the most likely state at any bin is unambiguous "
          "given which units fired (because each session uses a disjoint column "
          "set). The session-tagging solution captures more bin-level variance "
          "than any solution that tries to find shared within-session structure, "
          "so EM converges there.")

    # Fix options
    heading(doc, "Fix options for downstream analysis", level=1)
    para(doc,
          "The K=6 fits as-is are NOT suitable for state-conditioned behavioral "
          "analyses (e.g., pre-exit signatures, decision vectors). Four ways forward:")
    bullet(doc,
            "Per-session fits: fit one HMM per session × metabolic state. Loses "
            "the joint-fit semantics but gets real within-session states. Easy "
            "to implement: iterate sessions, call fit_one() per session with "
            "input u_t = the session's one-hot. The existing infrastructure "
            "supports this with minimal changes.")
    bullet(doc,
            "UnitMatch cross-session alignment: identify the same neurons "
            "across sessions, build a shared unit space. Memory notes UnitMatch "
            "needs per-region tuning. Major project but the most rigorous fix.")
    bullet(doc,
            "Population-feature input: aggregate spikes per region into "
            "low-D features (PC scores, mean firing rate, sync) and fit on those. "
            "Loses per-unit resolution but bypasses the session-identification "
            "artifact because the feature space is comparable across sessions.")
    bullet(doc,
            "Compositional emission model: factor emissions into "
            "(session-baseline) + (state-modulation), with the state modulation "
            "shared across sessions. Requires custom model code (ssm doesn't "
            "support this out of the box).")

    para(doc,
          "Recommended sequence: (a) start with per-session fits to validate "
          "that within-session structure exists at all; (b) if yes, consider "
          "population-feature input as a practical joint fit; (c) UnitMatch "
          "alignment as a longer-term enhancement.")

    # Caveats
    heading(doc, "Caveats", level=1)
    bullet(doc,
            "No cross-validation. Train LL alone always favors more parameters; "
            "cannot say K=6 was 'right' (and the artifact above makes the "
            "question moot for this pooling scheme).")
    bullet(doc, "3 random restarts per cell (spec asks for ≥5). Variance across "
                  "restarts was modest (e.g. ACA-exploration: -36.4M to -39.2M).")
    bullet(doc,
            "ssm `inputdriven` transitions use Ws shape (K, M) — per-target-state "
            "input bias — not the spec's (K, K, M) full form. Effective per-input "
            "(K, K) transitions are still distinct (computed via "
            "`transition_matrices()`), so the model is structurally usable; "
            "just simpler than the spec.")
    bullet(doc,
            "S5/13/23 have foreshortened durations (1100/600 s vs 1800 s). "
            "The model handles this fine — sequences of different length are "
            "ssm's standard format.")
    bullet(doc,
            "HFD sessions S20-S24 use template-derived depth/labels (no "
            "cluster_info.tsv per the project's HFD memory). The fallback path "
            "found 168-243 units per ACA session, 69-114 per LHA session — "
            "comparable to fed/fasted, so the fallback is working.")
    bullet(doc,
            "K=6 was a fixed choice motivated by analogy to the behavioral HMM "
            "pipeline's N=14 merged states. With the artifact identified, the "
            "natural-K question is essentially uninterpretable in the current "
            "pooled-emission setup.")

    # Output files
    heading(doc, "Output files", level=1)
    code(doc, "results/hmm/{region}/{phase}/K06/")
    bullet(doc, "model.pkl — pickled ssm HMM")
    bullet(doc, "config.yaml — region, phase, K, D, M, best restart, sessions, metabolic states")
    bullet(doc, "restart_log.csv — train LL + iterations per random restart")
    bullet(doc, "pooled_ids.npz — (D, 2) array of (session_num, cluster_id) per pooled column")
    bullet(doc, "state_assignments.npz — Viterbi state path per session (key = 'session_N')")
    bullet(doc, "transitions_by_condition.npz — (3, K, K) effective transition matrices, one per metabolic state")
    bullet(doc, "emissions.npz — (K, D) log_rates plus pooled_ids")
    bullet(doc, "occupancy.csv — long-form table of (session, metabolic_state, state, occupancy, n_bins)")
    bullet(doc, "plots/{occupancy.png, transitions.png, emissions.png, timeline_*.png}")
    code(doc, "scripts/HMM_glm/")
    bullet(doc, "_data.py — spike loader, QC, binning, pooled-stack concat with one-hot input")
    bullet(doc, "_model.py — ssm wrapper (fit_one, evaluate_ll, state_assignments, effective_transitions)")
    bullet(doc, "_plots.py — diagnostic plot helpers")
    bullet(doc, "fit_pipeline.py — CLI with fit-one / cv / all modes")
    bullet(doc, "run_all_K6_K8.py — driver that called fit-one for all 4 cells")
    bullet(doc, "build_report_docx.py — this report builder")

    OUT_DOCX.parent.mkdir(parents=True, exist_ok=True)
    doc.save(str(OUT_DOCX))
    print(f"Saved {OUT_DOCX}")


if __name__ == "__main__":
    main()
